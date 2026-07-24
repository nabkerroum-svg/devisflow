"""
Service de gestion des templates Word.

Trois opérations principales :
  1. analyser(fichier_docx) → liste des variables Jinja2 présentes
  2. annoter_auto(fichier_docx, mapping) → remplace les zones par des marqueurs
  3. generer_devis(template_path, data) → produit un .docx rempli
  4. docx_to_pdf(docx_path) → convertit via LibreOffice headless

Le projet est conçu pour que vous puissiez fournir DEUX types de templates :
  a) Templates DÉJÀ annotés (vous avez remplacé vous-même les zones par {{ VAR }}
     dans Word) — le système les utilise directement.
  b) Templates BRUTS (votre devis original sans modifications) — le système
     applique automatiquement les substitutions classiques (date, destinataire,
     numéro, adresse site) via une table de mapping configurable.
"""
import re
import subprocess
import shutil
import zipfile
import copy
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docxtpl import DocxTemplate, Listing

from config import TEMPLATES_DIR, GENERATED_DIR, SOFFICE_BIN


RYTHME_COPRO_REFERENCE = {
    "page_top_mm": 50,
    "body_after_pt": 3,
    "section_before_pt": 10,
    "section_after_pt": 5,
    "cgv_line_spacing": 1.08,
    "cgv_body_after_pt": 3,
    "cgv_title_before_pt": 8,
    "cgv_title_after_pt": 4,
    "cgv_cover_before_pt": 18,
    "cgv_cover_after_pt": 6,
    "bon_title_after_pt": 8,
    "bon_clause_before_pt": 2,
    "bon_clause_after_pt": 10,
    "bon_date_before_pt": 6,
    "bon_date_after_pt": 6,
    "bon_signature_before_pt": 4,
    "bon_signature_after_pt": 3,
}


# Ligne unique de départ du contenu sous le logo Marie-Eugénie pour les
# modèles bureaux_petit et ponctuels. Les pages concernées utilisent ce même
# espace avant le premier titre, sans calage page par page.
CONTENT_START_BEFORE_PT = 48
CONTENT_TITLE_AFTER_PT = 8
LOGO_FALLBACK_IMAGE_NAMES = {"perche_h2o_e03c5fe5.jpg"}
COPRO_PETITE_TEMPLATE_CODES = {"copro_petite"}
PONCTUEL_TEMPLATE_PREFIX = "ponctuel_"
COPRO_COVER_CONTACT_BLOCK_BEFORE_TWIPS = "720"
COPRO_COVER_AFTER_CONTACT_LINE_TWIPS = "160"
COPRO_PETITE_COMPANY_PRESENTATION_START_TWIPS = "40"


def _est_image_logo_entreprise(path: Any) -> bool:
    if not path:
        return False
    name = Path(str(path)).name.lower()
    if name in LOGO_FALLBACK_IMAGE_NAMES:
        return True
    return "logo" in name and ("marie" in name or "eugenie" in name or "eugénie" in name)


# ============================================================
# 0. CONVERSION .doc / .pdf → .docx (pour l'import de modèles)
# ============================================================

def convertir_en_docx(source: Path, out_dir: Optional[Path] = None) -> Path:
    """Convertit un fichier .doc (Word 97-2003) en .docx via LibreOffice headless.

    Les modèles maîtres de Marie Eugénie sont historiquement au format .doc ;
    cette fonction permet de les importer directement sans conversion manuelle.

    Returns:
        Le chemin du .docx produit.
    """
    out_dir = out_dir or source.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    result = subprocess.run(
        [SOFFICE_BIN, "--headless", "--convert-to", "docx",
         "--outdir", str(out_dir), str(source)],
        capture_output=True, text=True, timeout=120,
    )
    if result.returncode != 0:
        raise RuntimeError(f"Conversion .docx échouée : {result.stderr or result.stdout}")
    docx_path = out_dir / (source.stem + ".docx")
    if not docx_path.exists():
        raise RuntimeError(f".docx non créé : {docx_path}")
    return docx_path


# ============================================================
# 1. ANALYSE — extraire les variables présentes dans un template
# ============================================================

JINJA_VAR_PATTERN = re.compile(r"\{\{\s*([A-Z][A-Z0-9_]*)\s*\}\}")
# Variables utilisées dans les boucles : {%tr for opt in OPTIONS %} / {%p for p in PRESTATIONS %}
JINJA_LOOP_PATTERN = re.compile(r"\{%\s*(?:tr|p|r)?\s*for\s+\w+\s+in\s+([A-Z][A-Z0-9_]*)\s*%\}")


def paragraph_contains_media(paragraph) -> bool:
    """Retourne True si un paragraphe Word contient une image ou un objet graphique.

    Les anciens modèles .doc convertis en .docx utilisent souvent VML
    (`w:pict`, `v:shape`, `v:imagedata`) tandis que les modèles plus récents
    utilisent DrawingML (`w:drawing`). Ce helper protège ces deux familles sans
    faire planter le rendu si le paragraphe est vide ou incomplet.
    """
    if paragraph is None:
        return False
    element = getattr(paragraph, "_element", None)
    if element is None:
        return False
    try:
        return bool(
            element.xpath(
                ".//*[local-name()='drawing' or local-name()='pict' or local-name()='shape' or local-name()='imagedata']"
            )
        )
    except Exception:
        return False


def _remplir_tableau_financier_docx(doc, data: Dict) -> bool:
    """Remplit la proposition financiere avec une ou plusieurs lignes.

    Retourne False si aucune ligne financiere n'est fournie, pour laisser les
    anciens remplacements cellule par cellule prendre le relais.
    """
    lignes = data.get("FINANCIAL_LINES")
    if not isinstance(lignes, list) or not lignes:
        return False

    from docx.shared import Pt, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    cols_cfg = data.get("FINANCIAL_COLUMNS") or {}
    use_qty = bool(cols_cfg.get("quantite"))
    use_unit = bool(cols_cfg.get("unite"))
    use_unit_price = bool(cols_cfg.get("prix_unitaire_ht"))
    detailed = use_qty or use_unit or use_unit_price or len(lignes) > 1

    if detailed:
        columns = ["Désignation"]
        if use_qty:
            columns.append("Qté")
        if use_unit_price:
            columns.append("PU HT")
        columns += ["Total HT", "TVA", "Total TTC"]
    else:
        columns = ["Désignation", "Montant HT", "TVA", "Montant TTC"]

    def norm(txt: str) -> str:
        import unicodedata
        txt = (txt or "").lower().replace("\xa0", " ")
        txt = unicodedata.normalize("NFKD", txt)
        txt = "".join(ch for ch in txt if not unicodedata.combining(ch))
        return re.sub(r"\s+", " ", txt).strip()

    def table_score(table) -> int:
        txt = norm(" ".join(cell.text for row in table.rows for cell in row.cells))
        score = 0
        for token in ("prix", "tarif", "tva", "ttc", "ht"):
            if token in txt:
                score += 1
        if "proposition financiere" in txt:
            score += 3
        return score

    candidates = [(table_score(t), t) for t in doc.tables]
    candidates = [x for x in candidates if x[0] >= 3]
    if not candidates:
        return False
    table = max(candidates, key=lambda x: x[0])[1]

    def clear_cell(cell):
        cell.text = ""
        if not cell.paragraphs:
            cell.add_paragraph()

    def set_cell_text(cell, text, *, bold=False, italic=False, align=None, size=9):
        clear_cell(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if align is not None:
            p.alignment = align
        run = p.add_run(str(text or ""))
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic

    def set_cell_borders(cell):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        for side in ("top", "left", "bottom", "right"):
            border = tc_borders.find(qn(f"w:{side}"))
            if border is None:
                border = OxmlElement(f"w:{side}")
                tc_borders.append(border)
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "6")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")

    def mark_row(row, header=False):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if header and tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def set_table_widths(tbl):
        tbl_pr = tbl._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), "10000")
        tbl_w.set(qn("w:type"), "pct")
        layout = tbl_pr.first_child_found_in("w:tblLayout")
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")
        if detailed:
            widths = [66]
            if use_qty:
                widths.append(18)
            if use_unit_price:
                widths.append(20)
            widths += [20, 14, 22]
        else:
            widths = [74, 28, 22, 30]
        for row in tbl.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Mm(width)
                set_cell_borders(cell)
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_mar = tc_pr.first_child_found_in("w:tcMar")
                if tc_mar is None:
                    tc_mar = OxmlElement("w:tcMar")
                    tc_pr.append(tc_mar)
                for side, val in (("top", "80"), ("bottom", "80"), ("left", "90"), ("right", "90")):
                    mar = tc_mar.find(qn(f"w:{side}"))
                    if mar is None:
                        mar = OxmlElement(f"w:{side}")
                        tc_mar.append(mar)
                    mar.set(qn("w:w"), val)
                    mar.set(qn("w:type"), "dxa")
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width * 56.7)))
                tc_w.set(qn("w:type"), "dxa")

    def remove_table(tbl):
        parent = tbl._tbl.getparent()
        if parent is not None:
            parent.remove(tbl._tbl)

    if len(table.columns) != len(columns):
        new_table = doc.add_table(rows=1, cols=len(columns))
        new_table.style = table.style
        table._tbl.addprevious(new_table._tbl)
        remove_table(table)
        table = new_table
    else:
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)

    table.autofit = False
    set_table_widths(table)
    if table.rows:
        header = table.rows[0]
        mark_row(header, header=True)
        for cell, label in zip(header.cells, columns):
            set_cell_text(cell, label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=8.5)

    def cell_values(line):
        prefix = ""
        if line.get("is_option") or line.get("type_ligne") == "option":
            prefix = "Option - "
        elif line.get("type_ligne") == "remise":
            prefix = "Remise - "
        elif line.get("type_ligne") in {"information", "info"}:
            prefix = "Information - "
        designation = prefix + str(line.get("designation") or "")
        if line.get("description"):
            designation += "\n" + str(line.get("description"))
        values = [designation]
        if detailed:
            if use_qty:
                qty = str(line.get("quantite") or "")
                unit = str(line.get("unite") or "")
                values.append((qty + (" " + unit if unit else "")).strip())
            if use_unit_price:
                values.append(line.get("prix_unitaire_ht") or "")
        values += [line.get("total_ht_fmt") or "", line.get("taux_tva_fmt") or "", line.get("total_ttc_fmt") or ""]
        if line.get("excluded_from_total"):
            values[-1] = (values[-1] + "\n(hors total)").strip()
        return values

    for line in lignes:
        row = table.add_row()
        mark_row(row)
        values = cell_values(line)
        for idx, (cell, value) in enumerate(zip(row.cells, values)):
            align = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            set_cell_text(cell, value, italic=bool(line.get("excluded_from_total")), align=align, size=8.5)

    if detailed or len(lignes) > 1:
        row = table.add_row()
        mark_row(row)
        total_label = "Total général"
        total_values = [total_label]
        if detailed:
            if use_qty:
                total_values.append("")
            if use_unit_price:
                total_values.append("")
        total_values += [data.get("TOTAL_HT") or "", data.get("TOTAL_TVA") or "", data.get("TOTAL_TTC") or ""]
        for idx, (cell, value) in enumerate(zip(row.cells, total_values)):
            align = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            set_cell_text(cell, value, bold=True, align=align, size=8.5)

    set_table_widths(table)
    last_node = table._tbl
    notes = [str(data.get("FINANCIAL_NOTE") or "").strip(), str(data.get("TVA_NOTE") or "").strip()]
    for note in [n for n in notes if n]:
        p = OxmlElement("w:p")
        ppr = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "80")
        spacing.set(qn("w:after"), "0")
        ppr.append(spacing)
        p.append(ppr)
        for line_idx, line in enumerate(note.replace("\r\n", "\n").replace("\r", "\n").split("\n")):
            if line_idx:
                br_run = OxmlElement("w:r")
                br = OxmlElement("w:br")
                br_run.append(br)
                p.append(br_run)
            r = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "17")
            rpr.append(sz)
            r.append(rpr)
            t = OxmlElement("w:t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = line
            r.append(t)
            p.append(r)
        last_node.addnext(p)
        last_node = p
    return True


def analyser_template(docx_path: Path) -> List[str]:
    """
    Retourne la liste des variables Jinja2 ({{ VARIABLE }}) présentes dans le
    template. Lit directement le XML pour ne rien manquer (corps, headers,
    footers, tableaux).
    """
    variables = set()
    with zipfile.ZipFile(docx_path) as zf:
        for name in zf.namelist():
            if not name.endswith(".xml"):
                continue
            try:
                content = zf.read(name).decode("utf-8", errors="ignore")
            except Exception:
                continue
            for m in JINJA_VAR_PATTERN.finditer(content):
                variables.add(m.group(1))
            for m in JINJA_LOOP_PATTERN.finditer(content):
                variables.add(m.group(1))
    return sorted(variables)


# ============================================================
# 2. ANNOTATION AUTO — applique des substitutions classiques
# ============================================================

# Mapping par défaut : texte du devis original → marqueur Jinja2
# Vous pouvez le surcharger dans le back-office pour chaque template.
DEFAULT_SUBSTITUTIONS = [
    # (texte_original, marqueur, exact_match)
    ("Marseille, le 21 juin 2022",       "Marseille, le {{ DATE_EMISSION }}", False),
    ("client",                            "{{ DEST_LIGNE1 }}",                  True),
    ("adresse 1",                         "{{ DEST_LIGNE2 }}",                  False),
    ("adresse 2",                         "{{ DEST_LIGNE3 }}",                  False),
    ("contact",                           "{{ DEST_LIGNE4 }}",                  True),
    ("Proposition ME",                    "Proposition {{ NUMERO_DEVIS }}",     False),
    ("Entretien des parties communes",    "{{ TYPE_PRESTATION }}",              False),
    ("52 rue Louis Astruc,",              "{{ SITE_ADRESSE }},",                False),
    ("13005 MARSEILLE.",                  "{{ SITE_CP_VILLE }}.",               False),
]


def _replace_in_paragraph(para, old_text, new_text, exact_match=False):
    """Remplace une occurrence dans un paragraphe Word en préservant le formatage.

    Gère le cas fréquent où Word fragmente un mot en plusieurs runs : on
    reconstruit le texte complet du paragraphe, on effectue le remplacement, puis
    on réécrit le tout dans le premier run (en conservant son formatage) et on
    vide les runs suivants. C'est ce qui corrigeait l'échec silencieux sur la
    zone « contact » du modèle d'origine.

    Si exact_match=True, le paragraphe entier (après normalisation des espaces
    insécables et tabulations) doit être égal à old_text.
    """
    full_text = "".join(r.text for r in para.runs)
    if exact_match:
        normalized = full_text.replace("\xa0", " ").replace("\t", " ").strip()
        if normalized != old_text.replace("\xa0", " ").strip():
            return False
        new_full = new_text
    else:
        # tolérance aux espaces insécables : on compare aussi une version normalisée
        if old_text in full_text:
            new_full = full_text.replace(old_text, new_text, 1)
        else:
            norm = full_text.replace("\xa0", " ")
            if old_text in norm:
                new_full = norm.replace(old_text, new_text, 1)
            else:
                return False
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
        return True
    return False


def _coerce_substitutions(substitutions):
    """Normalise les substitutions en tuples (old, new, exact).

    Accepte soit des tuples (texte, marqueur, exact), soit des objets DB
    TemplateSubstitution (champs texte_origine / marqueur / exact_match).
    """
    if substitutions is None:
        return list(DEFAULT_SUBSTITUTIONS)
    out = []
    for s in substitutions:
        if isinstance(s, (tuple, list)):
            out.append((s[0], s[1], bool(s[2]) if len(s) > 2 else False))
        else:  # objet DB
            out.append((s.texte_origine, s.marqueur, bool(s.exact_match)))
    return out


def annoter_auto(docx_in: Path, docx_out: Path, substitutions=None) -> int:
    """
    Annote automatiquement un .docx brut avec des marqueurs Jinja2 selon la
    table de substitutions fournie. Retourne le nombre de substitutions
    effectuées.

    Le doc original est PRÉSERVÉ — on travaille sur une copie.
    """
    substitutions = _coerce_substitutions(substitutions)

    doc = Document(str(docx_in))
    count = 0

    # Traiter d'abord les substitutions à exact_match (plus prioritaires)
    sorted_subs = sorted(substitutions, key=lambda s: not s[2])

    for old, new, exact in sorted_subs:
        replaced = False
        # Corps du document
        for p in doc.paragraphs:
            if _replace_in_paragraph(p, old, new, exact_match=exact):
                count += 1
                replaced = True
                break
        # Tableaux
        if not replaced:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if _replace_in_paragraph(p, old, new, exact_match=exact):
                                count += 1
                                replaced = True
                                break
                        if replaced:
                            break
                    if replaced:
                        break
                if replaced:
                    break

    doc.save(str(docx_out))
    _nettoyer_doublons_zip(docx_out)
    return count


def _nettoyer_doublons_zip(docx_path: Path):
    """
    python-docx peut parfois produire un .docx avec des entrées ZIP dupliquées
    (notamment docProps/core.xml), ce qui empêche LibreOffice de l'ouvrir.
    Cette fonction nettoie en ne gardant que la première occurrence de chaque
    nom de fichier dans l'archive.
    """
    tmp_path = docx_path.with_suffix(".tmp.docx")
    seen = set()
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                if info.filename in seen:
                    continue
                seen.add(info.filename)
                zout.writestr(info, zin.read(info.filename))
    shutil.move(tmp_path, docx_path)


def _restaurer_image_couverture_bureaux(template_path: Path, docx_path: Path) -> None:
    """Restaure le 1er visuel de couverture Bureaux si son ancre Word disparaît."""
    target_media = "media/image1.jpeg"
    try:
        from lxml import etree
        with zipfile.ZipFile(template_path, "r") as zsrc:
            src_xml = zsrc.read("word/document.xml")
            src_rels = zsrc.read("word/_rels/document.xml.rels").decode("utf-8", errors="ignore")
        with zipfile.ZipFile(docx_path, "r") as zdst:
            dst_xml = zdst.read("word/document.xml")
            dst_rels = zdst.read("word/_rels/document.xml.rels").decode("utf-8", errors="ignore")

        def rel_id_for(rels: str) -> Optional[str]:
            m = re.search(r'<Relationship[^>]+Id="([^"]+)"[^>]+Target="' + re.escape(target_media) + r'"', rels)
            return m.group(1) if m else None

        src_rid = rel_id_for(src_rels)
        dst_rid = rel_id_for(dst_rels)
        first_page_end = dst_xml.find(b'w:type="page"')
        first_page_xml = dst_xml[:first_page_end if first_page_end > 0 else len(dst_xml)]
        if not src_rid or not dst_rid or f'r:id="{dst_rid}"'.encode("utf-8") in first_page_xml:
            return

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        parser = etree.XMLParser(remove_blank_text=False, recover=True)
        src_root = etree.fromstring(src_xml, parser)
        dst_root = etree.fromstring(dst_xml, parser)
        source_paragraph = None
        for p in src_root.xpath("//w:body/w:p", namespaces=ns):
            if f'r:id="{src_rid}"' in etree.tostring(p, encoding="unicode"):
                source_paragraph = copy.deepcopy(p)
                break
        if source_paragraph is None:
            return

        restored_xml = etree.tostring(source_paragraph, encoding="unicode").replace(
            f'r:id="{src_rid}"', f'r:id="{dst_rid}"'
        )
        restored_p = etree.fromstring(restored_xml.encode("utf-8"), parser)
        body = dst_root.find(".//w:body", namespaces=ns)
        if body is None:
            return

        insert_at = None
        for idx, child in enumerate(list(body)):
            if child.tag != f"{{{ns['w']}}}p":
                continue
            text = "".join(child.xpath(".//w:t/text()", namespaces=ns)).strip().lower()
            if text.startswith("proposition"):
                insert_at = idx
                break
        body.insert(insert_at if insert_at is not None else min(10, len(body)), restored_p)

        tmp_path = docx_path.with_suffix(".cover-image.tmp.docx")
        new_document_xml = etree.tostring(dst_root, encoding="UTF-8", xml_declaration=True, standalone=True)
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                payload = new_document_xml if info.filename == "word/document.xml" else zin.read(info.filename)
                zout.writestr(info, payload)
        shutil.move(tmp_path, docx_path)
    except Exception as exc:
        print(f"[bureaux_petit] Warning: restauration image couverture impossible: {exc}")


# ============================================================
# 3. GÉNÉRATION — remplir un template annoté avec des données
# ============================================================

def _ajuster_images_prestations_complementaires(docx_path: Path) -> None:
    """Applique les mesures du Word source aux 2 images fixes de la page 5.

    Idempotent et sans effet si le template ne contient pas ces images.
    """
    try:
        from lxml import etree
    except Exception:
        return

    ns_rel = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
    ns = {
        "v": "urn:schemas-microsoft-com:vml",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    styles = {
        "media/image8.jpeg": {
            "left_pt": 289.30,
            "top_pt": 182.60,
            "width_pt": 192.40,
            "height_pt": 111.95,
        },
        "media/image9.png": {
            "left_pt": 240.70,
            "top_pt": 370.85,
            "width_pt": 250.15,
            "height_pt": 192.85,
        },
    }
    wrap = {
        "mso-wrap-distance-left": 9.05,
        "mso-wrap-distance-right": 9.05,
        "mso-wrap-distance-top": 0.0,
        "mso-wrap-distance-bottom": 0.0,
    }

    def set_style_pt(style: str, key: str, value: float) -> str:
        val = f"{float(value):.2f}".rstrip("0").rstrip(".")
        match = re.search(rf"{re.escape(key)}:([\-0-9.]+)pt", style)
        if match:
            return style[:match.start(1)] + val + style[match.end(1):]
        return style + f";{key}:{val}pt"

    def w_tag(local_name: str) -> str:
        return f"{{{ns['w']}}}{local_name}"

    def normaliser_texte_xml(value: str) -> str:
        value = (value or "").replace("\u00a0", " ")
        return re.sub(r"\s+", " ", value).strip().lower()

    puces_action_nuisibles = {
        normaliser_texte_xml("Diagnostic de « circulation » des nuisibles"),
        normaliser_texte_xml("Bouchage, condamnation des accès"),
        normaliser_texte_xml("Services de dératisation, désinsectisation, ponctuels ou par contrat."),
        normaliser_texte_xml("Désinfection de locaux par nébulisation (DSVA) ou par contact"),
    }
    paragraphes_action_nuisibles = puces_action_nuisibles | {
        normaliser_texte_xml(
            "Chacune de ces prestations sera analysée séparément du contrat principal et donnera lieu à un devis spécifique."
        )
    }

    def texte_paragraphe(paragraph) -> str:
        return "".join(paragraph.xpath(".//w:t/text()", namespaces=ns)).strip()

    def remove_children(parent, local_name: str) -> None:
        for child in list(parent.findall(w_tag(local_name))):
            parent.remove(child)

    def ensure_child(parent, local_name: str, insert_at: int | None = None):
        child = parent.find(w_tag(local_name))
        if child is None:
            child = etree.Element(w_tag(local_name))
            if insert_at is None:
                parent.append(child)
            else:
                parent.insert(insert_at, child)
        return child

    def appliquer_format_reference_action_nuisibles(paragraph) -> bool:
        """Restaure l'interligne/retrait des puces depuis le Word source.

        Le numId est conserve car il depend du document genere; seules les
        proprietes directes du paragraphe, perdues par les normalisations,
        sont remises comme dans la page 5 du Word de reference.
        """
        txt = normaliser_texte_xml(texte_paragraphe(paragraph))
        if txt not in paragraphes_action_nuisibles:
            return False

        ppr = paragraph.find("w:pPr", namespaces=ns)
        if ppr is None:
            ppr = etree.Element(w_tag("pPr"))
            paragraph.insert(0, ppr)

        changed_local = False
        if txt in puces_action_nuisibles:
            pstyle = ppr.find(w_tag("pStyle"))
            if pstyle is None:
                pstyle = etree.Element(w_tag("pStyle"))
                ppr.insert(0, pstyle)
            if pstyle.get(w_tag("val")) != "Normal":
                pstyle.set(w_tag("val"), "Normal")
                changed_local = True

        spacing = ensure_child(ppr, "spacing")
        expected_spacing = (
            {"lineRule": "auto", "line": "360", "before": "0", "after": "0"}
            if txt in puces_action_nuisibles
            else {"lineRule": "auto", "line": "240", "before": "0", "after": "0"}
        )
        for attr, value in expected_spacing.items():
            if spacing.get(w_tag(attr)) != value:
                spacing.set(w_tag(attr), value)
                changed_local = True

        ind = ensure_child(ppr, "ind")
        expected_ind = (
            {"start": "720", "hanging": "360", "end": "-428"}
            if txt in puces_action_nuisibles
            else {"end": "-428"}
        )
        for attr in ("left", "right", "firstLine"):
            if ind.get(w_tag(attr)) is not None:
                del ind.attrib[w_tag(attr)]
                changed_local = True
        for attr, value in expected_ind.items():
            if ind.get(w_tag(attr)) != value:
                ind.set(w_tag(attr), value)
                changed_local = True

        remove_children(ppr, "tabs")
        tabs = etree.Element(w_tag("tabs"))
        clear_tab = etree.SubElement(tabs, w_tag("tab"))
        clear_tab.set(w_tag("val"), "clear")
        clear_tab.set(w_tag("pos"), "708")
        left_tab = etree.SubElement(tabs, w_tag("tab"))
        left_tab.set(w_tag("val"), "left")
        left_tab.set(w_tag("pos"), "0")
        left_tab.set(w_tag("leader"), "none")
        num_pr = ppr.find(w_tag("numPr"))
        insert_idx = list(ppr).index(num_pr) + 1 if num_pr is not None else len(ppr)
        ppr.insert(insert_idx, tabs)
        changed_local = True

        for keep_tag in ("keepNext", "keepLines"):
            node = ppr.find(w_tag(keep_tag))
            if node is not None:
                ppr.remove(node)
                changed_local = True
        return changed_local

    tmp_path = docx_path.with_suffix(".compl-images.tmp.docx")
    changed = False
    with zipfile.ZipFile(docx_path, "r") as zin:
        rel_map = {}
        try:
            rel_root = etree.fromstring(zin.read("word/_rels/document.xml.rels"))
            for rel in rel_root.xpath(".//rel:Relationship", namespaces=ns_rel):
                rel_map[rel.get("Id")] = rel.get("Target")
        except Exception:
            rel_map = {}
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                payload = zin.read(info.filename)
                if info.filename == "word/document.xml":
                    root = etree.fromstring(payload)
                    for paragraph in root.xpath(".//w:p", namespaces=ns):
                        if appliquer_format_reference_action_nuisibles(paragraph):
                            changed = True
                    for shape in root.xpath(".//v:shape[v:imagedata]", namespaces=ns):
                        image = shape.find("v:imagedata", namespaces=ns)
                        rid = image.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id") if image is not None else None
                        cfg = styles.get(rel_map.get(rid, ""))
                        if not cfg:
                            continue
                        style = shape.get("style") or ""
                        style_new = style
                        style_new = set_style_pt(style_new, "margin-left", cfg["left_pt"])
                        style_new = set_style_pt(style_new, "margin-top", cfg["top_pt"])
                        style_new = set_style_pt(style_new, "width", cfg["width_pt"])
                        style_new = set_style_pt(style_new, "height", cfg["height_pt"])
                        for key, val in wrap.items():
                            style_new = set_style_pt(style_new, key, val)
                        if style_new != style:
                            shape.set("style", style_new)
                            changed = True
                    if changed:
                        payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
                zout.writestr(info, payload)
    if changed:
        shutil.move(tmp_path, docx_path)
        _nettoyer_doublons_zip(docx_path)
    else:
        tmp_path.unlink(missing_ok=True)


def _supprimer_surlignages_jaunes_document(docx_path: Path) -> None:
    """Retire les repères jaunes hérités des modèles sans toucher aux images.

    Les .doc historiques contiennent des placeholders surlignés. Une fois les
    champs remplacés, ces repères ne doivent plus sortir dans le DOCX/PDF client.
    On nettoie donc uniquement les propriétés Word de surlignage/trame jaunes,
    dans une copie générée, en conservant bordures, titres orange et graphismes.
    """
    import xml.etree.ElementTree as ET

    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "v": "urn:schemas-microsoft-com:vml",
        "o": "urn:schemas-microsoft-com:office:office",
        "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    }
    for prefix, uri in namespaces.items():
        ET.register_namespace(prefix, uri)

    ns_w = namespaces["w"]
    qn_w = lambda name: f"{{{ns_w}}}{name}"
    yellow_values = {
        "YELLOW",
        "FFFF00",
        "FFF200",
        "FFF59D",
        "FEF9D9",
        "FFFF99",
        "FFEB3B",
        "FDE047",
        "FFD966",
    }
    placeholder_exact = {
        "<date>",
        "Nom du compte",
        "Code postal",
        "ADRESSE DE PRESTATION",
        "Adresse de prestation",
        "ME-XXXX",
        "ME XXXX",
    }

    def is_yellow(value: str | None) -> bool:
        if value is None:
            return False
        cleaned = str(value).strip().lstrip("#").upper()
        return cleaned in yellow_values

    def is_yellow_shading(element) -> bool:
        return any(
            is_yellow(element.get(qn_w(attr)) or element.get(attr))
            for attr in ("val", "fill", "color", "themeFill")
        )

    def clean_placeholder_text(text: str | None) -> str | None:
        if text is None:
            return text
        cleaned = text
        for placeholder in placeholder_exact:
            cleaned = cleaned.replace(placeholder, "")
        cleaned = re.sub(r"\bX{3,}\b", "", cleaned)
        cleaned = re.sub(r"[\u2026.]{4,}", "", cleaned)
        if re.fullmatch(r"[\s\u00a0.\u2026]+", cleaned or "") and re.search(r"[.\u2026]{3,}", text):
            return ""
        return cleaned

    tmp_path = docx_path.with_suffix(".yellow-clean.tmp.docx")
    changed = False
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                payload = zin.read(info.filename)
                if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                    try:
                        root = ET.fromstring(payload)
                    except Exception:
                        zout.writestr(info, payload)
                        continue

                    local_changed = False
                    parent_map = {child: parent for parent in root.iter() for child in parent}
                    for highlight in list(root.iter(qn_w("highlight"))):
                        val = highlight.get(qn_w("val")) or highlight.get("val")
                        if is_yellow(val):
                            parent = parent_map.get(highlight)
                            if parent is not None:
                                parent.remove(highlight)
                                local_changed = True

                    for shading in list(root.iter(qn_w("shd"))):
                        if is_yellow_shading(shading):
                            parent = parent_map.get(shading)
                            if parent is not None:
                                parent.remove(shading)
                                local_changed = True

                    for text_node in root.iter(qn_w("t")):
                        old_text = text_node.text
                        new_text = clean_placeholder_text(old_text)
                        if new_text != old_text:
                            text_node.text = new_text
                            local_changed = True

                    if local_changed:
                        payload = ET.tostring(root, encoding="UTF-8", xml_declaration=True)
                        changed = True
                zout.writestr(info, payload)
    if changed:
        shutil.move(tmp_path, docx_path)
        _nettoyer_doublons_zip(docx_path)
    else:
        tmp_path.unlink(missing_ok=True)


def _normaliser_images_word_compat(docx_path: Path) -> Path:
    """Normalise le DOCX FINAL pour compatibilité Microsoft Word (images).

    Les masters .doc historiques stockent leurs images en VML flottant
    (<w:pict>/<v:imagedata>). LibreOffice les affiche (d'où un aperçu correct),
    mais Word installé ne rend PAS ce VML flottant -> images « disparues » à
    l'ouverture dans Word. On ré-exporte donc le .docx généré via LibreOffice :
    le VML est converti en DrawingML (enveloppé en mc:AlternateContent), format
    rendu par TOUTES les versions de Word.

    Garanties :
      - positions, dimensions, ratios, habillages et mise en page préservés
        (ré-export fidèle de LibreOffice — pages fixes 3/4 incluses, vérifié) ;
      - les images DrawingML inline (photos uploadées par zone) restent intactes ;
      - AUCUNE modification des templates : agit uniquement sur le fichier généré ;
      - en cas d'échec de LibreOffice, on conserve le DOCX d'origine (jamais de
        régression sur le téléchargement Word/PDF déjà fonctionnel).
    """
    import tempfile
    if not docx_path.exists():
        return docx_path
    try:
        warmup_libreoffice()  # évite l'échec de 1re conversion à froid de LibreOffice
        with _SOFFICE_LOCK:
            with tempfile.TemporaryDirectory(prefix="lo_norm_out_") as outdir:
                with tempfile.TemporaryDirectory(prefix="lo_profile_") as profile:
                    result = subprocess.run(
                        [SOFFICE_BIN, "--headless",
                         f"-env:UserInstallation=file://{profile}",
                         "--convert-to", "docx",
                         "--outdir", str(outdir), str(docx_path)],
                        capture_output=True, text=True, timeout=120,
                    )
                produced = Path(outdir) / (docx_path.stem + ".docx")
                if result.returncode == 0 and produced.exists() and produced.stat().st_size > 0:
                    shutil.move(str(produced), str(docx_path))
                    _nettoyer_doublons_zip(docx_path)
                    return docx_path
                print(f"[normalisation Word] échec, DOCX conservé tel quel : "
                      f"{(result.stderr or result.stdout or 'sans message')[:200]}")
    except Exception as exc:
        print(f"[normalisation Word] exception, DOCX conservé tel quel : {exc}")
    return docx_path


def generer_devis(template_path: Path, data: Dict, output_path: Path) -> Path:
    """
    Génère un .docx en remplissant le template annoté avec les données fournies.

    Args:
        template_path : chemin vers le template Word annoté
        data : dictionnaire des variables Jinja2 à injecter
        output_path : chemin du .docx final

    Returns:
        Le chemin du fichier généré.
    """
    if not template_path.exists():
        raise FileNotFoundError(f"Template introuvable : {template_path}")

    template_stem = template_path.stem.lower()
    if template_stem in COPRO_PETITE_TEMPLATE_CODES:
        _generer_copro_petite_preserve_layout(template_path, data, output_path)
        return _normaliser_images_word_compat(output_path)
    if template_stem == "bureaux_petit":
        _generer_bureaux_petit_preserve_layout(template_path, data, output_path)
        return _normaliser_images_word_compat(output_path)
    if template_stem.startswith(PONCTUEL_TEMPLATE_PREFIX) and template_stem != "ponctuel_generique":
        _generer_ponctuel_source_preserve_layout(template_path, data, output_path)
        return _normaliser_images_word_compat(output_path)

    _filtrer_mentions_ponctuelles(data)
    doc = DocxTemplate(str(template_path))

    # Les photos sont injectées APRÈS le rendu (post-traitement python-docx),
    # méthode robuste qui garantit l'alignement et évite les soucis d'InlineImage.
    photos = data.pop("PRESTATIONS_PHOTOS", None) or data.pop("PHOTOS", None) or []

    # Garde-fou : injecter une valeur vide pour les variables absentes du payload
    # (évite les erreurs "undefined" de Jinja2)
    variables_template = analyser_template(template_path)
    LIST_VARS = {"OPTIONS", "PRESTATIONS", "LIGNES"}
    for var in variables_template:
        if var not in data:
            data[var] = [] if var in LIST_VARS else ""

    # Retours à la ligne : toute valeur texte multi-lignes est enveloppée dans un
    # Listing docxtpl, qui convertit chaque \n en véritable saut de ligne Word
    # (<w:br/>) et préserve les espaces. Sans ça, les \n seraient écrasés/collés.
    for k, val in list(data.items()):
        if isinstance(val, str) and ("\n" in val or "\r" in val):
            data[k] = Listing(val.replace("\r\n", "\n").replace("\r", "\n"))

    doc.render(data)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    _nettoyer_doublons_zip(output_path)
    _ajuster_images_prestations_complementaires(output_path)

    # Post-traitement : insérer les photos à l'emplacement du marqueur @@ZONE_PHOTOS@@
    # Pipeline ponctuel:
    # - le template Word porte la pagination principale;
    # - l'ancrage court ne reserve que la zone du logo pour les ruptures utiles
    #   (Bon pour accord, CGV, debordement de longues prestations);
    # - "Bon pour accord" est aussi rendu indivisible via keepNext/keepLines.
    # Chaque etape reste idempotente.
    _injecter_photos(output_path, photos, mode_ponctuel=template_stem.startswith(PONCTUEL_TEMPLATE_PREFIX))
    _injecter_encart_materiel(
        output_path,
        data.get("MATERIEL_SELECTIONNE") or [],
        mode_ponctuel=template_stem.startswith(PONCTUEL_TEMPLATE_PREFIX),
    )
    _supprimer_rouge_document(output_path)
    preserve_original_layout = template_stem in COPRO_PETITE_TEMPLATE_CODES
    famille_template = str(data.get("_FAMILLE_TEMPLATE", "") or "").lower()
    _ajouter_signature_cgv(output_path)
    if not preserve_original_layout:
        _normaliser_interligne_document(output_path)
        _mettre_en_page_devis_ponctuel(output_path)
        if template_stem.startswith(PONCTUEL_TEMPLATE_PREFIX):
            _normaliser_entete_ponctuel(output_path)
        _verrouiller_bloc_bon_pour_accord(output_path, data.get("DATE_SIGNATURE"))
        _ancrer_depart_pages_ponctuel(output_path)
        if template_stem.startswith(PONCTUEL_TEMPLATE_PREFIX):
            _normaliser_cgv_ponctuel(output_path)
        if famille_template == "contrat":
            _compacter_paragraphes_vides(output_path)
        _ajuster_images_prestations_complementaires(output_path)
        if template_stem.startswith(PONCTUEL_TEMPLATE_PREFIX):
            _supprimer_surlignages_jaunes_document(output_path)
    return _normaliser_images_word_compat(output_path)


def _filtrer_mentions_ponctuelles(data: Dict) -> None:
    """Evite qu'une mention metier ponctuelle fuie vers un autre type de devis."""
    modele_code = str(data.get("MODELE_CODE", "") or "").lower()
    if modele_code not in {"encombrants_caves", "encombrants_divers"}:
        data["MENTION_VALIDITE"] = ""
        data["MENTION_SPECIFIQUE"] = ""


def _generer_ponctuel_source_preserve_layout(template_path: Path, data: Dict, output_path: Path) -> Path:
    """Rend un devis ponctuel depuis son Word source metier.

    Le template physique est specifique au type (tag, tapis, vitrerie,
    encombrants...). On ne le transforme pas en mise en page generique : on
    remplace seulement les champs variables et la liste des prestations.
    """
    from docx import Document as _Doc
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
    from docx.shared import Pt

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(template_path, output_path)
    doc = _Doc(str(output_path))

    def norm(txt: str) -> str:
        import unicodedata
        txt = (txt or "").lower().replace("\xa0", " ")
        txt = unicodedata.normalize("NFKD", txt)
        txt = "".join(ch for ch in txt if not unicodedata.combining(ch))
        txt = re.sub(r"[\W_]+", " ", txt, flags=re.UNICODE)
        return re.sub(r"\s+", " ", txt).strip()

    def scalar(name: str) -> str:
        val = data.get(name, "")
        if isinstance(val, (list, dict, tuple)):
            return ""
        return str(val or "")

    def set_text(paragraph, text: str) -> None:
        xml_text_nodes = list(paragraph._element.iter(qn("w:t")))
        if xml_text_nodes:
            xml_text_nodes[0].text = text
            for node in xml_text_nodes[1:]:
                node.text = ""
            return
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)

    def remove_paragraph(paragraph) -> None:
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)

    def paragraph_contains_media(paragraph) -> bool:
        return bool(
            paragraph._element.xpath(
                ".//*[local-name()='drawing' or local-name()='pict' or local-name()='shape' or local-name()='imagedata']"
            )
        )

    def keep(paragraph, keep_next: bool = False) -> None:
        ppr = paragraph._p.get_or_add_pPr()
        if ppr.find(qn("w:keepLines")) is None:
            ppr.append(OxmlElement("w:keepLines"))
        if keep_next and ppr.find(qn("w:keepNext")) is None:
            ppr.append(OxmlElement("w:keepNext"))

    def all_paragraphs():
        for p in doc.paragraphs:
            yield p
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    def replace_split_vars(paragraph) -> None:
        txt = paragraph.text or ""
        if "{{" not in txt:
            return
        new = txt
        for key, val in data.items():
            if isinstance(val, (str, int, float)):
                new = new.replace("{{ " + key + " }}", str(val)).replace("{{" + key + "}}", str(val))
        if new != txt:
            set_text(paragraph, new)

    def replace_common_fields() -> None:
        non_empty = [(i, p, norm(p.text)) for i, p in enumerate(doc.paragraphs) if norm(p.text)]
        for i, p, txt in non_empty:
            raw = p.text.strip()
            if txt.startswith("marseille le"):
                set_text(p, scalar("DATE_EMISSION_LONGUE") or f"À Marseille, le {scalar('DATE_EMISSION') or scalar('DATE_SIGNATURE')}")
            elif txt.startswith("devis me") or txt == "devis me":
                set_text(p, f"Devis {scalar('NUMERO_DEVIS')}")
            elif txt in {"nom du client", "siga sas", "nh collection", "cabinet paul coudre"}:
                set_text(p, scalar("DEST_LIGNE1"))
            elif txt in {"adresse du client", "7 rue d italie"} or "boulevard des dames" in txt or "rue saint ferreol" in txt:
                set_text(p, scalar("DEST_LIGNE3") or scalar("SITE_ADRESSE"))
            elif "1300 marseille" in txt or "13007 arseille" in txt or "13001 marseille" in txt or "13002 marseille" in txt:
                set_text(p, scalar("DEST_LIGNE4") or scalar("SITE_CP_VILLE"))
            elif "@" in raw and "marie-eugenie" not in txt:
                set_text(p, scalar("DEST_LIGNE2"))
            elif txt in {"adresse", "adresse de prestation"}:
                set_text(p, scalar("SITE_ADRESSE") or scalar("SITE_CP_VILLE"))
            elif txt.startswith("le present contrat prend effet"):
                set_text(p, f"Le présent contrat prend effet au : {scalar('DATE_PRISE_EFFET')}")
            elif txt.startswith("fait a marseille"):
                set_text(p, f"Fait à Marseille, le {scalar('DATE_SIGNATURE') or scalar('DATE_EMISSION')}")
            replace_split_vars(p)

        # Type de prestation : premier paragraphe significatif apres "Devis".
        for idx, p in enumerate(doc.paragraphs):
            if norm(p.text).startswith("devis "):
                for q in doc.paragraphs[idx + 1: idx + 8]:
                    qnrm = norm(q.text)
                    if qnrm and qnrm not in {"au", "situee au", "située au"}:
                        set_text(q, scalar("TYPE_PRESTATION") or q.text)
                        break
                break

        # Adresse de prestation : paragraphe suivant "au" / "située au".
        for idx, p in enumerate(doc.paragraphs):
            if norm(p.text) in {"au", "situee au", "située au"}:
                for q in doc.paragraphs[idx + 1: idx + 5]:
                    if norm(q.text):
                        site = " ".join(x for x in [scalar("SITE_ADRESSE"), scalar("SITE_CP_VILLE")] if x).strip()
                        if site:
                            set_text(q, site)
                        break
                break

    def replace_prestations() -> None:
        prestations = data.get("PRESTATIONS") or []
        if not isinstance(prestations, list):
            prestations = [str(prestations)]
        prestations = [str(p).strip() for p in prestations if str(p).strip()]
        if not prestations:
            return

        start = None
        for i, p in enumerate(doc.paragraphs):
            if "detail des prestations" in norm(p.text):
                start = i
                break
        if start is None:
            return

        end = None
        stop_tokens = (
            "information", "informations", "photos", "photo",
            "materiel encart",
            "proposition financiere", "proposition financière",
            "bon pour accord", "conditions generales"
        )
        for j in range(start + 1, len(doc.paragraphs)):
            txt = norm(doc.paragraphs[j].text)
            if txt and any(tok in txt for tok in stop_tokens):
                end = j
                break
        if end is None:
            return

        anchor = doc.paragraphs[end]
        for p in reversed(doc.paragraphs[start + 1:end]):
            remove_paragraph(p)
        blank = anchor.insert_paragraph_before("\u00A0")
        blank.paragraph_format.space_before = Pt(0)
        blank.paragraph_format.space_after = Pt(0)
        blank.paragraph_format.line_spacing = 1
        for run in blank.runs:
            run.font.size = Pt(10)
        for op in prestations:
            anchor.insert_paragraph_before(f"-   {op}")

    def replace_table_amounts() -> None:
        if _remplir_tableau_financier_docx(doc, data):
            return
        values = [scalar("PRIX_HT") or scalar("FORFAIT_HT") or scalar("TOTAL_HT"),
                  scalar("TVA") or scalar("FORFAIT_TVA") or scalar("TOTAL_TVA"),
                  scalar("PRIX_TTC") or scalar("FORFAIT_TTC") or scalar("TOTAL_TTC")]
        for table in doc.tables:
            for row in table.rows:
                trpr = row._tr.get_or_add_trPr()
                if trpr.find(qn("w:cantSplit")) is None:
                    trpr.append(OxmlElement("w:cantSplit"))
            # Remplissage minimal : seulement les cellules déjà numériques/vides
            # en fin de tableau, sans reconstruire le tableau source.
            candidate_cells = []
            for row in table.rows:
                for cell in row.cells:
                    txt = norm(cell.text)
                    if not txt or re.search(r"\d+[,\.]\d+", cell.text or ""):
                        candidate_cells.append(cell)
            for cell, value in zip(candidate_cells[-3:], values):
                if value and cell.paragraphs:
                    set_text(cell.paragraphs[0], value)

    replace_common_fields()
    replace_prestations()
    replace_table_amounts()

    for p in doc.paragraphs:
        txt = norm(p.text)
        if "proposition financiere" in txt or "bon pour accord" in txt:
            keep(p, keep_next=True)
        elif "sas marie eugenie" in txt or "le client" in txt or "fait a marseille" in txt:
            keep(p, keep_next=True)

    template_has_materiel_marker = _docx_contient_marqueur(template_path, "@@MATERIEL_ENCART@@")

    doc.save(str(output_path))
    _nettoyer_doublons_zip(output_path)
    _ajuster_images_prestations_complementaires(output_path)
    _injecter_photos(output_path, data.get("PRESTATIONS_PHOTOS") or data.get("PHOTOS") or [], mode_ponctuel=True)
    if template_has_materiel_marker:
        _restaurer_marqueur_materiel_si_absent(output_path)
    _injecter_encart_materiel(output_path, data.get("MATERIEL_SELECTIONNE") or [], mode_ponctuel=True)
    _supprimer_rouge_document(output_path)
    _normaliser_entete_ponctuel(output_path)
    _verrouiller_bloc_bon_pour_accord(output_path, scalar("DATE_SIGNATURE") or scalar("DATE_EMISSION"))
    _ancrer_depart_pages_ponctuel(output_path)
    _normaliser_cgv_ponctuel(output_path)
    _ajuster_images_prestations_complementaires(output_path)
    _supprimer_surlignages_jaunes_document(output_path)
    return output_path


def _generer_bureaux_petit_preserve_layout(template_path: Path, data: Dict, output_path: Path) -> Path:
    """Rend le modèle Bureaux petite surface depuis son Word maître.

    Contrairement au chemin ponctuel, la pagination reste portée par le template
    Word. Les zones Bureaux sont seulement affichées/retirées et leurs opérations
    sont remplacées par les listes envoyées par l'interface.
    """
    from docx import Document as _Doc
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(template_path, output_path)
    doc = _Doc(str(output_path))

    zone_defs = [
        ("ACCUEIL_BUREAUX", "Accueil / Bureaux / Salle de réunion", "FREQ_ACCUEIL_BUREAUX",
         ("accueil bureaux salle reunion", "accueil bureaux salle de reunion")),
        ("CIRCULATION", "Circulations / Couloirs / Dégagements", "FREQ_CIRCULATION",
         ("circulation", "circulations couloirs degagements")),
        ("SANITAIRES", "Sanitaires", "FREQ_SANITAIRES", ("sanitaires",)),
        ("CUISINE", "Espace pause / Cafétéria", "FREQ_CUISINE", ("cuisine espace pause cafeteria", "cuisine")),
        ("VITRERIE", "Vitrerie intérieure / extérieure", "FREQ_VITRERIE",
         ("vitrerie exterieure interieure", "vitrerie interieure exterieure")),
        ("CONSOMMABLES", "Consommables sanitaires / Dimensionnement collaborateurs", None,
         ("fourniture et mise en place des consommables sanitaires", "consommables sanitaires")),
    ]

    def norm(txt: str) -> str:
        import unicodedata
        txt = (txt or "").lower().replace("\xa0", " ")
        txt = unicodedata.normalize("NFKD", txt)
        txt = "".join(ch for ch in txt if not unicodedata.combining(ch))
        txt = re.sub(r"[\W_]+", " ", txt, flags=re.UNICODE)
        return re.sub(r"\s+", " ", txt).strip()

    def scalar(name: str) -> str:
        val = data.get(name, "")
        if isinstance(val, (list, dict, tuple)):
            return ""
        value = str(val)
        cover_limits = {
            "DATE_EMISSION": 16,
            "DEST_LIGNE1": 18,
            "DEST_LIGNE3": 18,
            "DEST_LIGNE4": 18,
            "NUMERO_DEVIS": 12,
            "SITE_ADRESSE": 30,
        }
        limit = cover_limits.get(name)
        if limit:
            value = re.sub(r"\s+", " ", value).strip()
            if limit > 3 and len(value) > limit:
                value = value[:limit - 3].rstrip() + "..."
        return value

    def fixed_field(name: str, max_chars: int, fallback: str = "") -> str:
        """Limite un champ variable a sa zone du modele Bureaux verrouille."""
        value = re.sub(r"\s+", " ", scalar(name)).strip() or str(fallback or "").strip()
        if max_chars > 3 and len(value) > max_chars:
            return value[:max_chars - 3].rstrip() + "..."
        return value

    def set_paragraph_text(paragraph, text: str) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)

    def set_paragraph_text_styled(paragraph, text: str, *, bold=False, italic=False, underline=False) -> None:
        set_paragraph_text(paragraph, text)
        if not paragraph.runs:
            paragraph.add_run("")
        for run in paragraph.runs:
            if run.text:
                run.bold = bold
                run.italic = italic
                run.underline = underline

    def set_keep(paragraph, keep_next=True):
        ppr = paragraph._p.get_or_add_pPr()
        if ppr.find(qn("w:keepLines")) is None:
            ppr.append(OxmlElement("w:keepLines"))
        existing = ppr.find(qn("w:keepNext"))
        if keep_next and existing is None:
            ppr.append(OxmlElement("w:keepNext"))
        if not keep_next and existing is not None:
            ppr.remove(existing)

    def clear_list_format(paragraph) -> None:
        ppr = paragraph._p.get_or_add_pPr()
        num_pr = ppr.find(qn("w:numPr"))
        if num_pr is not None:
            ppr.remove(num_pr)
        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.first_line_indent = None

    def reset_page_start_title(paragraph, prefix: str = "", exact_text: str = "") -> None:
        """Replace fragile Word numbering by real text for page-start titles."""
        ppr = paragraph._p.get_or_add_pPr()
        for tag in ("w:numPr", "w:tabs"):
            node = ppr.find(qn(tag))
            if node is not None:
                ppr.remove(node)
        pf = paragraph.paragraph_format
        pf.left_indent = Pt(0)
        pf.first_line_indent = Pt(0)
        pf.right_indent = Pt(0)
        if exact_text:
            set_paragraph_text(paragraph, exact_text)
        elif prefix and paragraph.runs and not paragraph.text.strip().startswith(prefix.strip()):
            paragraph.runs[0].text = prefix + paragraph.runs[0].text

    def insert_blank_after(paragraph) -> None:
        """Insert one real blank paragraph after a paragraph, idempotently."""
        nxt = paragraph._p.getnext()
        if nxt is not None and nxt.tag == qn("w:p"):
            if not "".join(node.text or "" for node in nxt.iter(qn("w:t"))).strip():
                return
        blank = OxmlElement("w:p")
        ppr = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "210")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.append(spacing)
        blank.append(ppr)
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        text.text = "\u00A0"
        run.append(text)
        blank.append(run)
        paragraph._p.addnext(blank)

    def insert_paragraph_after(paragraph, text: str = ""):
        from docx.text.paragraph import Paragraph
        new_p = OxmlElement("w:p")
        new_r = OxmlElement("w:r")
        new_t = OxmlElement("w:t")
        new_t.text = text
        new_r.append(new_t)
        new_p.append(new_r)
        paragraph._p.addnext(new_p)
        return Paragraph(new_p, paragraph._parent)

    def force_page_break_before(paragraph) -> None:
        ppr = paragraph._p.get_or_add_pPr()
        if ppr.find(qn("w:pageBreakBefore")) is None:
            ppr.append(OxmlElement("w:pageBreakBefore"))

    def clear_page_break_before(paragraph) -> None:
        ppr = paragraph._p.get_or_add_pPr()
        for node in list(ppr.findall(qn("w:pageBreakBefore"))):
            ppr.remove(node)

    def style_main_section_title(paragraph) -> None:
        if not paragraph.runs:
            paragraph.add_run("")
        for run in paragraph.runs:
            if run.text:
                run.bold = True
                run.underline = True
                run.font.size = Pt(10)

    def remove_previous_empty_page_breaks(paragraph) -> None:
        parent = paragraph._p.getparent()
        if parent is None:
            return
        previous = paragraph._p.getprevious()
        while previous is not None and previous.tag == qn("w:p"):
            has_text = "".join(node.text or "" for node in previous.iter(qn("w:t"))).strip()
            has_media = bool(previous.xpath(".//*[local-name()='drawing' or local-name()='pict' or local-name()='shape' or local-name()='imagedata']"))
            if has_text or has_media:
                break
            to_remove = previous
            previous = previous.getprevious()
            parent.remove(to_remove)

    def apply_content_start(paragraph, *, force_break: bool = False) -> None:
        """Position unique de départ de contenu sous le logo Marie-Eugénie."""
        if force_break:
            remove_previous_empty_page_breaks(paragraph)
            force_page_break_before(paragraph)
        pf = paragraph.paragraph_format
        pf.left_indent = Pt(0)
        pf.first_line_indent = Pt(0)
        pf.right_indent = Pt(0)
        pf.space_before = Pt(CONTENT_START_BEFORE_PT)
        pf.space_after = Pt(CONTENT_TITLE_AFTER_PT)
        pf.keep_with_next = True
        style_main_section_title(paragraph)

    def insert_hard_page_break_before(paragraph) -> None:
        br_p = OxmlElement("w:p")
        br_r = OxmlElement("w:r")
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        br_r.append(br)
        br_p.append(br_r)
        paragraph._p.addprevious(br_p)

    def prepend_hard_page_break(paragraph) -> None:
        if paragraph._p.xpath('.//w:br[@w:type="page"]'):
            return
        br_r = OxmlElement("w:r")
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        br_r.append(br)
        insert_at = 1 if paragraph._p.find(qn("w:pPr")) is not None else 0
        paragraph._p.insert(insert_at, br_r)

    def replace_contains(paragraph, replacements: Dict[str, str]) -> None:
        txt = paragraph.text or ""
        for needle, repl in replacements.items():
            if needle in txt:
                set_paragraph_text(paragraph, txt.replace(needle, repl))
                return

    def replace_template_vars(paragraph) -> None:
        if "{{" in (paragraph.text or ""):
            full_text = paragraph.text
            def repl_full(match):
                return scalar(match.group(1).strip())
            rendered = re.sub(r"\{\{\s*([^}]+?)\s*\}\}", repl_full, full_text)
            if rendered != full_text:
                set_paragraph_text(paragraph, rendered)
                return
        for run in paragraph.runs:
            txt = run.text or ""
            if "{{" not in txt:
                continue
            def repl(match):
                return scalar(match.group(1).strip())
            run.text = re.sub(r"\{\{\s*([^}]+?)\s*\}\}", repl, txt)

    def remove_paragraph(paragraph) -> None:
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)

    def iter_all_paragraphs():
        for p in doc.paragraphs:
            yield p
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    def find_zone_ranges() -> Dict[str, tuple[int, int]]:
        starts: List[tuple[int, str]] = []
        for code, _title, _freq_var, aliases in zone_defs:
            alias_norm = [norm(a) for a in aliases]
            for idx, p in enumerate(doc.paragraphs):
                hay = norm(p.text)
                if hay and any(a and (hay == a or a in hay) for a in alias_norm):
                    starts.append((idx, code))
                    break
        starts.sort()
        ranges: Dict[str, tuple[int, int]] = {}
        for pos, (start, code) in enumerate(starts):
            next_start = starts[pos + 1][0] if pos + 1 < len(starts) else None
            end = (next_start - 1) if next_start is not None else start
            if code == "CONSOMMABLES":
                for j in range(start + 1, len(doc.paragraphs)):
                    if "prestations complementaires" in norm(doc.paragraphs[j].text):
                        end = j - 1
                        break
            ranges[code] = (start, end)
        return ranges

    def show_zone(code: str, freq_var: Optional[str]) -> bool:
        key = f"SHOW_{code}"
        if key in data:
            return bool(data.get(key))
        if freq_var:
            return bool(str(data.get(freq_var, "") or "").strip())
        return False

    def render_zone(code: str, title: str, freq_var: Optional[str], start: int, end: int) -> None:
        paras = doc.paragraphs[start:end + 1]
        if not paras:
            return
        ops_key = f"OPS_{code}"
        operations = data.get(ops_key)
        if not isinstance(operations, list):
            operations = []
        operations = [str(op).replace("{{ NB_COLLABORATEURS }}", scalar("NB_COLLABORATEURS")).strip()
                      for op in operations if str(op).strip()]
        # Prestations libres (descriptives, sans prix) ajoutées par l'utilisateur :
        # elles s'affichent comme des puces d'opération dans la zone, mais restent
        # hors de tout calcul (voir routes_devis : LIBRE_* n'alimente pas le tarif).
        libres = data.get(f"LIBRE_{code}")
        if isinstance(libres, list):
            operations = operations + [str(x).strip() for x in libres if str(x).strip()]
        freq = str(data.get(freq_var, "") or "").strip() if freq_var else ""
        desired = [(title, "title"), ("", "blank")]
        desired.extend((f"-   {op}", "operation") for op in operations)
        if freq:
            desired.extend([("", "blank"), (f"Fréquence : {freq}", "frequency")])
        usable = list(paras)
        for idx, (text, role) in enumerate(desired):
            if idx < len(usable):
                if role == "title":
                    set_paragraph_text_styled(usable[idx], text, bold=True, underline=True)
                elif role == "frequency":
                    set_paragraph_text_styled(usable[idx], text, bold=True, italic=True)
                else:
                    set_paragraph_text(usable[idx], text)
            else:
                new_p = insert_paragraph_after(usable[-1], text)
                usable.append(new_p)
                if role == "title":
                    set_paragraph_text_styled(new_p, text, bold=True, underline=True)
                elif role == "frequency":
                    set_paragraph_text_styled(new_p, text, bold=True, italic=True)
            current = usable[idx]
            current.paragraph_format.space_before = Pt(0)
            current.paragraph_format.space_after = Pt(4)
            if role == "blank":
                clear_list_format(current)
                set_paragraph_text(current, "\u00A0")
                for run in current.runs:
                    run.font.size = Pt(10)
                current.paragraph_format.line_spacing = 1
                current.paragraph_format.space_after = Pt(0)
                set_keep(current, keep_next=(idx < len(desired) - 1))
            elif role == "title":
                clear_list_format(current)
                current.paragraph_format.space_after = Pt(14)
                set_keep(current, keep_next=True)
            elif role == "operation":
                clear_list_format(current)
                current.paragraph_format.left_indent = Pt(14)
                current.paragraph_format.first_line_indent = Pt(-7)
                current.paragraph_format.space_after = Pt(2)
                set_keep(current, keep_next=(idx < len(desired) - 1))
            elif role == "frequency":
                clear_list_format(current)
                current.paragraph_format.space_before = Pt(12)
                current.paragraph_format.space_after = Pt(10)
                set_keep(current, keep_next=False)
        for extra in usable[len(desired):]:
            remove_paragraph(extra)

    def trim_empty_runs_between_dynamic_blocks() -> None:
        marker_start = "detail et frequences des prestations"
        marker_end = "prestations complementaires"
        active = False
        empty_buffer = []
        for p in list(doc.paragraphs):
            txt = norm(p.text)
            if marker_start in txt:
                active = True
                empty_buffer = []
                continue
            if active and marker_end in txt:
                for ep in empty_buffer:
                    remove_paragraph(ep)
                break
            if not active:
                continue
            if txt:
                for ep in empty_buffer:
                    remove_paragraph(ep)
                empty_buffer = []
            else:
                empty_buffer.append(p)

    def remove_cover_phone_overflow_paragraph() -> None:
        """Evite qu'une ligne telephone de la couverture remonte en page 2.

        Le modele Word contient deja les coordonnees completes dans le pied de
        page et dans les objets de couverture. Apres sauvegarde par python-docx,
        le paragraphe texte "T : 04 91 47 14 38" peut deborder au-dessus de la
        page "Presentation". On le retire uniquement dans la zone de couverture,
        avant le titre de la page 2.
        """
        for p in list(doc.paragraphs):
            txt = norm(p.text)
            if "presentation de la societe marie eugenie" in txt:
                break
            if txt in {"t 04 91 47 14 38", "tel 04 91 47 14 38"}:
                remove_paragraph(p)
                break

    def normalize_presentation_media_wrap() -> None:
        """Evite que l'image de presentation se superpose au texte sous LibreOffice.

        Certains ancrages herites du .doc source sont interpretes comme des
        objets flottants sans habillage lors du passage DOCX -> PDF. On force un
        habillage carre sur les images ancrees apres la page de couverture.
        """
        body_children = list(doc.element.body)
        seen_presentation = False
        for child in body_children:
            if child.tag != qn("w:p"):
                continue
            text = "".join(node.text or "" for node in child.iter(qn("w:t")))
            ntext = norm(text)
            if "presentation de la societe marie eugenie" in ntext:
                seen_presentation = True
            if not seen_presentation:
                continue
            for anchor in child.findall(".//" + qn("wp:anchor")):
                for wrap_tag in ("wp:wrapNone", "wp:wrapTopAndBottom", "wp:wrapTight", "wp:wrapThrough"):
                    for wrap in list(anchor.findall(qn(wrap_tag))):
                        anchor.remove(wrap)
                if anchor.find(qn("wp:wrapSquare")) is None:
                    wrap_square = OxmlElement("wp:wrapSquare")
                    wrap_square.set("wrapText", "bothSides")
                    extent = anchor.find(qn("wp:extent"))
                    insert_at = list(anchor).index(extent) + 1 if extent is not None else 0
                    anchor.insert(insert_at, wrap_square)

    def constrain_presentation_text_column() -> None:
        """Reserve la colonne droite a l'image de presentation Bureaux."""
        active = False
        for p in doc.paragraphs:
            txt = norm(p.text)
            if "presentation de la societe marie eugenie" in txt:
                active = True
                continue
            if active and txt.startswith("depuis pres de 50 ans"):
                break
            if active and txt:
                p.paragraph_format.right_indent = Pt(235)

    def tighten_presentation_start() -> None:
        """Remonte la page Presentation juste sous le logo, sans toucher au logo."""
        seen_title = False
        for idx, p in enumerate(list(doc.paragraphs)):
            txt = norm(p.text)
            if "presentation de la societe marie eugenie" not in txt:
                if seen_title and txt:
                    p.paragraph_format.space_before = Pt(0)
                    break
                continue
            seen_title = True
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            removed = 0
            j = idx - 1
            while j >= 0 and removed < 2:
                prev = doc.paragraphs[j]
                prev_txt = norm(prev.text)
                if prev_txt or paragraph_contains_media(prev):
                    break
                remove_paragraph(prev)
                removed += 1
                j -= 1
            continue

    def lock_presentation_on_single_page() -> None:
        """Compacte uniquement la page Presentation pour qu'elle tienne sur une page."""
        in_page = False
        presentation_paras = []
        for p in list(doc.paragraphs):
            txt = norm(p.text)
            if "presentation de la societe marie eugenie" in txt:
                in_page = True
            if in_page and "typologie des locaux" in txt:
                reset_page_start_title(p, "1. ")
                clear_page_break_before(p)
                apply_content_start(p, force_break=True)
                break
            if in_page:
                presentation_paras.append(p)

        previous_empty = False
        for p in list(presentation_paras):
            txt = norm(p.text)
            is_empty = not txt and not paragraph_contains_media(p)
            if is_empty and previous_empty:
                remove_paragraph(p)
                continue
            previous_empty = is_empty

        for p in presentation_paras:
            if p._element.getparent() is None:
                continue
            txt = norm(p.text)
            pf = p.paragraph_format
            if "presentation de la societe marie eugenie" not in txt:
                clear_page_break_before(p)
            clear_page_break_before(p)
            pf.keep_with_next = False
            pf.keep_together = False
            pf.space_before = Pt(0)
            pf.space_after = Pt(2)
            pf.line_spacing = 1.04
            if "presentation de la societe marie eugenie" in txt:
                pf.space_before = Pt(6)
                pf.space_after = Pt(2)
                pf.keep_with_next = True
            elif txt.startswith(("depuis pres", "notre vraie richesse", "une entreprise familiale", "les chiffres cles")):
                pf.space_before = Pt(5)
                pf.space_after = Pt(1)
                pf.keep_with_next = True
            elif txt.startswith("effectifs"):
                pf.space_before = Pt(3)
                pf.keep_with_next = True
            elif txt.startswith(("12 vehicules", "mono brosses")):
                pf.space_before = Pt(3)
                pf.space_after = Pt(0)
            for run in p.runs:
                if not run.text:
                    continue
                if txt.startswith(("depuis pres", "notre vraie richesse", "une entreprise familiale", "les chiffres cles")):
                    run.font.size = Pt(9.8)
                elif "presentation de la societe marie eugenie" in txt:
                    run.font.size = Pt(10)
                else:
                    run.font.size = Pt(9.35)

        for p in list(doc.paragraphs):
            txt = norm(p.text)
            if (
                "presentation de la societe marie eugenie" in txt
                or txt.startswith(("depuis pres", "notre vraie richesse", "une entreprise familiale", "les chiffres cles"))
            ):
                insert_blank_after(p)

    def lock_static_page_starts() -> None:
        """Force uniquement les annexes qui n'ont pas deja un saut Word fiable."""
        for p in doc.paragraphs:
            raw = (p.text or "").strip()
            txt = norm(raw)
            if txt.startswith("conditions generales de vente") and raw.upper() == raw:
                p.paragraph_format.page_break_before = True

    def normalize_bureaux_section_numbers() -> None:
        """Stabilise les numeros de titres Bureaux sans dépendre des listes Word."""
        specs = [
            ("presentation de la societe marie eugenie", " - ", "Présentation de la société Marie Eugénie"),
            ("typologie des locaux", ". ", "Typologie des locaux"),
            ("detail et frequences des prestations", ". ", "Détail et fréquences des prestations"),
            ("prestations complementaires", " - ", "Prestations complémentaires"),
            ("proposition financiere", ". ", "Proposition financière"),
            ("date de prise d effet", ". ", "Date de prise d’effet"),
            ("conditions generales de vente", ". ", "Conditions générales de vente"),
        ]
        done = set()
        number = 0
        for p in doc.paragraphs:
            txt = norm(p.text)
            if not txt:
                continue
            for token, sep, label in specs:
                if token in txt and token not in done:
                    number += 1
                    title = f"{number}{sep}{label}"
                    reset_page_start_title(p, exact_text=title)
                    style_main_section_title(p)
                    if token in {"typologie des locaux", "prestations complementaires", "proposition financiere"}:
                        apply_content_start(
                            p,
                            force_break=True,
                        )
                    elif token == "conditions generales de vente":
                        clear_page_break_before(p)
                        pf = p.paragraph_format
                        pf.space_before = Pt(12)
                        pf.space_after = Pt(8)
                        pf.keep_with_next = True
                    if token in {"detail et frequences des prestations", "typologie des locaux"}:
                        insert_blank_after(p)
                    done.add(token)
                    break

    def keep_bureaux_short_dynamic_blocks() -> None:
        """Evite les fins de zones courtes orphelines avant les pages fixes.

        La pagination reste portée par Word : on ne déplace rien en code et on
        n'insère aucun blanc. On indique seulement que le petit bloc
        consommables doit rester solidaire s'il ne tient pas en bas de page.
        """
        in_consumables = False
        consumables: List[Any] = []
        for p in doc.paragraphs:
            txt = norm(p.text)
            if "fourniture et mise en place des consommables sanitaires" in txt:
                in_consumables = True
                consumables = [p]
                continue
            if in_consumables and "prestations complementaires" in txt:
                break
            if in_consumables:
                if txt:
                    consumables.append(p)
                elif consumables:
                    # On garde le premier blanc suivant dans le bloc pour
                    # conserver l'aération du modèle source.
                    consumables.append(p)
        for idx, p in enumerate(consumables):
            set_keep(p, keep_next=idx < len(consumables) - 1)

    def paginate_bureaux_zone_blocks() -> None:
        """Evite qu'une zone continue naturellement a cote du logo.

        Si le bloc estime ne plus tenir dans la zone utile restante, le titre de
        zone porte le saut de page et le meme depart sous logo que les pages
        fixes. Aucun blanc de calage n'est ajoute.
        """
        ranges_now = find_zone_ranges()
        ordered = [
            (ranges_now[code][0], ranges_now[code][1], code)
            for code, _title, freq_var, _aliases in zone_defs
            if code in ranges_now and show_zone(code, freq_var)
        ]
        ordered.sort()
        if not ordered:
            return

        def line_count(text: str) -> int:
            raw = (text or "").strip()
            if not raw:
                return 1
            if raw == "\u00A0":
                return 1
            return max(1, (len(raw) // 78) + 1)

        def block_height(start: int, end: int) -> int:
            return max(3, sum(line_count(p.text) for p in doc.paragraphs[start:end + 1]))

        used = 16
        capacity = 43
        has_zone = False
        for start, end, _code in ordered:
            height = block_height(start, end)
            title = doc.paragraphs[start]
            if has_zone and used + height > capacity:
                apply_content_start(title, force_break=True)
                title.paragraph_format.space_after = Pt(14)
                used = height
            else:
                used += height
            has_zone = True

    # Variables simples et typologie.
    scalar_replacements = {
        "Marseille, le <date>": f"Marseille, le {scalar('DATE_EMISSION')}",
        "Nom du compte": scalar("DEST_LIGNE1"),
        "Adresse": scalar("DEST_LIGNE3") or scalar("SITE_ADRESSE"),
        "Code postal": scalar("DEST_LIGNE4") or scalar("SITE_CP_VILLE"),
        "Proposition ME …": f"Proposition {scalar('NUMERO_DEVIS')}",
        "Nettoyage des bureaux": scalar("TYPE_PRESTATION") or "Nettoyage des bureaux",
        "adresse de prestation": scalar("SITE_ADRESSE") or scalar("SITE_CP_VILLE"),
        "Surface : environ 140 m2": f"Surface : environ {scalar('SURFACE_LOCAUX') or '140 m2'}",
        "XXX bloc sanitaire": f"{scalar('NB_BLOCS_SANITAIRES') or 'XXX'} bloc sanitaire",
        "….. Collaborateurs travaillent dans ces bureaux": f"{scalar('NB_COLLABORATEURS') or '.....'} collaborateurs travaillent dans ces bureaux",
        "Bureau …..": f"Bureau {scalar('REVETEMENT_BUREAU')}",
        "Sanitaire….": f"Sanitaire {scalar('REVETEMENT_SANITAIRE')}",
        "Le présent contrat prend effet au :………………….": f"Le présent contrat prend effet au : {scalar('DATE_PRISE_EFFET')}",
        "Fait à Marseille, le": f"Fait à Marseille, le {scalar('DATE_SIGNATURE') or scalar('DATE_EMISSION')}",
    }
    scalar_replacements.update({
        "Marseille, le <date>": f"Marseille, le {fixed_field('DATE_EMISSION', 16)}",
        "Nom du compte": fixed_field("DEST_LIGNE1", 18),
        "Adresse": fixed_field("DEST_LIGNE3", 18, scalar("SITE_ADRESSE")),
        "Code postal": fixed_field("DEST_LIGNE4", 18, scalar("SITE_CP_VILLE")),
        "Nettoyage des bureaux": fixed_field("TYPE_PRESTATION", 42, "Nettoyage des bureaux"),
        "adresse de prestation": fixed_field("SITE_ADRESSE", 30, scalar("SITE_CP_VILLE")),
    })
    for key in list(scalar_replacements):
        if key.startswith("Proposition ME"):
            scalar_replacements[key] = f"Proposition {fixed_field('NUMERO_DEVIS', 12)}"
    for p in iter_all_paragraphs():
        replace_template_vars(p)
        replace_contains(p, scalar_replacements)
    remove_cover_phone_overflow_paragraph()
    lock_static_page_starts()

    # Tableau financier.
    if not _remplir_tableau_financier_docx(doc, data):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        txt = p.text
                        if "180,00" in txt:
                            set_paragraph_text(p, scalar("TOTAL_HT") or txt)
                        elif "36,00" in txt:
                            set_paragraph_text(p, scalar("TOTAL_TVA") or txt)
                        elif "216,00" in txt:
                            set_paragraph_text(p, scalar("TOTAL_TTC") or txt)

    def lock_financial_section() -> None:
        """Garde le titre Proposition financière et son tableau ensemble."""
        body_children = list(doc.element.body)
        for idx, child in enumerate(body_children):
            if child.tag != qn("w:p"):
                continue
            txt = "".join(node.text or "" for node in child.iter(qn("w:t")))
            if "Proposition financière" not in txt:
                continue
            title_p = None
            for p in doc.paragraphs:
                if p._p is child:
                    title_p = p
                    break
            if title_p is not None:
                set_keep(title_p, keep_next=True)
            # Chainer les paragraphes vides entre le titre et le tableau.
            for next_child in body_children[idx + 1:]:
                if next_child.tag == qn("w:p"):
                    txt2 = "".join(node.text or "" for node in next_child.iter(qn("w:t"))).strip()
                    if txt2:
                        break
                    for p in doc.paragraphs:
                        if p._p is next_child:
                            set_keep(p, keep_next=True)
                            break
                elif next_child.tag == qn("w:tbl"):
                    tbl_pr = next_child.find(qn("w:tblPr"))
                    if tbl_pr is not None:
                        cant_autofit = tbl_pr.find(qn("w:tblLayout"))
                        if cant_autofit is None:
                            cant_autofit = OxmlElement("w:tblLayout")
                            tbl_pr.append(cant_autofit)
                        cant_autofit.set(qn("w:type"), "fixed")
                    for tr in next_child.findall(qn("w:tr")):
                        tr_pr = tr.find(qn("w:trPr"))
                        if tr_pr is None:
                            tr_pr = OxmlElement("w:trPr")
                            tr.insert(0, tr_pr)
                        if tr_pr.find(qn("w:cantSplit")) is None:
                            tr_pr.append(OxmlElement("w:cantSplit"))
                        for p_el in tr.findall(".//" + qn("w:p")):
                            for p in doc.paragraphs:
                                if p._p is p_el:
                                    set_keep(p, keep_next=False)
                                    break
                    break
            break

    def raise_complementary_fixed_images(docx_path: Path) -> None:
        try:
            from lxml import etree
        except Exception:
            return
        ns_rel = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
        ns = {
            "v": "urn:schemas-microsoft-com:vml",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }
        targets = {
            # Styles dedies extraits de la page "Prestations complementaires"
            # du Word source 13_BUREAUX petit (1).doc, converti en DOCX.
            # Ces visuels n'heritent jamais des tailles de photos de chantier.
            "media/image8.jpeg": {
                "style_name": "action_nuisibles_logo",
                "left_pt": 289.30,
                "top_pt": 182.60,
                "target_width_pt": 192.40,
                "target_height_pt": 111.95,
                "wrap_left_pt": 9.05,
                "wrap_right_pt": 9.05,
                "wrap_top_pt": 0.0,
                "wrap_bottom_pt": 0.0,
            },
            "media/image9.png": {
                "style_name": "habilitation_electrique_image",
                "left_pt": 240.70,
                "top_pt": 370.85,
                "target_width_pt": 250.15,
                "target_height_pt": 192.85,
                "wrap_left_pt": 9.05,
                "wrap_right_pt": 9.05,
                "wrap_top_pt": 0.0,
                "wrap_bottom_pt": 0.0,
            },
        }
        tmp_path = docx_path.with_suffix(".images-up.tmp.docx")
        changed = False
        with zipfile.ZipFile(docx_path, "r") as zin:
            rel_map = {}
            try:
                rel_root = etree.fromstring(zin.read("word/_rels/document.xml.rels"))
                for rel in rel_root.xpath(".//rel:Relationship", namespaces=ns_rel):
                    rel_map[rel.get("Id")] = rel.get("Target")
            except Exception:
                rel_map = {}
            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for info in zin.infolist():
                    payload = zin.read(info.filename)
                    if info.filename == "word/document.xml":
                        root = etree.fromstring(payload)
                        for shape in root.xpath(".//v:shape[v:imagedata]", namespaces=ns):
                            image = shape.find("v:imagedata", namespaces=ns)
                            rid = image.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id") if image is not None else None
                            cfg = targets.get(rel_map.get(rid, ""))
                            if not cfg:
                                continue
                            style = shape.get("style") or ""
                            match = re.search(r"margin-top:([\-0-9.]+)pt", style)
                            style_new = style
                            if "top_pt" in cfg:
                                top_value = f"{float(cfg['top_pt']):.2f}".rstrip("0").rstrip(".")
                                if match:
                                    style_new = style_new[:match.start(1)] + top_value + style_new[match.end(1):]
                                else:
                                    style_new = style_new + f";margin-top:{top_value}pt"
                            elif match:
                                old = float(match.group(1))
                                new = max(0.0, old - float(cfg.get("shift_pt", 0)))
                                style_new = style_new[:match.start(1)] + f"{new:.2f}".rstrip("0").rstrip(".") + style_new[match.end(1):]
                            if "left_pt" in cfg:
                                left_match = re.search(r"margin-left:([\-0-9.]+)pt", style_new)
                                left_value = f"{float(cfg['left_pt']):.2f}".rstrip("0").rstrip(".")
                                if left_match:
                                    style_new = style_new[:left_match.start(1)] + left_value + style_new[left_match.end(1):]
                                else:
                                    style_new = style_new + f";margin-left:{left_value}pt"
                            for wrap_key, css_key in (
                                ("wrap_left_pt", "mso-wrap-distance-left"),
                                ("wrap_right_pt", "mso-wrap-distance-right"),
                                ("wrap_top_pt", "mso-wrap-distance-top"),
                                ("wrap_bottom_pt", "mso-wrap-distance-bottom"),
                            ):
                                if wrap_key not in cfg:
                                    continue
                                wrap_value = f"{float(cfg[wrap_key]):.2f}".rstrip("0").rstrip(".")
                                wrap_match = re.search(rf"{re.escape(css_key)}:([\-0-9.]+)pt", style_new)
                                if wrap_match:
                                    style_new = style_new[:wrap_match.start(1)] + wrap_value + style_new[wrap_match.end(1):]
                                else:
                                    style_new = style_new + f";{css_key}:{wrap_value}pt"
                            width_match = re.search(r"width:([0-9.]+)pt", style_new)
                            height_match = re.search(r"height:([0-9.]+)pt", style_new)
                            if cfg.get("target_width_pt") and width_match and height_match:
                                old_width = float(width_match.group(1))
                                old_height = float(height_match.group(1))
                                ratio = old_height / old_width if old_width else 1.0
                                new_width = float(cfg["target_width_pt"])
                                new_height = float(cfg.get("target_height_pt") or (new_width * ratio))
                                width_value = f"{new_width:.2f}".rstrip("0").rstrip(".")
                                height_value = f"{new_height:.2f}".rstrip("0").rstrip(".")
                                style_new = style_new[:width_match.start(1)] + width_value + style_new[width_match.end(1):]
                                height_match = re.search(r"height:([0-9.]+)pt", style_new)
                                if height_match:
                                    style_new = style_new[:height_match.start(1)] + height_value + style_new[height_match.end(1):]
                            else:
                                for prop, max_key in (("width", "max_width_pt"), ("height", "max_height_pt")):
                                    dim = re.search(rf"{prop}:([0-9.]+)pt", style_new)
                                    if not dim:
                                        continue
                                    old_dim = float(dim.group(1))
                                    max_dim = float(cfg.get(max_key, old_dim))
                                    new_dim = min(max_dim, old_dim * float(cfg.get("scale", 1.0)))
                                    style_new = style_new[:dim.start(1)] + f"{new_dim:.2f}".rstrip("0").rstrip(".") + style_new[dim.end(1):]
                            if style_new != style:
                                shape.set("style", style_new)
                                changed = True
                        if changed:
                            payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
                    zout.writestr(info, payload)
        if changed:
            shutil.move(tmp_path, docx_path)
            _nettoyer_doublons_zip(docx_path)
        else:
            tmp_path.unlink(missing_ok=True)

    def protect_bureaux_overflow_page_start() -> None:
        """Evite qu'un bloc dynamique long commence a cote du logo en haut de page."""
        visible = [code for code, _title, freq_var, _aliases in zone_defs if show_zone(code, freq_var)]
        target = None
        for code in visible:
            ops = data.get(f"OPS_{code}")
            if isinstance(ops, list) and len([op for op in ops if str(op).strip()]) > 12:
                target = code
                break
        if target is None and len(visible) >= 5 and "VITRERIE" in visible:
            target = "VITRERIE"
        if target is None and len(visible) >= 6 and "CONSOMMABLES" in visible:
            target = "CONSOMMABLES"
        if target is None:
            return
        ranges_now = find_zone_ranges()
        if target not in ranges_now:
            return
        start, _end = ranges_now[target]
        title = doc.paragraphs[start]
        apply_content_start(title, force_break=True)
        title.paragraph_format.space_after = Pt(14)

    ranges = find_zone_ranges()
    # Le document Word source reste maitre : pour les zones cochees, on ne
    # reconstruit pas les paragraphes. On remplace seulement les variables deja
    # presentes. Les zones non cochees sont retirees en bloc.
    for start, _end, code, title, freq_var in sorted(
        (
            (ranges[code][0], ranges[code][1], code, title, freq_var)
            for code, title, freq_var, _aliases in zone_defs
            if code in ranges and show_zone(code, freq_var)
        ),
        reverse=True,
    ):
        render_zone(code, title, freq_var, start, ranges[code][1])
    ranges = find_zone_ranges()
    for code, _title, freq_var, _aliases in reversed(zone_defs):
        if code in ranges and not show_zone(code, freq_var):
            start, end = ranges[code]
            for p in reversed(doc.paragraphs[start:end + 1]):
                if paragraph_contains_media(p):
                    continue
                remove_paragraph(p)
    trim_empty_runs_between_dynamic_blocks()
    keep_bureaux_short_dynamic_blocks()
    paginate_bureaux_zone_blocks()
    protect_bureaux_overflow_page_start()

    # Pagination portée par le template : vrai départ de CGV.
    for p in doc.paragraphs:
        raw = (p.text or "").strip()
        if norm(raw).startswith("conditions generales de vente") and raw.upper() == raw:
            p.paragraph_format.page_break_before = True
            break

    # Bloc signature du modèle Bureaux : indivisible, sans repositionnement.
    in_signature = False
    signature_block = []
    for p in doc.paragraphs:
        txt = norm(p.text)
        if txt.startswith("toutes les autres clauses contractuelles"):
            in_signature = True
        if in_signature:
            signature_block.append(p)
            if txt == "laurent prevert":
                break
    for idx, p in enumerate(signature_block):
        set_keep(p, keep_next=(idx < len(signature_block) - 1))

    lock_financial_section()
    normalize_presentation_media_wrap()
    tighten_presentation_start()
    lock_presentation_on_single_page()
    constrain_presentation_text_column()
    normalize_bureaux_section_numbers()

    doc.save(str(output_path))
    _nettoyer_doublons_zip(output_path)
    _ajuster_images_prestations_complementaires(output_path)
    _injecter_photos(output_path, data.get("PRESTATIONS_PHOTOS") or data.get("PHOTOS") or [])
    _supprimer_rouge_document(output_path)
    _restaurer_image_couverture_bureaux(template_path, output_path)
    _ajuster_images_prestations_complementaires(output_path)
    _supprimer_surlignages_jaunes_document(output_path)
    return output_path


def _generer_copro_petite_preserve_layout(template_path: Path, data: Dict, output_path: Path) -> Path:
    """Rend le modèle Copro petite sans reconstruire le document avec docxtpl.

    Le fichier de référence contient beaucoup d'images fixes et d'ancrages Word
    hérités du .doc original. Le rendu docxtpl peut déplacer/supprimer certains
    de ces objets. Ici on conserve le document maître et on remplace seulement
    les variables simples, puis on développe le tableau financier.
    """
    from docx import Document as _Doc

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(template_path, output_path)

    doc = _Doc(str(output_path))
    photos = data.get("PRESTATIONS_PHOTOS") or data.get("PHOTOS") or []
    zone_defs = [
        ("HALL", "Hall d'Entrée", ("hall", "hall entree", "hall d entree")),
        ("ASCENSEUR", "Nettoyage de la cabine d’ascenseur", ("ascenseur", "cabine ascenseur")),
        ("ESCALIERS", "Cages d’escaliers et paliers", ("escaliers", "cages escaliers", "paliers")),
        ("CAVES", "Caves et descentes garage", ("caves", "descentes garage")),
        ("GARAGE", "Garage", ("garage",)),
        ("ABORDS", "Abords & accès", ("abords", "acces")),
        ("CONTENEUR", "Local conteneur / vide ordure", ("conteneur", "vide ordure", "local conteneur")),
        ("OM", "Ordures Ménagères", ("ordures", "ordures menageres")),
    ]

    def scalar_value(name: str) -> str:
        val = data.get(name, "")
        if isinstance(val, (list, dict, tuple)):
            return ""
        return str(val)

    def replace_in_runs(paragraph):
        for run in paragraph.runs:
            txt = run.text or ""
            if "{{" not in txt:
                continue
            def repl(match):
                expr = match.group(1).strip()
                if expr.startswith("opt."):
                    return match.group(0)
                return scalar_value(expr)
            run.text = re.sub(r"\{\{\s*([^}]+?)\s*\}\}", repl, txt)

    def norm(txt: str) -> str:
        import unicodedata
        txt = (txt or "").lower().replace("\xa0", " ")
        txt = unicodedata.normalize("NFKD", txt)
        txt = "".join(ch for ch in txt if not unicodedata.combining(ch))
        txt = re.sub(r"[\W_]+", " ", txt, flags=re.UNICODE)
        return re.sub(r"\s+", " ", txt).strip()

    def remove_paragraph(paragraph):
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)

    def show_zone(code: str) -> bool:
        key = f"SHOW_{code}"
        if key in data:
            return bool(data.get(key))
        freq = str(data.get(f"FREQ_{code}", "") or "").strip()
        return bool(freq)

    def find_zone_ranges():
        starts = []
        for code, title, _aliases in zone_defs:
            candidates = [norm(title), *(norm(a) for a in _aliases)]
            for i, p in enumerate(doc.paragraphs):
                hay = norm(p.text)
                if hay and any(c and (hay == c or (len(c) > 8 and c in hay)) for c in candidates):
                    starts.append((i, code))
                    break
        starts.sort()
        ranges = {}
        for pos, (start, code) in enumerate(starts):
            end = starts[pos + 1][0] - 1 if pos + 1 < len(starts) else start
            # La dernière zone s'arrête à sa ligne Fréquence, pour ne pas toucher
            # aux pages fixes qui suivent dans le modèle Word.
            for j in range(start, min(len(doc.paragraphs), start + 40)):
                if f"FREQ_{code}" in (doc.paragraphs[j].text or ""):
                    end = j
                    break
            ranges[code] = (start, end)
        return ranges

    def apply_zone_visibility():
        to_remove = []
        for code, (start, end) in find_zone_ranges().items():
            if show_zone(code):
                continue
            to_remove.extend(doc.paragraphs[start:end + 1])
        seen = set()
        for p in to_remove:
            key = id(p._element)
            if key in seen:
                continue
            seen.add(key)
            remove_paragraph(p)

    def filter_photos_for_visible_zones(raw_photos):
        visible_codes = {code for code, _title, _aliases in zone_defs if show_zone(code)}
        if not visible_codes:
            return []

        def detect_zone_code(label: str):
            label_n = norm(label)
            if not label_n:
                return None
            for code, title, aliases in zone_defs:
                title_n = norm(title)
                alias_ns = [norm(a) for a in aliases]
                if label_n == title_n or title_n in label_n:
                    return code
                for alias in alias_ns:
                    if not alias:
                        continue
                    if code == "GARAGE":
                        if label_n == alias:
                            return code
                    elif len(alias) > 3 and alias in label_n:
                        return code
            return None

        kept = []
        for ph in raw_photos or []:
            if not isinstance(ph, dict):
                continue
            label = norm(str(ph.get("libelle", "") or ph.get("prestation", "")))
            detected = detect_zone_code(label)
            if detected in visible_codes:
                kept.append(ph)
        return kept

    def _flag_to_bool(value) -> bool:
        if isinstance(value, bool):
            return value
        if isinstance(value, (int, float)):
            return value != 0
        return str(value).strip().lower() not in {"0", "false", "faux", "non", "no", "off"}

    def apply_visible_zone_operation_filters():
        """Supprime les prestations recurrentes decochees dans les blocs fixes Copro."""
        ranges = find_zone_ranges()
        for code, (start, end) in ranges.items():
            if not show_zone(code):
                continue
            flags_raw = data.get(f"OPS_ENABLED_{code}")
            flags = [_flag_to_bool(v) for v in flags_raw] if isinstance(flags_raw, list) else None
            active_raw = data.get(f"OPS_{code}")
            active_norms = (
                {norm(str(op)) for op in active_raw if str(op).strip()}
                if isinstance(active_raw, list)
                else None
            )
            if flags is None and active_norms is None:
                continue

            op_index = 0
            for p in list(doc.paragraphs[start + 1:end + 1]):
                txt_raw = p.text or ""
                txt = norm(txt_raw)
                if not txt:
                    continue
                if "frequence" in txt or f"freq {code.lower()}" in txt:
                    continue
                keep = True
                if flags is not None:
                    if op_index < len(flags):
                        keep = flags[op_index]
                    elif active_norms is not None and not active_norms:
                        keep = False
                elif active_norms is not None:
                    keep = txt in active_norms
                op_index += 1
                if not keep:
                    remove_paragraph(p)

    def trim_empty_paragraphs_after_selected_zones():
        """Retire les paragraphes vides devenus inutiles dans la zone prestations."""
        selected_ranges = [rng for code, rng in find_zone_ranges().items() if show_zone(code)]
        if not selected_ranges:
            return
        last_end = max(end for _start, end in selected_ranges)
        # Après suppression, les indices peuvent avoir bougé : on repart du texte.
        trailing = []
        in_dynamic_tail = False
        for p in list(doc.paragraphs):
            txt = norm(p.text)
            if any(norm(title) == txt or (len(norm(title)) > 8 and norm(title) in txt)
                   for code, title, _aliases in zone_defs if show_zone(code)):
                in_dynamic_tail = True
                trailing = []
                continue
            if txt and any(marker in txt for marker in (
                "prestations complementaires", "tracabilite", "proposition financiere",
                "conditions generales"
            )):
                break
            if in_dynamic_tail:
                if txt:
                    trailing = []
                else:
                    trailing.append(p)
        for p in trailing:
            remove_paragraph(p)

    def paginate_visible_zone_blocks():
        """Evite qu'un bloc zone/frequence commence ou finisse trop haut."""
        from docx.shared import Pt

        ranges = sorted(
            (start, end, code)
            for code, (start, end) in find_zone_ranges().items()
            if show_zone(code)
        )
        if not ranges:
            return

        def line_count(text: str) -> int:
            text = (text or "").strip()
            if not text:
                return 0
            return max(1, (len(text) // 82) + 1)

        def block_height(start: int, end: int) -> int:
            height = 0
            for p in doc.paragraphs[start:end + 1]:
                height += line_count(p.text)
            return max(3, height)

        used = 8  # titre "2.", intro et sous-titre deja presents sur la premiere page.
        capacity = 39
        next_capacity = 39
        has_block_on_page = False
        for start, end, _code in ranges:
            height = block_height(start, end)
            if has_block_on_page and used + height > capacity:
                title = doc.paragraphs[start]
                title.paragraph_format.page_break_before = True
                title.paragraph_format.space_before = Pt(36)
                title.paragraph_format.space_after = Pt(0)
                used = height
                capacity = next_capacity
            else:
                used += height
            has_block_on_page = True

    def replace_paragraph_text(paragraph, text: str):
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)

    def normalize_cover_date() -> None:
        date_longue = scalar_value("DATE_EMISSION_LONGUE")
        if not date_longue:
            return
        for p in iter_paragraphs():
            txt = norm(p.text)
            if txt.startswith("marseille le") or txt.startswith("a marseille le"):
                replace_paragraph_text(p, date_longue)

    def apply_bureau_zone_content():
        """Conserve le gabarit Copro mais remplace les blocs dynamiques par du bureau."""
        modele = str(data.get("MODELE_CODE", "") or "").lower()
        bureau_blocks = {
            "bureaux_petit": {
                "HALL": {
                    "title": "Bureaux et postes de travail",
                    "ops": [
                        "Dépoussiérage des bureaux, plans de travail et surfaces accessibles",
                        "Vidage des corbeilles et remplacement des sacs si nécessaire",
                        "Aspiration et lavage des sols des bureaux",
                        "Nettoyage des points de contact : poignées, interrupteurs et portes",
                    ],
                },
                "ASCENSEUR": {
                    "title": "Sanitaires et kitchenette",
                    "ops": [
                        "Nettoyage des sanitaires et points d'eau",
                        "Réapprovisionnement si prévu au devis",
                        "Nettoyage de la kitchenette ou zone café si présente",
                    ],
                },
            },
            "bureaux_important": {
                "HALL": {
                    "title": "Bureaux, open spaces et salles de réunion",
                    "ops": [
                        "Dépoussiérage complet des bureaux, postes de travail et surfaces accessibles",
                        "Vidage des corbeilles, tri et évacuation des déchets courants",
                        "Aspiration et lavage des sols des bureaux et salles de réunion",
                        "Nettoyage des points de contact : poignées, interrupteurs, portes et claviers accessibles",
                    ],
                },
                "ASCENSEUR": {
                    "title": "Sanitaires, points d'eau et zones de pause",
                    "ops": [
                        "Nettoyage approfondi des sanitaires et points d'eau",
                        "Réapprovisionnement si prévu au devis",
                        "Nettoyage de la kitchenette, des plans de travail et zones de pause",
                    ],
                },
                "ESCALIERS": {
                    "title": "Circulations internes et accès bureaux",
                    "ops": [
                        "Aspiration et lavage des circulations internes",
                        "Nettoyage des portes, poignées et interrupteurs des circulations",
                        "Essuyage des vitrages intérieurs accessibles si prévu au devis",
                    ],
                },
            },
        }.get(modele)
        if not bureau_blocks:
            return

        ranges = find_zone_ranges()
        range_paragraphs = {
            code: list(doc.paragraphs[start:end + 1])
            for code, (start, end) in ranges.items()
        }
        for code, block in bureau_blocks.items():
            if code not in range_paragraphs or not show_zone(code):
                continue
            paras = range_paragraphs[code]
            if not paras:
                continue
            replace_paragraph_text(paras[0], block["title"])
            body_paras = paras[1:]
            desired = list(block["ops"])
            freq = str(data.get(f"FREQ_{code}", "") or "").strip()
            if freq:
                desired.append(f"Fréquence : {freq}")
            for idx, text in enumerate(desired):
                if idx < len(body_paras):
                    replace_paragraph_text(body_paras[idx], text)
            for extra in body_paras[len(desired):]:
                remove_paragraph(extra)

    def normalize_bureau_cgv_page_break():
        """Evite une page filigrane vide avant les CGV sur les contrats bureaux."""
        modele = str(data.get("MODELE_CODE", "") or "").lower()
        if modele not in {"bureaux_petit", "bureaux_important"}:
            return
        for p in doc.paragraphs:
            if not norm(p.text).startswith("conditions generales"):
                continue
            for br in list(p._element.xpath('.//w:br[@w:type="page"]')):
                parent = br.getparent()
                if parent is not None:
                    parent.remove(br)
            p.paragraph_format.page_break_before = False
            return

    def lock_fixed_section_starts():
        """Force les sections officielles à commencer sur une nouvelle page."""
        fixed_titles = ("prestations complementaires",)
        for p in doc.paragraphs:
            txt = norm(p.text)
            if any(marker in txt for marker in fixed_titles):
                p.paragraph_format.page_break_before = True

    def iter_paragraphs():
        for p in doc.paragraphs:
            yield p
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
        for section in doc.sections:
            for part in (section.header, section.first_page_header,
                         section.footer, section.first_page_footer):
                for p in part.paragraphs:
                    yield p
                for table in part.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                yield p

    def row_text(row) -> str:
        return "\n".join(cell.text for cell in row.cells)

    def fill_option_row(row, opt):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    txt = run.text or ""
                    txt = txt.replace("{{ opt.libelle }}", str(opt.get("libelle", "")))
                    txt = txt.replace("{{ opt.ht }}", str(opt.get("ht", "")))
                    txt = txt.replace("{{ opt.tva }}", str(opt.get("tva", "")))
                    txt = txt.replace("{{ opt.ttc }}", str(opt.get("ttc", "")))
                    run.text = txt

    def expand_options_tables():
        options = data.get("OPTIONS") or []
        if not isinstance(options, list):
            options = []
        for table in doc.tables:
            rows = list(table.rows)
            for idx, row in enumerate(rows):
                if "{%tr for opt in OPTIONS %}" not in row_text(row):
                    continue
                if idx + 2 >= len(rows):
                    continue
                template_row = rows[idx + 1]
                template_tr = copy.deepcopy(template_row._tr)
                tbl = table._tbl
                opts = [opt for opt in options if isinstance(opt, dict)] or [{}]
                fill_option_row(template_row, opts[0])
                if idx + 2 < len(rows):
                    tbl.remove(rows[idx + 2]._tr)
                tbl.remove(rows[idx]._tr)
                for opt in opts[1:]:
                    new_tr = copy.deepcopy(template_tr)
                    tbl.append(new_tr)
                    fill_option_row(table.rows[-1], opt)
                return

    def insert_free_prestations():
        """Insère les prestations LIBRES (descriptives, SANS prix) dans le détail
        de chaque zone visible, juste avant la ligne Fréquence. Le format est cloné
        d'une puce existante de la zone pour préserver strictement la mise en page.
        Ces libellés sont purement descriptifs : ils n'entrent jamais dans le calcul
        ni dans la proposition financière (voir routes_devis : LIBRE_* est distinct
        des lignes financières)."""
        W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        title_by_code = {code: title for code, title, _al in zone_defs}
        for code, (start, end) in find_zone_ranges().items():
            if not show_zone(code):
                continue
            libres = data.get(f"LIBRE_{code}")
            if not isinstance(libres, list) or not libres:
                continue
            zone_paras = doc.paragraphs[start:end + 1]
            title_norm = norm(title_by_code.get(code, ""))
            freq_para = None
            ref_bullet = None
            for p in zone_paras:
                t = norm(p.text)
                if not t:
                    continue
                if "frequence" in t or f"freq {code.lower()}" in t:
                    freq_para = p
                    break
                if t != title_norm and (not title_norm or title_norm not in t):
                    ref_bullet = p
            anchor = freq_para or (zone_paras[-1] if zone_paras else None)
            if anchor is None:
                continue
            prefix = ""
            if ref_bullet is not None:
                m = re.match(r"^[\s\-–• ]+", ref_bullet.text or "")
                prefix = m.group(0) if m else ""
            # On insère À LA SUITE de la dernière prestation existante (et non avant
            # « Fréquence »). Sinon le paragraphe d'espacement qui précède la fréquence
            # se retrouve intercalé -> saut de ligne parasite + fréquence collée.
            # Résultat visé : ...dernière puce -> libres -> [espacement] -> Fréquence.
            cursor = ref_bullet._element if ref_bullet is not None else None
            for op in libres:
                op = str(op).strip()
                if not op:
                    continue
                if cursor is not None:
                    new_el = copy.deepcopy(ref_bullet._element)
                    ts = list(new_el.iter(W + "t"))
                    if ts:
                        ts[0].text = prefix + op
                        for t in ts[1:]:
                            t.text = ""
                        cursor.addnext(new_el)
                        cursor = new_el
                        continue
                anchor.insert_paragraph_before(f"-   {op}")

    expand_options_tables()
    apply_zone_visibility()
    apply_visible_zone_operation_filters()
    insert_free_prestations()
    trim_empty_paragraphs_after_selected_zones()
    paginate_visible_zone_blocks()
    apply_bureau_zone_content()
    normalize_bureau_cgv_page_break()
    lock_fixed_section_starts()
    photos = filter_photos_for_visible_zones(photos)
    for p in iter_paragraphs():
        replace_in_runs(p)
    normalize_cover_date()

    doc.save(str(output_path))
    _aligner_depart_pages_copro(output_path)
    _injecter_photos(output_path, photos)
    _supprimer_rouge_document(output_path)
    _supprimer_surlignages_jaunes_document(output_path)
    return output_path


def _aligner_depart_pages_copro(docx_path: Path):
    """Aligne le début des pages Copro sur l'ancrage vertical du modèle.

    Les anciens .doc convertis en .docx contiennent des titres que python-docx
    expose mal selon les révisions Word. On corrige donc directement le XML en
    ajoutant une marge haute stable aux titres qui démarraient trop haut.
    """
    try:
        from lxml import etree
    except Exception:
        return

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    qn_w = lambda name: f"{{{ns['w']}}}{name}"
    titre_prestations_complementaires = "3 - Prestations complémentaires"
    titres_a_abaisser = {
        # Reglage dedie a la page "Presentation" de copro_petite : ancienne
        # valeur 280 twips remontee d'une ligne de 12 pt, soit 240 twips.
        "1 - Présentation de la société Marie Eugénie": COPRO_PETITE_COMPANY_PRESENTATION_START_TWIPS,
        "2. Détail et fréquences des prestations": "360",
        "CONDITIONS GENERALES DE VENTE": "360",
    }
    titre_bloc_coordonnees_couverture = "Société Marie Eugénie"

    def texte_para(p):
        return "".join(p.xpath(".//w:t/text()", namespaces=ns)).strip()

    def est_bloc_coordonnees_couverture(paragraphs, index):
        suivants = [texte_para(candidate) for candidate in paragraphs[index + 1:index + 6]]
        return (
            len(suivants) >= 5
            and suivants[0] == "Représentée par"
            and suivants[1] == "Laurent PREVERT"
            and "lprevert@marie-eugenie.fr" in suivants[2]
            and suivants[3] == "1 rue Raspail 13004 Marseille"
            and "04 91 47 14 38" in suivants[4]
        )

    def compenser_hauteur_apres_coordonnees(paragraphs, start_index):
        telephone_seen = False
        for candidate in paragraphs[start_index + 1:]:
            txt = texte_para(candidate)
            if txt.startswith("1 - Présentation de la société"):
                return False
            if "04 91 47 14 38" in txt:
                telephone_seen = True
                continue
            if not telephone_seen or txt:
                continue
            ppr = candidate.find("w:pPr", namespaces=ns)
            if ppr is None:
                ppr = etree.Element(qn_w("pPr"))
                candidate.insert(0, ppr)
            spacing = ppr.find("w:spacing", namespaces=ns)
            if spacing is None:
                spacing = etree.Element(qn_w("spacing"))
                ppr.append(spacing)
            spacing.set(qn_w("before"), "0")
            spacing.set(qn_w("after"), "0")
            spacing.set(qn_w("line"), COPRO_COVER_AFTER_CONTACT_LINE_TWIPS)
            spacing.set(qn_w("lineRule"), "exact")
            for size in candidate.xpath(".//w:sz | .//w:szCs", namespaces=ns):
                size.set(qn_w("val"), str(int(COPRO_COVER_AFTER_CONTACT_LINE_TWIPS) // 10))
            return True
        return False

    tmp_path = docx_path.with_suffix(".anchor.tmp.docx")
    changed = False
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/document.xml":
                    root = etree.fromstring(data)
                    paragraphs = root.xpath(".//w:p", namespaces=ns)
                    for idx, p in enumerate(paragraphs):
                        txt = texte_para(p)
                        if titre_prestations_complementaires in txt:
                            parent = p.getparent()
                            if parent is not None:
                                ppr = p.find("w:pPr", namespaces=ns)
                                if ppr is None:
                                    ppr = etree.Element(qn_w("pPr"))
                                    p.insert(0, ppr)
                                for page_break in list(ppr.findall("w:pageBreakBefore", namespaces=ns)):
                                    ppr.remove(page_break)
                                spacing = ppr.find("w:spacing", namespaces=ns)
                                if spacing is None:
                                    spacing = etree.Element(qn_w("spacing"))
                                    ppr.append(spacing)
                                spacing.set(qn_w("before"), "0")
                                spacing.set(qn_w("after"), "0")

                                previous = p.getprevious()
                                if (
                                    previous is not None
                                    and not texte_para(previous)
                                    and previous.find(".//w:pageBreakBefore", namespaces=ns) is not None
                                ):
                                    parent.remove(previous)

                                # Le titre "3" est le premier paragraphe apres le saut de page.
                                # Dans LibreOffice/PDF, un simple space_before reste dans la zone
                                # d'habillage du logo flottant et le titre se cale a droite du logo.
                                # On insere donc un vrai paragraphe d'ancrage invisible, avec une
                                # ligne non imprimante, pour faire demarrer le bloc sous le logo.
                                def build_spacer(with_page_break=False):
                                    spacer = etree.Element(qn_w("p"))
                                    spacer_ppr = etree.SubElement(spacer, qn_w("pPr"))
                                    if with_page_break:
                                        etree.SubElement(spacer_ppr, qn_w("pageBreakBefore"))
                                    spacer_spacing = etree.SubElement(spacer_ppr, qn_w("spacing"))
                                    spacer_spacing.set(qn_w("before"), "0")
                                    spacer_spacing.set(qn_w("after"), "0")
                                    spacer_r = etree.SubElement(spacer, qn_w("r"))
                                    spacer_rpr = etree.SubElement(spacer_r, qn_w("rPr"))
                                    spacer_sz = etree.SubElement(spacer_rpr, qn_w("sz"))
                                    spacer_sz.set(qn_w("val"), "24")
                                    spacer_color = etree.SubElement(spacer_rpr, qn_w("color"))
                                    spacer_color.set(qn_w("val"), "FFFFFF")
                                    spacer_text = etree.SubElement(spacer_r, qn_w("t"))
                                    spacer_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                                    spacer_text.text = " "
                                    return spacer

                                insert_at = parent.index(p)
                                parent.insert(insert_at, build_spacer(with_page_break=True))
                                parent.insert(insert_at + 1, build_spacer())
                                parent.insert(insert_at + 2, build_spacer())
                                changed = True
                            continue
                        before_value = next(
                            (value for title, value in titres_a_abaisser.items() if title in txt),
                            None,
                        )
                        if txt == titre_bloc_coordonnees_couverture and est_bloc_coordonnees_couverture(paragraphs, idx):
                            ppr = p.find("w:pPr", namespaces=ns)
                            if ppr is None:
                                ppr = etree.Element(qn_w("pPr"))
                                p.insert(0, ppr)
                            spacing = ppr.find("w:spacing", namespaces=ns)
                            if spacing is None:
                                spacing = etree.Element(qn_w("spacing"))
                                ppr.append(spacing)
                            spacing.attrib.pop(qn_w("beforeLines"), None)
                            spacing.set(qn_w("before"), COPRO_COVER_CONTACT_BLOCK_BEFORE_TWIPS)
                            spacing.set(qn_w("after"), "0")
                            compenser_hauteur_apres_coordonnees(paragraphs, idx)
                            changed = True
                            continue
                        if before_value is None:
                            continue
                        ppr = p.find("w:pPr", namespaces=ns)
                        if ppr is None:
                            ppr = etree.Element(qn_w("pPr"))
                            p.insert(0, ppr)
                        spacing = ppr.find("w:spacing", namespaces=ns)
                        if spacing is None:
                            spacing = etree.Element(qn_w("spacing"))
                            ppr.append(spacing)
                        # 240 twips = 12 pt, 360 twips = 18 pt. Ces valeurs
                        # replacent les titres sur l'ancrage visuel des pages
                        # 4/5 sans créer de saut de page supplémentaire.
                        spacing.set(qn_w("before"), before_value)
                        spacing.set(qn_w("after"), "0")
                        changed = True
                    for idx, p in enumerate(paragraphs):
                        txt = texte_para(p)
                        if not txt.startswith("- Mono brosses"):
                            continue
                        if idx > 0 and not texte_para(paragraphs[idx - 1]):
                            parent = paragraphs[idx - 1].getparent()
                            if parent is not None:
                                parent.remove(paragraphs[idx - 1])
                                changed = True
                        break
                    if changed:
                        data = etree.tostring(
                            root, xml_declaration=True, encoding="UTF-8", standalone="yes"
                        )
                zout.writestr(info, data)
    if changed:
        shutil.move(tmp_path, docx_path)
    else:
        tmp_path.unlink(missing_ok=True)



def _normaliser_interligne_document(docx_path: Path):
    """Force l'interligne simple et les espacements avant/apres a 0 pt."""
    from docx import Document as _Doc
    from docx.shared import Pt
    from docx.oxml.ns import qn

    d = _Doc(str(docx_path))

    def iter_paragraphes():
        for p in d.paragraphs:
            yield p
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
        for section in d.sections:
            for part in (section.header, section.first_page_header,
                         section.footer, section.first_page_footer):
                for p in part.paragraphs:
                    yield p
                for table in part.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                yield p

    for p in iter_paragraphes():
        fmt = p.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1
        p_pr = p._p.get_or_add_pPr()
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is not None:
            spacing.set(qn("w:before"), "0")
            spacing.set(qn("w:after"), "0")
            spacing.attrib.pop(qn("w:line"), None)
            spacing.attrib.pop(qn("w:lineRule"), None)
        for child in list(p_pr):
            if child.tag.endswith("}contextualSpacing"):
                p_pr.remove(child)

    for style in d.styles:
        try:
            fmt = getattr(style, "paragraph_format", None)
            if fmt:
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(0)
                fmt.line_spacing = 1
                p_pr = style.element.get_or_add_pPr()
                spacing = p_pr.find(qn("w:spacing"))
                if spacing is not None:
                    spacing.set(qn("w:before"), "0")
                    spacing.set(qn("w:after"), "0")
                    spacing.attrib.pop(qn("w:line"), None)
                    spacing.attrib.pop(qn("w:lineRule"), None)
        except Exception:
            pass

    d.save(str(docx_path))
    _nettoyer_xml_mise_en_page(docx_path)


def _mettre_en_page_devis_ponctuel(docx_path: Path):
    """Applique une mise en page plus aeree aux devis ponctuels."""
    from docx import Document as _Doc
    from docx.shared import Pt, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    d = _Doc(str(docx_path))
    rythme = RYTHME_COPRO_REFERENCE

    for section in d.sections:
        section.top_margin = Mm(rythme["page_top_mm"])

    def norm(txt):
        import unicodedata
        txt = (txt or "").lower().replace("\xa0", " ")
        txt = unicodedata.normalize("NFKD", txt)
        txt = "".join(ch for ch in txt if not unicodedata.combining(ch))
        return re.sub(r"\s+", " ", txt).strip()

    def set_font(paragraph, size=10.5):
        for run in paragraph.runs:
            if not run.text:
                continue
            run.font.name = "Arial"
            run.font.size = Pt(size)
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            for attr in ("ascii", "hAnsi", "cs"):
                rfonts.set(qn(f"w:{attr}"), "Arial")

    def paragraphes():
        for p in d.paragraphs:
            yield p
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    in_cgv = False
    for p in paragraphes():
        txt = (p.text or "").strip()
        n = norm(txt)
        if n.startswith("conditions generales"):
            in_cgv = True
        pf = p.paragraph_format
        pf.line_spacing = rythme["cgv_line_spacing"] if in_cgv else 1
        pf.space_before = Pt(0)
        pf.space_after = Pt(rythme["cgv_body_after_pt"] if in_cgv else rythme["body_after_pt"])
        set_font(p, 9 if in_cgv else 10)

        if not txt:
            continue

        is_main_title = n.startswith("devis ")
        is_section_title = (
            n.startswith("1 -")
            or n.startswith("2-")
            or n.startswith("2 -")
            or n.startswith("3 -")
            or n.startswith("conditions generales")
            or n.startswith("article ")
            or n == "information :"
        )
        if is_main_title:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(12)
            pf.space_after = Pt(5)
            pf.keep_with_next = True
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(14)
        elif is_section_title:
            pf.space_before = Pt(rythme["cgv_title_before_pt"] if in_cgv else rythme["section_before_pt"])
            pf.space_after = Pt(rythme["cgv_title_after_pt"] if in_cgv else rythme["section_after_pt"])
            pf.keep_with_next = True
            if not in_cgv and n.startswith(("2-", "2 -", "3 -")):
                pf.space_before = Pt(18)
            if n.startswith("conditions generales"):
                pf.space_before = Pt(rythme["cgv_cover_before_pt"])
                pf.space_after = Pt(rythme["cgv_cover_after_pt"])
            for run in p.runs:
                if run.text.strip():
                    run.bold = True
        elif txt.startswith("-"):
            pf.left_indent = Mm(8)
            pf.first_line_indent = Mm(-3)
            pf.space_after = Pt(2)
        elif "Suite à votre demande" in txt or "Suite a votre demande" in txt:
            pf.space_before = Pt(8)
            pf.space_after = Pt(6)
        elif "Compte tenu" in txt:
            pf.space_before = Pt(8)
            pf.space_after = Pt(4)
            for run in p.runs:
                run.bold = True

    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_mar = tc_pr.first_child_found_in("w:tcMar")
                if tc_mar is None:
                    tc_mar = OxmlElement("w:tcMar")
                    tc_pr.append(tc_mar)
                for side, width in (("top", "60"), ("bottom", "60"), ("left", "80"), ("right", "80")):
                    mar = tc_mar.find(qn(f"w:{side}"))
                    if mar is None:
                        mar = OxmlElement(f"w:{side}")
                        tc_mar.append(mar)
                    mar.set(qn("w:w"), width)
                    mar.set(qn("w:type"), "dxa")
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1
                    set_font(p)

    d.save(str(docx_path))
    _nettoyer_doublons_zip(docx_path)


def _normaliser_entete_ponctuel(docx_path: Path):
    """Normalise l'en-tete de premiere page des devis ponctuels.

    La regle reste commune a tous les ponctuels : numero centre, logo/image
    d'en-tete aligne a droite et legerement descendu, sans toucher aux CGV ni
    aux contrats recurrents.
    """
    from docx import Document as _Doc
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    d = _Doc(str(docx_path))

    def norm(txt: str) -> str:
        import unicodedata
        txt = (txt or "").lower().replace("\xa0", " ")
        txt = unicodedata.normalize("NFKD", txt)
        txt = "".join(ch for ch in txt if not unicodedata.combining(ch))
        txt = re.sub(r"[\W_]+", " ", txt, flags=re.UNICODE)
        return re.sub(r"\s+", " ", txt).strip()

    def contains_media(paragraph) -> bool:
        return bool(
            paragraph._element.xpath(
                ".//*[local-name()='drawing' or local-name()='pict' or local-name()='shape' or local-name()='imagedata']"
            )
        )

    changed = False
    before_detail = True
    logo_done = False
    numero_seen = False
    for p in d.paragraphs[:45]:
        n = norm(p.text)
        if "detail des prestations" in n or "proposition financiere" in n:
            before_detail = False
        if before_detail and not numero_seen and n:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = None
            p.paragraph_format.right_indent = None
            p.paragraph_format.first_line_indent = None
            changed = True
        if before_detail and n.startswith("devis "):
            numero_seen = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            for run in p.runs:
                if run.text.strip():
                    run.bold = True
                    run.underline = True
            changed = True
        if before_detail and not logo_done and contains_media(p):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            ppr = p._p.get_or_add_pPr()
            if ppr.find(qn("w:keepNext")) is None:
                ppr.append(OxmlElement("w:keepNext"))
            logo_done = True
            changed = True

    if changed:
        d.save(str(docx_path))
        _nettoyer_doublons_zip(docx_path)

    _positionner_logo_premiere_page_ponctuel(docx_path)


def _positionner_logo_premiere_page_ponctuel(docx_path: Path) -> None:
    """Place le logo de premiere page ponctuelle a droite via le XML Word.

    Plusieurs sources ponctuelles stockent le logo dans l'en-tete en VML
    (`w:pict/v:shape`). Un simple alignement de paragraphe ne deplace donc rien.
    Cette passe centralisee ajuste uniquement le header de premiere page.
    """
    import tempfile
    import zipfile
    from xml.etree import ElementTree as ET

    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
        "v": "urn:schemas-microsoft-com:vml",
    }
    ET.register_namespace("w", ns["w"])
    ET.register_namespace("r", ns["r"])
    ET.register_namespace("v", ns["v"])

    def style_to_dict(value: str) -> dict:
        out = {}
        for part in (value or "").split(";"):
            if ":" not in part:
                continue
            key, val = part.split(":", 1)
            out[key.strip()] = val.strip()
        return out

    def dict_to_style(style: dict) -> str:
        preferred = [
            "position", "margin-left", "margin-top", "width", "height",
            "z-index", "visibility", "mso-position-horizontal-relative",
            "mso-position-vertical-relative", "mso-width-relative",
            "mso-height-relative",
        ]
        keys = preferred + [k for k in style.keys() if k not in preferred]
        return ";".join(f"{k}:{style[k]}" for k in keys if k in style)

    def pt_value(value: str):
        m = re.search(r"(-?\d+(?:[.,]\d+)?)pt", value or "")
        return float(m.group(1).replace(",", ".")) if m else None

    try:
        with zipfile.ZipFile(docx_path) as zin:
            document_xml = ET.fromstring(zin.read("word/document.xml"))
            rels_xml = ET.fromstring(zin.read("word/_rels/document.xml.rels"))
            sect = document_xml.find(".//w:sectPr", ns)
            if sect is None:
                return
            first_header = sect.find("./w:headerReference[@w:type='first']", ns)
            if first_header is None:
                return
            rel_id = first_header.get(f"{{{ns['r']}}}id")
            rel_target = None
            for rel in rels_xml:
                if rel.get("Id") == rel_id:
                    rel_target = rel.get("Target")
                    break
            if not rel_target:
                return
            header_name = "word/" + rel_target.lstrip("/")
            header_xml = ET.fromstring(zin.read(header_name))

            pg_sz = sect.find("./w:pgSz", ns)
            pg_mar = sect.find("./w:pgMar", ns)
            page_w_pt = float(pg_sz.get(f"{{{ns['w']}}}w", "11906")) / 20.0
            left_pt = float(pg_mar.get(f"{{{ns['w']}}}left", "1418")) / 20.0
            right_pt = float(pg_mar.get(f"{{{ns['w']}}}right", "1418")) / 20.0
            content_w_pt = page_w_pt - left_pt - right_pt

            changed = False
            parent_map = {child: parent for parent in header_xml.iter() for child in parent}
            shapes = []
            for shape in header_xml.findall(".//v:shape", ns):
                if shape.find("./v:imagedata", ns) is None:
                    continue
                style = style_to_dict(shape.get("style", ""))
                width = pt_value(style.get("width", ""))
                height = pt_value(style.get("height", ""))
                if width and height:
                    shapes.append((shape, style, width, height))

            # Les anciens ponctuels ont souvent un bandeau horizontal VML dans le
            # meme header. Il chevauche le logo quand celui-ci est place a droite.
            for shape, style, width, height in list(shapes):
                if width > 280 and height < 40:
                    parent = parent_map.get(shape)
                    if parent is not None:
                        parent.remove(shape)
                        changed = True

            logo_candidates = [
                (shape, style, width, height)
                for shape, style, width, height in shapes
                if 80 <= width <= 180 and 55 <= height <= 150
            ]
            if not logo_candidates:
                return
            shape, style, width, _height = max(logo_candidates, key=lambda item: item[2] * item[3])
            new_left = max(0.0, content_w_pt - width)
            current_left = pt_value(style.get("margin-left", ""))
            current_top = pt_value(style.get("margin-top", ""))
            if current_left is None or abs(current_left - new_left) > 0.5:
                style["margin-left"] = f"{new_left:.1f}pt"
                changed = True
            # Relative a la marge haute : -22 pt place le haut du logo environ
            # 17 mm sous le bord de page sur A4 avec les marges actuelles.
            new_top = -22.0
            if current_top is None or abs(current_top - new_top) > 0.5:
                style["margin-top"] = f"{new_top:.1f}pt"
                changed = True
            style["mso-position-horizontal-relative"] = "margin"
            style["mso-position-vertical-relative"] = "margin"
            shape.set("style", dict_to_style(style))

            if not changed:
                return

            tmp = Path(tempfile.mkstemp(suffix=".docx", dir=str(docx_path.parent))[1])
            try:
                with zipfile.ZipFile(docx_path, "r") as zin2, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
                    for item in zin2.infolist():
                        if item.filename == header_name:
                            zout.writestr(item, ET.tostring(header_xml, encoding="utf-8", xml_declaration=True))
                        else:
                            zout.writestr(item, zin2.read(item.filename))
                shutil.copyfile(str(tmp), str(docx_path))
            finally:
                if tmp.exists():
                    tmp.unlink(missing_ok=True)
    except Exception as exc:
        print(f"[warn] positionnement logo ponctuel ignore : {exc}")


def _normaliser_cgv_ponctuel(docx_path: Path) -> None:
    """Applique une mise en page CGV commune a tous les devis ponctuels.

    Les CGV des modeles ponctuels proviennent de sources Word differentes. Cette
    passe ne touche pas au texte juridique : elle harmonise uniquement les
    espacements, l'habillage du logo des headers et les proprietes de pagination.
    """
    from docx import Document as _Doc
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    _corriger_habillage_logo_cgv_ponctuel(docx_path)

    d = _Doc(str(docx_path))

    def norm(txt: str) -> str:
        import unicodedata
        txt = (txt or "").lower().replace("\xa0", " ")
        txt = unicodedata.normalize("NFKD", txt)
        txt = "".join(ch for ch in txt if not unicodedata.combining(ch))
        txt = re.sub(r"[\W_]+", " ", txt, flags=re.UNICODE)
        return re.sub(r"\s+", " ", txt).strip()

    def body_paragraphs():
        return list(d.paragraphs)

    def contains_media(paragraph) -> bool:
        return bool(
            paragraph._element.xpath(
                ".//*[local-name()='drawing' or local-name()='pict' or local-name()='shape' or local-name()='imagedata']"
            )
        )

    def has_section_properties(paragraph) -> bool:
        ppr = paragraph._p.find(qn("w:pPr"))
        return ppr is not None and ppr.find(qn("w:sectPr")) is not None

    def remove_paragraph(paragraph) -> None:
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)

    def remove_child(ppr, tag_name: str) -> None:
        for node in list(ppr.findall(qn(f"w:{tag_name}"))):
            ppr.remove(node)

    def set_flag(ppr, tag_name: str, enabled: bool) -> None:
        remove_child(ppr, tag_name)
        if enabled:
            ppr.append(OxmlElement(f"w:{tag_name}"))

    def set_widow_control(ppr) -> None:
        if ppr.find(qn("w:widowControl")) is None:
            ppr.append(OxmlElement("w:widowControl"))

    def remove_hard_page_breaks(paragraph) -> None:
        for br in list(paragraph._element.xpath(".//w:br[@w:type='page']")):
            parent = br.getparent()
            if parent is not None:
                parent.remove(br)

    cgv_start = None
    for idx, p in enumerate(body_paragraphs()):
        n = norm(p.text)
        if n.startswith("conditions generales de vente") or n == "conditions generales de vente":
            cgv_start = idx
            break
    if cgv_start is None:
        return

    # Les paragraphes vides du modele produisent des trous de plusieurs lignes en
    # PDF. On les remplace par des espacements explicites sur les vrais titres.
    for p in list(body_paragraphs()[cgv_start + 1:]):
        if (p.text or "").strip():
            continue
        if contains_media(p) or has_section_properties(p):
            continue
        remove_paragraph(p)

    paragraphs = body_paragraphs()
    for idx, p in enumerate(paragraphs):
        if idx < cgv_start:
            continue
        txt = (p.text or "").strip()
        n = norm(txt)
        if not txt:
            continue

        is_cgv_title = n.startswith("conditions generales de vente")
        is_company_title = n == "societe marie eugenie"
        is_article_title = bool(re.match(r"^article\s+\d+\b", n))
        is_signature_text = "signature du client" in n or "lu et approuve" in n

        pf = p.paragraph_format
        ppr = p._p.get_or_add_pPr()
        remove_hard_page_breaks(p)
        remove_child(ppr, "keepNext")
        remove_child(ppr, "keepLines")
        if not is_cgv_title:
            remove_child(ppr, "pageBreakBefore")
        set_widow_control(ppr)

        pf.left_indent = None
        pf.right_indent = None
        pf.first_line_indent = None
        pf.line_spacing = 1.02

        if is_cgv_title:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(CONTENT_START_BEFORE_PT)
            pf.space_after = Pt(4)
            pf.keep_with_next = True
            set_flag(ppr, "pageBreakBefore", True)
            set_flag(ppr, "keepNext", True)
            set_flag(ppr, "keepLines", True)
        elif is_company_title:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(0)
            pf.space_after = Pt(13)
            pf.keep_with_next = True
            set_flag(ppr, "keepNext", True)
            set_flag(ppr, "keepLines", True)
        elif is_article_title:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.space_before = Pt(6)
            pf.space_after = Pt(2)
            pf.keep_with_next = True
            set_flag(ppr, "keepNext", True)
            set_flag(ppr, "keepLines", True)
        elif is_signature_text:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            set_flag(ppr, "keepLines", True)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.space_before = Pt(0)
            pf.space_after = Pt(1.5)
            # Garde les petits sous-articles indivisibles sans forcer de longs
            # blocs entiers a la page suivante.
            if len(txt) < 185:
                set_flag(ppr, "keepLines", True)

        for run in p.runs:
            if not run.text:
                continue
            run.font.name = "Arial"
            run.font.size = Pt(10)
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            for attr in ("ascii", "hAnsi", "cs"):
                rfonts.set(qn(f"w:{attr}"), "Arial")
            if is_cgv_title or is_company_title or is_article_title:
                run.bold = True

    for table in d.tables:
        if not any("signature du client" in norm(p.text) for row in table.rows for cell in row.cells for p in cell.paragraphs):
            continue
        for row in table.rows:
            trpr = row._tr.get_or_add_trPr()
            if trpr.find(qn("w:cantSplit")) is None:
                trpr.append(OxmlElement("w:cantSplit"))
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_together = True
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)

    d.save(str(docx_path))
    _nettoyer_doublons_zip(docx_path)


def _corriger_habillage_logo_cgv_ponctuel(docx_path: Path) -> None:
    """Empêche le corps des pages CGV de tourner autour du logo d'en-tete."""
    import tempfile
    import zipfile
    from xml.etree import ElementTree as ET

    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
        "v": "urn:schemas-microsoft-com:vml",
        "w10": "urn:schemas-microsoft-com:office:word",
    }
    ET.register_namespace("w", ns["w"])
    ET.register_namespace("r", ns["r"])
    ET.register_namespace("v", ns["v"])
    ET.register_namespace("w10", ns["w10"])

    def style_to_dict(value: str) -> dict:
        out = {}
        for part in (value or "").split(";"):
            if ":" not in part:
                continue
            key, val = part.split(":", 1)
            out[key.strip()] = val.strip()
        return out

    def pt_value(value: str):
        m = re.search(r"(-?\d+(?:[.,]\d+)?)pt", value or "")
        return float(m.group(1).replace(",", ".")) if m else None

    def header_targets(zin) -> set:
        targets = set()
        document_xml = ET.fromstring(zin.read("word/document.xml"))
        rels_xml = ET.fromstring(zin.read("word/_rels/document.xml.rels"))
        rel_map = {rel.get("Id"): rel.get("Target") for rel in rels_xml}
        for ref in document_xml.findall(".//w:sectPr/w:headerReference", ns):
            ref_type = ref.get(f"{{{ns['w']}}}type")
            if ref_type == "first":
                continue
            rel_id = ref.get(f"{{{ns['r']}}}id")
            target = rel_map.get(rel_id)
            if target:
                targets.add("word/" + target.lstrip("/"))
        return targets

    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            targets = header_targets(zin)
            if not targets:
                return
            updated_payloads = {}
            for name in targets:
                if name not in zin.namelist():
                    continue
                root = ET.fromstring(zin.read(name))
                changed = False
                for shape in root.findall(".//v:shape", ns):
                    if shape.find("./v:imagedata", ns) is None:
                        continue
                    style = style_to_dict(shape.get("style", ""))
                    width = pt_value(style.get("width", ""))
                    height = pt_value(style.get("height", ""))
                    if not (width and height and 80 <= width <= 180 and 55 <= height <= 150):
                        continue
                    wrap = shape.find("./w10:wrap", ns)
                    if wrap is None:
                        wrap = ET.Element(f"{{{ns['w10']}}}wrap")
                        shape.append(wrap)
                    if wrap.get("type") != "topAndBottom":
                        wrap.set("type", "topAndBottom")
                        wrap.set("anchorx", "margin")
                        wrap.set("anchory", "margin")
                        changed = True
                if changed:
                    updated_payloads[name] = ET.tostring(root, encoding="utf-8", xml_declaration=True)

            if not updated_payloads:
                return

            tmp = Path(tempfile.mkstemp(suffix=".docx", dir=str(docx_path.parent))[1])
            try:
                seen = set()
                with zipfile.ZipFile(docx_path, "r") as zin2, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
                    for item in zin2.infolist():
                        if item.filename in seen:
                            continue
                        seen.add(item.filename)
                        if item.filename in updated_payloads:
                            zout.writestr(item, updated_payloads[item.filename])
                        else:
                            zout.writestr(item, zin2.read(item.filename))
                shutil.copyfile(str(tmp), str(docx_path))
            finally:
                if tmp.exists():
                    tmp.unlink(missing_ok=True)
    except Exception as exc:
        print(f"[warn] habillage logo CGV ponctuel ignore : {exc}")


def _verrouiller_bloc_bon_pour_accord(docx_path: Path, date_signature: Optional[str] = None):
    """Garde le bloc Bon pour accord/signatures indivisible dans le rendu Word/PDF."""
    from docx import Document as _Doc
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    d = _Doc(str(docx_path))

    def norm(txt):
        import unicodedata
        txt = (txt or "").lower().replace("\xa0", " ")
        txt = unicodedata.normalize("NFKD", txt)
        txt = "".join(ch for ch in txt if not unicodedata.combining(ch))
        return re.sub(r"\s+", " ", txt).strip()

    def set_keep(paragraph, keep_next=False):
        ppr = paragraph._p.get_or_add_pPr()
        if ppr.find(qn("w:keepLines")) is None:
            ppr.append(OxmlElement("w:keepLines"))
        existing = ppr.find(qn("w:keepNext"))
        if keep_next and existing is None:
            ppr.append(OxmlElement("w:keepNext"))
        if not keep_next and existing is not None:
            ppr.remove(existing)

    in_block = False
    block = []
    for p in d.paragraphs:
        n = norm(p.text)
        if "bon pour accord" in n:
            in_block = True
        if in_block and n.startswith("conditions generales"):
            break
        if in_block:
            block.append(p)

    for idx, p in enumerate(block):
        n = norm(p.text)
        set_keep(p, keep_next=(idx < len(block) - 1))
        if "toutes les autres clauses contractuelles" in n:
            _remplacer_texte_paragraphe(
                p,
                "Toutes les autres clauses contractuelles sont définies dans les conditions générales de ventes \n"
                "annexées à la présente offre",
            )
        elif n.startswith("fait a marseille"):
            if date_signature:
                _remplacer_texte_paragraphe(p, f"Fait à Marseille, le {date_signature}")

    if block:
        d.save(str(docx_path))
        _nettoyer_doublons_zip(docx_path)


def _remplacer_texte_paragraphe(paragraph, text: str):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _ancrer_depart_pages_ponctuel(docx_path: Path):
    """Normalise le depart des pages ponctuelles sous le logo.

    Le logo reste reserve dans sa zone haute. Le depart est porte par le
    paragraphe titre lui-meme : pageBreakBefore + la constante
    CONTENT_START_BEFORE_PT. On n'ajoute aucun paragraphe blanc invisible.
    """
    try:
        from lxml import etree
    except Exception:
        return

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    qn_w = lambda name: f"{{{ns['w']}}}{name}"
    content_start_twips = str(CONTENT_START_BEFORE_PT * 20)
    content_after_twips = str(CONTENT_TITLE_AFTER_PT * 20)

    def norm(txt: str) -> str:
        import unicodedata
        txt = (txt or "").lower().replace("\xa0", " ")
        txt = unicodedata.normalize("NFKD", txt)
        txt = "".join(ch for ch in txt if not unicodedata.combining(ch))
        txt = re.sub(r"[\W_]+", " ", txt, flags=re.UNICODE)
        return re.sub(r"\s+", " ", txt).strip()

    def texte_para(p):
        return "".join(p.xpath(".//w:t/text()", namespaces=ns)).strip()

    def has_page_break_before(p):
        return p.find(".//w:pageBreakBefore", namespaces=ns) is not None

    def has_hard_page_break(p):
        return p.find('.//w:br[@w:type="page"]', namespaces=ns) is not None

    def set_page_break_before(p, enabled=True):
        ppr = p.find("w:pPr", namespaces=ns)
        if ppr is None:
            ppr = etree.Element(qn_w("pPr"))
            p.insert(0, ppr)
        for br in list(ppr.findall("w:pageBreakBefore", namespaces=ns)):
            ppr.remove(br)
        if enabled:
            etree.SubElement(ppr, qn_w("pageBreakBefore"))
        return ppr

    def set_spacing_zero(p):
        ppr = p.find("w:pPr", namespaces=ns)
        if ppr is None:
            ppr = etree.Element(qn_w("pPr"))
            p.insert(0, ppr)
        spacing = ppr.find("w:spacing", namespaces=ns)
        if spacing is None:
            spacing = etree.SubElement(ppr, qn_w("spacing"))
        spacing.set(qn_w("before"), "0")
        spacing.set(qn_w("after"), "0")

    def is_anchor_spacer(p):
        texts = p.xpath(".//w:t/text()", namespaces=ns)
        if "".join(texts) != " ":
            return False
        color = p.find(".//w:color", namespaces=ns)
        return color is not None and color.get(qn_w("val")) == "FFFFFF"

    def anchor_before(paragraph, spacer_count=2, force_page_break=True):
        parent = paragraph.getparent()
        if parent is None:
            return False

        previous = paragraph.getprevious()
        while previous is not None and is_anchor_spacer(previous):
            to_remove = previous
            previous = previous.getprevious()
            parent.remove(to_remove)
        if (
            previous is not None
            and not texte_para(previous)
            and (has_page_break_before(previous) or has_hard_page_break(previous))
        ):
            parent.remove(previous)

        ppr = set_page_break_before(paragraph, force_page_break)
        spacing = ppr.find("w:spacing", namespaces=ns)
        if spacing is None:
            spacing = etree.SubElement(ppr, qn_w("spacing"))
        spacing.set(qn_w("before"), content_start_twips)
        spacing.set(qn_w("after"), content_after_twips)
        return True

    tmp_path = docx_path.with_suffix(".ponctuel-anchor.tmp.docx")
    changed = False
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/document.xml":
                    root = etree.fromstring(data)
                    paragraphs = root.xpath(".//w:p", namespaces=ns)
                    in_prestations = False
                    in_cgv_flow = False
                    prestation_count = 0
                    for p in paragraphs:
                        n = norm(texte_para(p))
                        if in_cgv_flow and not n and has_hard_page_break(p):
                            parent = p.getparent()
                            if parent is not None:
                                parent.remove(p)
                                changed = True
                            continue
                        if not n:
                            continue
                        if n.startswith("1 detail des prestations"):
                            in_prestations = True
                            prestation_count = 0
                            continue
                        if in_prestations and (n.startswith("information") or n.startswith("2 proposition financiere")):
                            in_prestations = False
                        if in_prestations:
                            prestation_count += 1
                            if prestation_count > 24 and (prestation_count - 1) % 24 == 0:
                                changed = anchor_before(p, spacer_count=1, force_page_break=True) or changed
                        is_bon_pour_accord = "bon pour accord" in n
                        is_cgv_title = n.startswith("conditions generales")
                        if is_cgv_title:
                            in_cgv_flow = True
                        if is_bon_pour_accord:
                            changed = anchor_before(p, spacer_count=1, force_page_break=True) or changed
                        elif is_cgv_title:
                            changed = anchor_before(p, spacer_count=1, force_page_break=True) or changed
                    if changed:
                        data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
                zout.writestr(info, data)
    if changed:
        shutil.move(tmp_path, docx_path)
    else:
        tmp_path.unlink(missing_ok=True)


def _supprimer_rouge_document(docx_path: Path):
    """Supprime les traits, formes, bordures et styles rouges non desires."""
    _nettoyer_xml_mise_en_page(docx_path, supprimer_rouge=True)
    _nettoyer_images_entete(docx_path)


def _nettoyer_xml_mise_en_page(docx_path: Path, supprimer_rouge: bool = False):
    """Nettoyage XML final, y compris zones de texte/VML/drawing peu exposes."""
    tmp_path = docx_path.with_suffix(".layout.tmp.docx")
    red_values = (
        "FF0000", "ff0000", "#FF0000", "#ff0000",
        "F00", "f00", "#F00", "#f00",
        "E61B1B", "e61b1b", "#E61B1B", "#e61b1b",
        "E00000", "e00000", "#E00000", "#e00000",
        "E31B23", "e31b23", "D00000", "d00000",
        "C00000", "c00000", "C000000", "c000000",
        "B42318", "b42318", "red", "RED", "rouge",
    )
    # Ne pas utiliser le mot nu "red" dans les regex larges : il apparaît dans
    # des textes XML Word ordinaires (lastRenderedPageBreak, redevable, etc.) et
    # ferait supprimer des formes/images VML qui n'ont aucune couleur rouge.
    red_re = r"(?:#?)(?:FF0000|F00|E61B1B|E00000|E31B23|D00000|C00000|B42318)"

    def clean_xml(text: str) -> str:
        text = re.sub(r'\s+w:lineRule="[^"]*"', "", text)
        text = re.sub(r'\s+w:line="[^"]*"', "", text)
        if not supprimer_rouge:
            return text

        # Les traits rouges du modele peuvent etre de simples bordures de
        # paragraphe dans les headers, ou des shapes VML/DrawingML. On supprime
        # l'objet graphique complet quand il porte une couleur rouge connue.
        text = re.sub(
            rf"<w:pBdr\b(?=[\s\S]*?{red_re})[\s\S]*?</w:pBdr>",
            "",
            text,
            flags=re.I,
        )
        text = re.sub(
            rf"<w:tblBorders\b(?=[\s\S]*?{red_re})[\s\S]*?</w:tblBorders>",
            "",
            text,
            flags=re.I,
        )
        text = re.sub(
            rf"<w:(?:top|left|bottom|right|insideH|insideV)\b(?=[^>]{{0,300}}{red_re})[^>]*/>",
            "",
            text,
            flags=re.I,
        )
        text = re.sub(
            rf"<v:(?:line|rect|shape|oval)\b(?=[\s\S]*?{red_re})[\s\S]*?</v:(?:line|rect|shape|oval)>",
            "",
            text,
            flags=re.I,
        )
        text = re.sub(
            rf"<v:(?:line|rect|shape|oval)\b(?=[^>]{{0,800}}{red_re})[^>]*/>",
            "",
            text,
            flags=re.I,
        )
        text = re.sub(
            rf"<a:ln\b(?=[\s\S]{{0,1400}}{red_re})[\s\S]*?</a:ln>",
            "",
            text,
            flags=re.I,
        )
        for color in red_values:
            text = text.replace(f'w:val="{color}"', 'w:val="000000"')
            text = text.replace(f'w:color="{color}"', 'w:color="auto"')
            text = text.replace(f'w:fill="{color}"', 'w:fill="auto"')
            text = text.replace(f'color="{color}"', 'color="auto"')
            text = text.replace(f'fillcolor="{color}"', 'fillcolor="auto"')
            text = text.replace(f'strokecolor="{color}"', 'strokecolor="auto"')
            text = re.sub(r'<w:(?:top|left|bottom|right|insideH|insideV)[^>]+' + re.escape(color) + r'[^>]*/>', '', text)
            text = re.sub(r'<v:(?:line|rect|shape|oval)[^>]+' + re.escape(color) + r'[^>]*/>', '', text, flags=re.I)
            text = re.sub(r'<v:(?:line|rect|shape|oval)[^>]+' + re.escape(color) + r'[^>]*>.*?</v:(?:line|rect|shape|oval)>', '', text, flags=re.I | re.S)
            text = re.sub(r'<a:ln\b(?=[\s\S]{0,900}' + re.escape(color) + r')[\s\S]*?</a:ln>', '', text, flags=re.I)
        text = re.sub(r'<w:p\b(?=[\s\S]{0,900}<w:color w:val="000000"[\s\S]{0,900}</w:p>)(?=[\s\S]{0,900}<w:t/>)[\s\S]*?</w:p>', '', text)
        return text

    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                    data = clean_xml(data.decode("utf-8", errors="ignore")).encode("utf-8")
                zout.writestr(info, data)
    shutil.move(tmp_path, docx_path)


def _nettoyer_images_entete(docx_path: Path):
    """Efface les longues regles graphiques dans les bandeaux d'en-tete.

    Certains modeles contiennent un petit PNG d'en-tete avec des filets
    horizontaux. On conserve le texte central du bandeau et on blanchit seulement
    les longs traits noirs/gris sur les cotes.
    """
    try:
        from PIL import Image
        import io
    except Exception:
        return

    tmp_path = docx_path.with_suffix(".media.tmp.docx")
    changed = False
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename.startswith("word/media/") and info.filename.lower().endswith((".png", ".jpg", ".jpeg")):
                    try:
                        img = Image.open(io.BytesIO(data)).convert("RGBA")
                        w, h = img.size
                        if w >= 700 and h <= 180:
                            px = img.load()
                            altered = False
                            for y in range(h):
                                dark = [x for x in range(w) if px[x, y][3] and px[x, y][0] < 130 and px[x, y][1] < 130 and px[x, y][2] < 130]
                                if len(dark) < int(w * 0.18):
                                    continue
                                runs = []
                                start = prev = dark[0]
                                for x in dark[1:]:
                                    if x == prev + 1:
                                        prev = x
                                    else:
                                        runs.append((start, prev))
                                        start = prev = x
                                runs.append((start, prev))
                                for x1, x2 in runs:
                                    if x2 - x1 < 80:
                                        continue
                                    for x in range(x1, x2 + 1):
                                        if x < int(w * 0.42) or x > int(w * 0.82):
                                            px[x, y] = (255, 255, 255, px[x, y][3])
                                            altered = True
                            if altered:
                                out = io.BytesIO()
                                fmt = "PNG" if info.filename.lower().endswith(".png") else "JPEG"
                                if fmt == "JPEG":
                                    img = img.convert("RGB")
                                img.save(out, format=fmt)
                                data = out.getvalue()
                                changed = True
                    except Exception:
                        pass
                zout.writestr(info, data)
    if changed:
        shutil.move(tmp_path, docx_path)
    else:
        tmp_path.unlink(missing_ok=True)


def _compacter_paragraphes_vides(docx_path: Path, max_consecutifs: int = 0):
    """Supprime les suites de paragraphes vides pour eviter les grands blancs."""
    from docx import Document as _Doc
    from docx.oxml.ns import qn
    d = _Doc(str(docx_path))
    body = d.element.body

    def est_vide(p):
        for r in p.findall(qn("w:r")):
            if (r.findall(qn("w:t")) or r.findall(qn("w:drawing"))
                    or r.findall(qn("w:br")) or r.findall(qn("w:pict"))):
                return False
        return True

    consec = 0
    a_suppr = []
    for child in list(body):
        if child.tag == qn("w:p"):
            if est_vide(child):
                consec += 1
                if consec > max_consecutifs:
                    a_suppr.append(child)
            else:
                consec = 0
        elif child.tag == qn("w:tbl"):
            consec = 0
    for p_el in a_suppr:
        body.remove(p_el)
    if a_suppr:
        d.save(str(docx_path))


def _docx_contient_marqueur(docx_path: Path, marqueur: str) -> bool:
    try:
        d = Document(str(docx_path))
        def iter_paragraphs(container):
            for p in container.paragraphs:
                yield p
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from iter_paragraphs(cell)
        return any(marqueur in (p.text or "") for p in iter_paragraphs(d))
    except Exception:
        return False


def _restaurer_marqueur_materiel_si_absent(docx_path: Path) -> None:
    """Conserve le marqueur officiel s'il a été retiré par le bloc dynamique.

    Cette fonction n'est appelée que lorsque le template source contenait déjà
    @@MATERIEL_ENCART@@. Elle ne crée donc aucun emplacement arbitraire pour les
    templates non annotés.
    """
    MARQUEUR = "@@MATERIEL_ENCART@@"
    if _docx_contient_marqueur(docx_path, MARQUEUR):
        return
    from docx.oxml import OxmlElement

    d = Document(str(docx_path))

    def norm(txt):
        import unicodedata
        txt = (txt or "").lower().replace("\xa0", " ")
        txt = unicodedata.normalize("NFKD", txt)
        txt = "".join(ch for ch in txt if not unicodedata.combining(ch))
        txt = re.sub(r"[\W_]+", " ", txt, flags=re.UNICODE)
        return re.sub(r"\s+", " ", txt).strip()

    target = None
    for p in d.paragraphs:
        if "proposition financiere" in norm(p.text):
            target = p
            break
    if target is None:
        print(f"[warn] marqueur {MARQUEUR} retiré mais proposition financière introuvable dans {docx_path.name}")
        d.save(str(docx_path))
        return

    new_p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = MARQUEUR
    r.append(t)
    new_p.append(r)
    target._p.addprevious(new_p)
    d.save(str(docx_path))


def _injecter_encart_materiel(docx_path: Path, materiels: list, mode_ponctuel: bool = False):
    """Insère l'encart matériel uniquement sur le marqueur @@MATERIEL_ENCART@@.

    Aucun fallback n'est volontairement prévu : si le template ne contient pas
    le marqueur, la mise en page reste intacte et un avertissement est journalisé.
    """
    MARQUEUR = "@@MATERIEL_ENCART@@"
    if not materiels:
        materiels = []
    try:
        from docx import Document as _Doc
        from docx.shared import Cm, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception as exc:
        print(f"[warn] encart matériel impossible : import python-docx échoué ({exc})")
        return

    doc = _Doc(str(docx_path))

    def iter_paragraphs(container):
        for p in container.paragraphs:
            yield p
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_paragraphs(cell)

    def norm(txt):
        import unicodedata
        txt = (txt or "").lower().replace("\xa0", " ")
        txt = unicodedata.normalize("NFKD", txt)
        txt = "".join(ch for ch in txt if not unicodedata.combining(ch))
        txt = re.sub(r"[\W_]+", " ", txt, flags=re.UNICODE)
        return re.sub(r"\s+", " ", txt).strip()

    def creer_marqueur_fallback_ponctuel():
        if not mode_ponctuel or not materiels:
            return None
        target_before = None
        for p in doc.paragraphs:
            n = norm(p.text)
            if "proposition financiere" in n or "bon pour accord" in n or "conditions generales" in n:
                target_before = p
                break
        if target_before is None:
            return None
        new_p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = MARQUEUR
        r.append(t)
        new_p.append(r)
        target_before._p.addprevious(new_p)
        for p in doc.paragraphs:
            if p._p is new_p:
                return p
        return None

    target = None
    for p in iter_paragraphs(doc):
        if MARQUEUR in (p.text or ""):
            target = p
            break

    if target is None:
        target = creer_marqueur_fallback_ponctuel()

    if target is None:
        if materiels:
            print(f"[warn] marqueur {MARQUEUR} absent dans {docx_path.name} : encart matériel non injecté")
        return

    def remove_paragraph(paragraph):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    def body_text(child) -> str:
        return "".join(node.text or "" for node in child.iter(qn("w:t")))

    def child_has_media(child) -> bool:
        return bool(child.xpath(".//*[local-name()='drawing' or local-name()='pict' or local-name()='shape' or local-name()='imagedata']"))

    def child_has_page_break(child) -> bool:
        if child.find(".//" + qn("w:pageBreakBefore")) is not None:
            return True
        for br in child.iter(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
        return child.find(".//" + qn("w:sectPr")) is not None

    def nettoyer_blancs_avant_finance_ponctuel():
        if not mode_ponctuel:
            return
        finance = None
        for child in list(doc.element.body):
            n = norm(body_text(child))
            if "proposition financiere" in n or "conditions de reglement" in n:
                finance = child
                break
        if finance is None:
            return
        current = finance.getprevious()
        while current is not None and current.tag == qn("w:p"):
            text = body_text(current).strip()
            compact = re.sub(r"\s+", " ", text).strip()
            marker = compact in {MARQUEUR, "PHOTO", "PHOTOS", "Photo", "Photos", "photo", "photos"}
            if (text and not marker) or child_has_media(current):
                break
            nxt = current.getprevious()
            parent = current.getparent()
            if parent is not None:
                parent.remove(current)
            current = nxt

    if not materiels:
        remove_paragraph(target)
        nettoyer_blancs_avant_finance_ponctuel()
        doc.save(str(docx_path))
        return

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(2.8)
    table.columns[1].width = Cm(12.4)

    def set_cell_border(cell, color="b08b5c"):
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ("top", "left", "bottom", "right"):
            tag = "w:" + edge
            node = borders.find(qn(tag))
            if node is None:
                node = OxmlElement(tag)
                borders.append(node)
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "6")
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), color)

    def set_table_fixed_width(tbl):
        tbl.autofit = False
        tbl_pr = tbl._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), "10000")
        tbl_w.set(qn("w:type"), "pct")
        layout = tbl_pr.first_child_found_in("w:tblLayout")
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")

    def keep_row_together(row, repeat_header=False):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if repeat_header and tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))

    def text_from_child(child) -> str:
        return "".join(node.text or "" for node in child.iter(qn("w:t")))

    def child_starts_new_page(child) -> bool:
        if child.find(".//" + qn("w:pageBreakBefore")) is not None:
            return True
        for br in child.iter(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
        return child.find(".//" + qn("w:sectPr")) is not None

    def spacing_pt(child) -> float:
        spacing = child.find(".//" + qn("w:spacing"))
        if spacing is None:
            return 0.0
        total = 0.0
        for attr in ("before", "after"):
            raw = spacing.get(qn(f"w:{attr}"))
            if raw and raw.isdigit():
                total += int(raw) / 20.0
        return total

    def estimate_child_height_pt(child) -> float:
        if child.tag == qn("w:tbl"):
            rows = len(child.findall(".//" + qn("w:tr")))
            media = bool(child.xpath(".//*[local-name()='drawing' or local-name()='pict' or local-name()='imagedata']"))
            # Les grilles photo ponctuelles sont rendues beaucoup plus hautes par
            # Word/LibreOffice que leur simple nombre de lignes XML. Cette valeur
            # prudente evite de demarrer le tableau materiel dans un espace trop
            # court puis de le couper entre deux pages.
            return 26.0 + rows * (150.0 if media else 30.0)
        if child.tag != qn("w:p"):
            return 0.0
        txt = text_from_child(child).strip()
        has_media = bool(child.xpath(".//*[local-name()='drawing' or local-name()='pict' or local-name()='imagedata']"))
        if has_media:
            return 95.0
        if not txt:
            return max(4.0, spacing_pt(child))
        lines = max(1, (len(txt) + 88) // 89)
        return 11.5 * lines + spacing_pt(child)

    def estimate_remaining_height_before(paragraph) -> Optional[float]:
        if paragraph._p.getparent() is not doc.element.body:
            return None
        children = list(doc.element.body)
        try:
            idx = children.index(paragraph._p)
        except ValueError:
            return None
        used = 0.0
        for child in children[:idx]:
            if child_starts_new_page(child):
                used = 0.0
            used += estimate_child_height_pt(child)
        # A4 utile approximatif avec zone logo/pied deja reserves par le modele.
        return max(0.0, 650.0 - used)

    def estimate_remaining_height_after_child(anchor_child) -> Optional[float]:
        if anchor_child is None or anchor_child.getparent() is not doc.element.body:
            return None
        children = list(doc.element.body)
        try:
            idx = children.index(anchor_child)
        except ValueError:
            return None
        used = 0.0
        for child in children[:idx + 1]:
            if child_starts_new_page(child):
                used = 0.0
            used += estimate_child_height_pt(child)
        return max(0.0, 650.0 - used)

    def required_material_height_pt(count: int, has_photos: bool) -> float:
        row_h = 58.0 if has_photos else 34.0
        return 24.0 + row_h * max(1, count)

    def page_spacer_element():
        spacer = OxmlElement("w:p")
        ppr = OxmlElement("w:pPr")
        ppr.append(OxmlElement("w:pageBreakBefore"))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), str(int(max(CONTENT_START_BEFORE_PT, 120) * 20)))
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "1")
        spacing.set(qn("w:lineRule"), "exact")
        ppr.append(spacing)
        spacer.append(ppr)
        return spacer

    def block_spacing_element(before_pt: float = 6.0, after_pt: float = 2.0):
        spacer = OxmlElement("w:p")
        ppr = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), str(int(before_pt * 20)))
        spacing.set(qn("w:after"), str(int(after_pt * 20)))
        spacing.set(qn("w:line"), "1")
        spacing.set(qn("w:lineRule"), "exact")
        ppr.append(spacing)
        spacer.append(ppr)
        return spacer

    def clear_marker_as_page_spacer(paragraph):
        for run in paragraph.runs:
            run.text = ""
        if not paragraph.runs:
            paragraph.add_run("")
        pf = paragraph.paragraph_format
        pf.page_break_before = True
        pf.space_before = Pt(CONTENT_START_BEFORE_PT)
        pf.space_after = Pt(0)
        pf.line_spacing = 1

    def ponctuel_material_anchor_child():
        if not mode_ponctuel:
            return None
        children = list(doc.element.body)
        start_idx = None
        for idx, child in enumerate(children):
            if child.tag != qn("w:p") and child.tag != qn("w:tbl"):
                continue
            if "detail des prestations" in norm(text_from_child(child)):
                start_idx = idx
                break
        if start_idx is None:
            return None

        last_content = None
        for child in children[start_idx:]:
            if child.tag not in {qn("w:p"), qn("w:tbl")}:
                break
            text = text_from_child(child).strip()
            n = norm(text)
            has_media = bool(child.xpath(".//*[local-name()='drawing' or local-name()='pict' or local-name()='shape' or local-name()='imagedata']"))
            if last_content is not None and (
                "materiel encart" in n
                or "proposition financiere" in n
                or "bon pour accord" in n
                or "conditions generales" in n
                or "conditions de reglement" in n
                or "contact " in n
                or "contact@" in text.lower()
                or n.startswith("2 ")
                or n.startswith("3 ")
            ):
                break
            if text or has_media:
                last_content = child
        return last_content

    def remove_blank_body_paragraphs_after(anchor_child):
        if anchor_child is None or anchor_child.getparent() is not doc.element.body:
            return
        current = anchor_child.getnext()
        while current is not None and current.tag == qn("w:p"):
            text = text_from_child(current).strip()
            has_media = bool(current.xpath(".//*[local-name()='drawing' or local-name()='pict' or local-name()='shape' or local-name()='imagedata']"))
            has_break = child_starts_new_page(current)
            if text or has_media or has_break:
                break
            nxt = current.getnext()
            doc.element.body.remove(current)
            current = nxt

    def count_media_in_child(child) -> int:
        if child is None:
            return 0
        return len(child.xpath(".//*[local-name()='drawing' or local-name()='pict' or local-name()='shape' or local-name()='imagedata']"))

    header_cells = table.rows[0].cells
    set_table_fixed_width(table)
    keep_row_together(table.rows[0], repeat_header=True)
    header_cells[0].merge(header_cells[1])
    hp = header_cells[0].paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run("Matériel mobilisé")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(70, 58, 44)
    set_cell_border(header_cells[0])

    for mat in materiels:
        row = table.add_row()
        keep_row_together(row)
        cells = row.cells
        cells[0].width = Cm(2.8)
        cells[1].width = Cm(12.4)
        for c in cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(c, "d6c7ad")
        photo = mat.get("photo_path")
        if photo and _est_image_logo_entreprise(photo):
            photo = None
        if photo and Path(photo).exists():
            try:
                pr = cells[0].paragraphs[0].add_run()
                pr.add_picture(str(photo), width=Cm(2.35))
            except Exception:
                cells[0].paragraphs[0].add_run("Photo")
        else:
            ph = cells[0].paragraphs[0]
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = ph.add_run("Photo\nnon renseignée")
            rr.font.size = Pt(8.5)
            rr.font.color.rgb = RGBColor(120, 110, 98)

        body = cells[1]
        p_title = body.paragraphs[0]
        r_title = p_title.add_run(str(mat.get("label") or mat.get("code") or "Matériel"))
        r_title.bold = True
        r_title.font.size = Pt(10)
        r_title.font.color.rgb = RGBColor(31, 36, 48)
        categorie = str(mat.get("categorie") or "").strip()
        if categorie:
            r_cat = p_title.add_run(f"  |  {categorie}")
            r_cat.font.size = Pt(8.5)
            r_cat.font.color.rgb = RGBColor(120, 110, 98)
        desc = str(mat.get("description") or "").strip()
        if desc:
            p_desc = body.add_paragraph()
            p_desc.paragraph_format.space_before = Pt(2)
            p_desc.paragraph_format.space_after = Pt(1)
            r_desc = p_desc.add_run(desc)
            r_desc.font.size = Pt(9)
            r_desc.font.color.rgb = RGBColor(55, 55, 55)

    has_material_photos = any(
        mat.get("photo_path")
        and not _est_image_logo_entreprise(mat.get("photo_path"))
        and Path(mat.get("photo_path")).exists()
        for mat in materiels
    )
    ponctuel_anchor = ponctuel_material_anchor_child()
    remaining = (
        estimate_remaining_height_after_child(ponctuel_anchor)
        if ponctuel_anchor is not None
        else estimate_remaining_height_before(target)
    )
    required = required_material_height_pt(len(materiels), has_material_photos)
    table_fits_one_page = required <= 510.0

    def split_material_table_elements(tbl, rows_per_chunk: int):
        rows = list(tbl._tbl.findall(qn("w:tr")))
        if len(rows) <= rows_per_chunk + 1:
            return [tbl._tbl]
        header = rows[0]
        data_rows = rows[1:]
        chunks = [data_rows[i:i + rows_per_chunk] for i in range(0, len(data_rows), rows_per_chunk)]
        elements = []
        for chunk in chunks:
            new_tbl = copy.deepcopy(tbl._tbl)
            for row_el in list(new_tbl.findall(qn("w:tr"))):
                new_tbl.remove(row_el)
            new_tbl.append(copy.deepcopy(header))
            for row_el in chunk:
                new_tbl.append(copy.deepcopy(row_el))
            elements.append(new_tbl)
        return elements

    rows_per_chunk = 5 if has_material_photos else 9
    table_elements = [table._tbl]
    if mode_ponctuel and not table_fits_one_page:
        table_elements = split_material_table_elements(table, rows_per_chunk)
        parent = table._tbl.getparent()
        if parent is not None:
            parent.remove(table._tbl)

    def insert_table_sequence_after(anchor_xml, elements, *, start_on_new_page=False, add_pre_spacing=True):
        previous = anchor_xml
        if start_on_new_page:
            spacer = page_spacer_element()
            previous.addnext(spacer)
            previous = spacer
        elif add_pre_spacing:
            pre = block_spacing_element(6.0, 1.0)
            previous.addnext(pre)
            previous = pre
        for idx, element in enumerate(elements):
            if idx > 0:
                spacer = page_spacer_element()
                previous.addnext(spacer)
                previous = spacer
            previous.addnext(element)
            previous = element
        previous.addnext(block_spacing_element(4.0, 2.0))

    move_to_next_page = False
    if remaining is not None:
        if table_fits_one_page:
            move_to_next_page = remaining < required + 20.0
        else:
            # Tableau long : on segmente le tableau. Le premier segment ne doit
            # pas etre laisse a Word s'il ne tient pas dans l'espace restant.
            first_segment_required = required_material_height_pt(
                min(len(materiels), rows_per_chunk),
                has_material_photos,
            )
            move_to_next_page = has_material_photos or remaining < first_segment_required + 20.0
    if (
        mode_ponctuel
        and ponctuel_anchor is not None
        and ponctuel_anchor.tag == qn("w:tbl")
        and count_media_in_child(ponctuel_anchor) >= 2
        and table_fits_one_page
    ):
        # Une grille de plusieurs photos consomme une hauteur que Word ajuste au
        # rendu final. Pour un tableau court, on le deplace entierement afin de
        # ne jamais laisser l'en-tete ou une premiere ligne seuls en bas de page.
        move_to_next_page = True

    if ponctuel_anchor is not None:
        remove_blank_body_paragraphs_after(ponctuel_anchor)
        insert_table_sequence_after(
            ponctuel_anchor,
            table_elements,
            start_on_new_page=move_to_next_page,
            add_pre_spacing=not move_to_next_page,
        )
        if target._p.getparent() is not None:
            remove_paragraph(target)
    elif move_to_next_page:
        clear_marker_as_page_spacer(target)
        insert_table_sequence_after(target._p, table_elements, start_on_new_page=False, add_pre_spacing=False)
    else:
        insert_table_sequence_after(target._p, table_elements, start_on_new_page=False, add_pre_spacing=True)
        remove_paragraph(target)
    doc.save(str(docx_path))


def _injecter_photos(docx_path: Path, photos: list, mode_ponctuel: bool = False):
    """Injecte les photos de devis.

    En mode ponctuel, les repÃ¨res PHOTO/PHOTOS sont des marqueurs de template :
    ils sont toujours supprimÃ©s et les photos, si prÃ©sentes, sont placÃ©es aprÃ¨s
    le dÃ©tail des prestations dans une grille centrÃ©e. Hors ponctuel, on garde
    l'ancrage par zone afin de ne pas modifier les contrats dÃ©jÃ  validÃ©s.
    """
    from docx import Document as _Doc
    from docx.shared import Mm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    MARQUEUR = "@@ZONE_PHOTOS@@"
    MARQUEURS_PHOTOS = (MARQUEUR, "PHOTOS", "Photos", "photos", "PHOTO", "Photo", "photo")
    d = _Doc(str(docx_path))

    def tous_paragraphes(doc):
        for p in doc.paragraphs:
            yield p, None
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p, cell

    def supprimer_table_marqueur_si_vide(cell) -> bool:
        if cell is None:
            return False
        node = cell._tc
        table_node = None
        while node is not None:
            if node.tag == qn("w:tbl"):
                table_node = node
                break
            node = node.getparent()
        if table_node is None:
            return False
        texts = [
            re.sub(r"\s+", " ", (t.text or "").strip())
            for t in table_node.iter(qn("w:t"))
            if (t.text or "").strip()
        ]
        has_media = bool(table_node.xpath(".//*[local-name()='drawing' or local-name()='pict' or local-name()='shape' or local-name()='imagedata']"))
        marker_texts = {m.upper() for m in MARQUEURS_PHOTOS}
        only_marker = bool(texts) and all(t.upper() in marker_texts for t in texts)
        if has_media or not only_marker:
            return False
        parent = table_node.getparent()
        if parent is not None:
            parent.remove(table_node)
            return True
        return False

    def supprimer_paragraphe(paragraph, cell=None):
        if cell is not None and mode_ponctuel and supprimer_table_marqueur_si_vide(cell):
            return
        if cell is not None:
            for run in paragraph.runs:
                run.text = ""
            if not paragraph.runs:
                paragraph.add_run("")
            return
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)

    def retirer_marqueurs():
        marqueurs = []
        for p, cell in tous_paragraphes(d):
            full_text = p.text or "".join(r.text for r in p.runs)
            compact = re.sub(r"\s+", " ", (full_text or "").strip())
            upper_compact = compact.upper()
            if MARQUEUR in full_text or (mode_ponctuel and upper_compact in {"PHOTO", "PHOTOS"}):
                marqueurs.append((p, cell, compact))
                markers_to_remove = MARQUEURS_PHOTOS if mode_ponctuel else (MARQUEUR,)
                for r in p.runs:
                    for marker in markers_to_remove:
                        r.text = r.text.replace(marker, "")
        if not marqueurs:
            return None, None, []
        return marqueurs[0][0], marqueurs[0][1], marqueurs

    def norm(txt):
        import unicodedata
        txt = (txt or "").lower().replace("\xa0", " ")
        txt = unicodedata.normalize("NFKD", txt)
        txt = "".join(ch for ch in txt if not unicodedata.combining(ch))
        txt = re.sub(r"[\W_]+", " ", txt, flags=re.UNICODE)
        return re.sub(r"\s+", " ", txt).strip()

    def trouver_ancre(libelle):
        needle = norm(libelle)
        if not needle:
            return None, None
        mots = [m for m in needle.split() if len(m) > 2]

        def score_para(p):
            hay = norm(p.text)
            if not hay or MARQUEUR in p.text:
                return 0
            score = 0
            if needle in hay:
                score = 100
            elif mots:
                score = sum(1 for m in mots[:5] if m in hay)
                if score == 1 and mots[0] in hay:
                    score = 50
            return score

        best = (None, None, 0)
        # Priorite au corps du document : les tableaux financiers peuvent
        # contenir le meme libelle, mais ils ne sont pas le bon point d'ancrage.
        for p in d.paragraphs:
            score = score_para(p)
            if score > best[2]:
                best = (p, None, score)
        if best[2] >= 1:
            return best[0], best[1]

        for p, cell in tous_paragraphes(d):
            score = score_para(p)
            if score > best[2]:
                best = (p, cell, score)
        return (best[0], best[1]) if best[2] >= 1 else (None, None)

    def set_cell_borderless(cell):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for side in ("top", "left", "bottom", "right"):
            mar = tc_mar.find(qn(f"w:{side}"))
            if mar is None:
                mar = OxmlElement(f"w:{side}")
                tc_mar.append(mar)
            mar.set(qn("w:w"), "35")
            mar.set(qn("w:type"), "dxa")
        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge = tc_borders.find(qn(f"w:{edge_name}"))
            if edge is None:
                edge = OxmlElement(f"w:{edge_name}")
                tc_borders.append(edge)
            edge.set(qn("w:val"), "nil")

    def set_table_borderless(tbl):
        tbl_pr = tbl._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), "10000")
        tbl_w.set(qn("w:type"), "pct")
        tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
        if tbl_borders is None:
            tbl_borders = OxmlElement("w:tblBorders")
            tbl_pr.append(tbl_borders)
        for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge = tbl_borders.find(qn(f"w:{edge_name}"))
            if edge is None:
                edge = OxmlElement(f"w:{edge_name}")
                tbl_borders.append(edge)
            edge.set(qn("w:val"), "nil")

    def est_titre_ou_limite(p):
        txt = (p.text or "").strip()
        if not txt:
            return False
        norm_txt = norm(txt)
        limites = (
            "conditions generales", "proposition financiere", "tarif",
            "validite", "signature", "total", "tva", "modalites",
            "presentation de la societe", "tracabilite", "habilitation",
        )
        if any(x in norm_txt for x in limites):
            return True
        if len(txt) < 90:
            try:
                if any(r.bold or r.underline for r in p.runs if r.text.strip()):
                    return True
            except Exception:
                pass
        return False

    def paragraphes_apres(ancre):
        body = d.element.body
        enfants = list(body)
        try:
            idx = enfants.index(ancre._p)
        except ValueError:
            return []
        out = []
        last_blank = False
        for child in enfants[idx + 1:]:
            if child.tag != qn("w:p"):
                break
            p = None
            for para in d.paragraphs:
                if para._p is child:
                    p = para
                    break
            if p is None:
                break
            txt = (p.text or "").strip()
            if not txt:
                if not last_blank:
                    out.append(p)
                last_blank = True
                continue
            last_blank = False
            if est_titre_ou_limite(p):
                break
            out.append(p)
            if len(out) >= 35:
                break
        return out

    def vider_cellule(cell):
        for child in list(cell._tc):
            if child.tag == qn("w:p"):
                cell._tc.remove(child)

    def append_paragraph_copy(cell, paragraph):
        cell._tc.append(copy.deepcopy(paragraph._p))

    def inserer_table_apres(tbl, ancre, cell):
        if cell is not None:
            node = cell._tc
            while node is not None and not node.tag.endswith("}tbl"):
                node = node.getparent()
            if node is not None:
                node.addnext(tbl._tbl)
                return
        ancre._p.addnext(tbl._tbl)

    def inserer_table_avant(tbl, ancre, cell):
        if cell is not None:
            node = cell._tc
            while node is not None and not node.tag.endswith("}tbl"):
                node = node.getparent()
            if node is not None:
                node.addprevious(tbl._tbl)
                return
        ancre._p.addprevious(tbl._tbl)

    def trouver_ancre_photos_ponctuel():
        in_detail = False
        last_content = None
        for p in d.paragraphs:
            n = norm(p.text)
            raw = (p.text or "").strip()
            if not in_detail and "detail des prestations" in n:
                in_detail = True
                last_content = p
                continue
            if not in_detail:
                continue
            if n and (
                "proposition financiere" in n
                or "bon pour accord" in n
                or "conditions generales" in n
                or n.startswith("2 ")
                or n.startswith("2 proposition")
                or n.startswith("information")
            ):
                break
            if raw:
                last_content = p
        return last_content

    def bloc_detail_ponctuel():
        body_children = list(d.element.body)
        paragraph_by_el = {p._p: p for p in d.paragraphs}
        start_idx = None
        for idx, child in enumerate(body_children):
            if child.tag != qn("w:p"):
                continue
            p = paragraph_by_el.get(child)
            if p is not None and "detail des prestations" in norm(p.text):
                start_idx = idx
                break
        if start_idx is None:
            return []

        bloc = []
        for child in body_children[start_idx:]:
            if child.tag != qn("w:p"):
                break
            p = paragraph_by_el.get(child)
            if p is None:
                break
            marker_elements = {mp._p for mp, _cell, _compact in marqueurs}
            if bloc and p._p in marker_elements:
                break
            raw = (p.text or "").strip()
            n = norm(raw)
            compact = re.sub(r"\s+", " ", raw)
            if bloc and (
                compact.upper() in {"PHOTO", "PHOTOS"}
                or compact == MARQUEUR
                or "materiel encart" in n
                or "proposition financiere" in n
                or "bon pour accord" in n
                or "conditions generales" in n
                or n.startswith("2 ")
            ):
                break
            bloc.append(p)
        return bloc

    def bloc_detail_trop_long(paragraphes) -> bool:
        text = "\n".join((p.text or "") for p in paragraphes)
        bullets = sum(1 for p in paragraphes if (p.text or "").lstrip().startswith(("-", "•")))
        return len(text) > 900 or bullets > 8 or len(paragraphes) > 14

    def ajouter_ligne_vide_apres_titre_cell(cell):
        if len(cell.paragraphs) < 2:
            return
        first = cell.paragraphs[0]
        second = cell.paragraphs[1]
        if not (second.text or "").strip():
            return
        blank = OxmlElement("w:p")
        ppr = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "210")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.append(spacing)
        blank.append(ppr)
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        text.text = "\u00A0"
        run.append(text)
        blank.append(run)
        first._p.addnext(blank)

    def creer_tableau_photo_unique_ponctuel(paragraphes, chemin):
        tbl = d.add_table(rows=1, cols=2)
        set_table_borderless(tbl)
        tbl.autofit = False
        row = tbl.rows[0]
        trpr = row._tr.get_or_add_trPr()
        if trpr.find(qn("w:cantSplit")) is None:
            trpr.append(OxmlElement("w:cantSplit"))
        c_txt, c_img = row.cells[0], row.cells[1]
        for cell in (c_txt, c_img):
            set_cell_borderless(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        c_txt.width = Mm(112)
        c_img.width = Mm(53)
        vider_cellule(c_txt)
        for p in paragraphes:
            append_paragraph_copy(c_txt, p)
        ajouter_ligne_vide_apres_titre_cell(c_txt)
        for p in c_txt.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1
        para = c_img.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        width_mm = 49.0
        height_max_mm = 82.0
        kwargs = {"width": Mm(width_mm)}
        try:
            from PIL import Image
            with Image.open(str(chemin)) as im:
                w_px, h_px = im.size
            if w_px and h_px and width_mm * (h_px / w_px) > height_max_mm:
                kwargs = {"height": Mm(height_max_mm)}
        except Exception:
            pass
        try:
            para.add_run().add_picture(str(chemin), **kwargs)
        except Exception as e:
            print(f"[warn] photo ponctuelle non injectee : {e}")
        return tbl

    def creer_grille_photos_ponctuel(groupes):
        chemins = []
        seen = set()
        for grp in groupes:
            for chemin in grp.get("photos", []):
                key = str(chemin)
                if key not in seen:
                    chemins.append(chemin)
                    seen.add(key)
        if not chemins:
            return None

        cols = 1 if len(chemins) == 1 else 2
        rows = (len(chemins) + cols - 1) // cols
        tbl = d.add_table(rows=rows, cols=cols)
        set_table_borderless(tbl)
        tbl.autofit = False

        for row in tbl.rows:
            trpr = row._tr.get_or_add_trPr()
            if trpr.find(qn("w:cantSplit")) is None:
                trpr.append(OxmlElement("w:cantSplit"))
            for cell in row.cells:
                set_cell_borderless(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                cell.width = Mm(165 if cols == 1 else 80)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.line_spacing = 1

        for idx, chemin in enumerate(chemins):
            cell = tbl.rows[idx // cols].cells[idx % cols]
            para = cell.paragraphs[0]
            try:
                width_mm = 128.0 if cols == 1 else 72.0
                height_max_mm = 118.0 if cols == 1 else 70.0
                kwargs = {"width": Mm(width_mm)}
                try:
                    from PIL import Image
                    with Image.open(str(chemin)) as im:
                        w_px, h_px = im.size
                    if w_px and h_px:
                        projected_h = width_mm * (h_px / w_px)
                        if projected_h > height_max_mm:
                            kwargs = {"height": Mm(height_max_mm)}
                except Exception:
                    pass
                para.add_run().add_picture(str(chemin), **kwargs)
            except Exception as e:
                print(f"[warn] photo ponctuelle non injectee : {e}")

        for idx in range(len(chemins), rows * cols):
            cell = tbl.rows[idx // cols].cells[idx % cols]
            for p in cell.paragraphs:
                for run in p.runs:
                    run.text = ""
        return tbl

    def page_start_from_anchor(paragraph):
        ppr = paragraph._p.pPr
        has_break = ppr is not None and ppr.find(qn("w:pageBreakBefore")) is not None
        before = paragraph.paragraph_format.space_before
        before_pt = before.pt if before is not None else 0
        return has_break or before_pt >= max(1, CONTENT_START_BEFORE_PT - 1)

    def media_before_anchor(paragraph):
        body = d.element.body
        enfants = list(body)
        try:
            idx = enfants.index(paragraph._p)
        except ValueError:
            return False
        for child in reversed(enfants[max(0, idx - 8):idx]):
            if child.tag != qn("w:p"):
                continue
            has_media = bool(child.xpath(".//*[local-name()='drawing' or local-name()='pict' or local-name()='shape' or local-name()='imagedata']"))
            text = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
            if has_media:
                return True
            if text:
                return False
        return False

    def inserer_table_apres_ancre(tbl, ancre, cell):
        if cell is not None or not (page_start_from_anchor(ancre) or media_before_anchor(ancre)):
            inserer_table_apres(tbl, ancre, cell)
            return

        spacer = OxmlElement("w:p")
        ppr = OxmlElement("w:pPr")
        ppr.append(OxmlElement("w:pageBreakBefore"))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), str(max(CONTENT_START_BEFORE_PT, 86) * 20))
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "1")
        spacing.set(qn("w:lineRule"), "exact")
        ppr.append(spacing)
        spacer.append(ppr)
        ancre._p.addnext(spacer)
        spacer.addnext(tbl._tbl)

    def creer_tableau(grp, paragraphes=None):
        size = str(grp.get("image_size") or "grande").strip().lower().replace(" ", "_")
        align = str(grp.get("image_align") or "droite").strip().lower().replace(" ", "_")
        try:
            width_pct = float(str(grp.get("image_width_pct") or 40).replace(",", "."))
        except Exception:
            width_pct = 40.0
        width_pct = max(20.0, min(100.0, width_pct))

        size_defaults = {
            "petite": 28.0,
            "moyenne": 40.0,
            "grande": 50.0,
            "pleine_largeur": 162.0,
            "pleine-largeur": 162.0,
        }
        full_width = size in {"pleine_largeur", "pleine-largeur"} or align in {"pleine_largeur", "pleine-largeur", "full", "centre", "centrée", "centree"}
        cols = 1 if full_width else 2
        tbl = d.add_table(rows=1, cols=cols)
        set_table_borderless(tbl)
        tbl.autofit = False
        row = tbl.rows[0]

        def remplir_texte(cell):
            set_cell_borderless(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if paragraphes:
                vider_cellule(cell)
                for p in paragraphes:
                    append_paragraph_copy(cell, p)
                if cell.paragraphs:
                    first = cell.paragraphs[0]
                    nxt = first._p.getnext()
                    is_blank = False
                    if nxt is not None and nxt.tag == qn("w:p"):
                        is_blank = not "".join(node.text or "" for node in nxt.iter(qn("w:t"))).strip()
                    if not is_blank:
                        blank = OxmlElement("w:p")
                        ppr = OxmlElement("w:pPr")
                        spacing = OxmlElement("w:spacing")
                        spacing.set(qn("w:before"), "0")
                        spacing.set(qn("w:after"), "0")
                        spacing.set(qn("w:line"), "210")
                        spacing.set(qn("w:lineRule"), "auto")
                        ppr.append(spacing)
                        blank.append(ppr)
                        run = OxmlElement("w:r")
                        text = OxmlElement("w:t")
                        text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                        text.text = "\u00A0"
                        run.append(text)
                        blank.append(run)
                        first._p.addnext(blank)
            else:
                cell.text = str(grp["libelle"])
            for p in cell.paragraphs:
                if not (p.text or "").strip():
                    if not p.runs:
                        p.add_run("\u00A0")
                    else:
                        p.runs[0].text = "\u00A0"
                        for extra in p.runs[1:]:
                            extra.text = ""
                    for run in p.runs:
                        run.font.size = Pt(10)
                p.paragraph_format.space_before = 0
                p.paragraph_format.space_after = 0
                p.paragraph_format.line_spacing = 1
                p.paragraph_format.page_break_before = False

        if cols == 1:
            c_txt = row.cells[0]
            remplir_texte(c_txt)
            c_txt.width = Mm(165)
            image_cell = c_txt
            para = image_cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            largeur = size_defaults.get(size, min(162.0, max(70.0, width_pct * 1.65)))
            hauteur_max = 95.0
        else:
            if align in {"gauche", "left"}:
                c_img, c_txt = row.cells[0], row.cells[1]
            else:
                c_txt, c_img = row.cells[0], row.cells[1]
            remplir_texte(c_txt)
            set_cell_borderless(c_img)
            c_txt.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            c_img.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            image_col_mm = max(48.0, min(68.0, 165.0 * (width_pct / 100.0)))
            text_col_mm = max(95.0, 165.0 - image_col_mm)
            c_txt.width = Mm(text_col_mm)
            c_img.width = Mm(image_col_mm)
            image_cell = c_img
            para = image_cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if align in {"gauche", "left"} else WD_ALIGN_PARAGRAPH.RIGHT
            largeur = min(image_col_mm - 4.0, size_defaults.get(size, 50.0))
            hauteur_max = 62.0 if size == "grande" else (48.0 if size == "moyenne" else 38.0)

        para.paragraph_format.space_before = 0
        para.paragraph_format.space_after = 0
        para.paragraph_format.line_spacing = 1
        n = len(grp["photos"])
        if n > 1 and not full_width:
            largeur = min(largeur, 30.0 if n == 2 else 25.0)
            hauteur_max = min(hauteur_max, 42.0 if n == 2 else 34.0)
        max_photos = min(n, 8)
        for i, chemin in enumerate(grp["photos"][:max_photos]):
            try:
                kwargs = {"width": Mm(largeur)}
                try:
                    from PIL import Image
                    with Image.open(str(chemin)) as im:
                        w_px, h_px = im.size
                    if w_px and h_px:
                        projected_h = largeur * (h_px / w_px)
                        if projected_h > hauteur_max:
                            kwargs = {"height": Mm(hauteur_max)}
                except Exception:
                    pass
                para.add_run().add_picture(str(chemin), **kwargs)
                if i < max_photos - 1:
                    if n >= 3 and i % 2 == 1:
                        para.add_run().add_break()
                    else:
                        para.add_run(" ")
            except Exception as e:
                print(f"[warn] photo non injectee : {e}")
        return tbl

    groupes = []
    for ph in photos:
        if not isinstance(ph, dict):
            continue
        chemins = ph.get("photos")
        if chemins is None and ph.get("photo_path"):
            chemins = [ph["photo_path"]]
        chemins = [
            c for c in (chemins or [])
            if c and Path(c).exists() and not _est_image_logo_entreprise(c)
        ]
        if chemins:
            groupes.append({
                "libelle": ph.get("libelle", ""),
                "photos": chemins,
                "image_size": ph.get("image_size") or "grande",
                "image_align": ph.get("image_align") or "droite",
                "image_width_pct": ph.get("image_width_pct") or 40,
            })

    def nettoyer_blancs_avant_finance_ponctuel():
        if not mode_ponctuel:
            return
        finance = None
        for child in list(d.element.body):
            n = norm("".join(node.text or "" for node in child.iter(qn("w:t"))))
            if "proposition financiere" in n or "conditions de reglement" in n:
                finance = child
                break
        if finance is None:
            return
        current = finance.getprevious()
        while current is not None and current.tag == qn("w:p"):
            text = "".join(node.text or "" for node in current.iter(qn("w:t"))).strip()
            compact = re.sub(r"\s+", " ", text)
            has_media = bool(current.xpath(".//*[local-name()='drawing' or local-name()='pict' or local-name()='shape' or local-name()='imagedata']"))
            marker = compact in {MARQUEUR, "PHOTO", "PHOTOS", "Photo", "Photos", "photo", "photos"}
            if (text and not marker) or has_media:
                break
            previous = current.getprevious()
            parent = current.getparent()
            if parent is not None:
                parent.remove(current)
            current = previous

    marqueur, marqueur_cell, marqueurs = retirer_marqueurs()
    if not groupes:
        for p, cell, compact in marqueurs:
            if compact.upper() in {"PHOTO", "PHOTOS"} or compact == MARQUEUR:
                supprimer_paragraphe(p, cell)
        nettoyer_blancs_avant_finance_ponctuel()
        d.save(str(docx_path))
        return

    if mode_ponctuel:
        chemins_uniques = []
        seen_paths = set()
        for grp in groupes:
            for chemin in grp.get("photos", []):
                key = str(chemin)
                if key not in seen_paths:
                    chemins_uniques.append(chemin)
                    seen_paths.add(key)

        if len(chemins_uniques) == 1:
            bloc = bloc_detail_ponctuel()
            if bloc and not bloc_detail_trop_long(bloc):
                tbl = creer_tableau_photo_unique_ponctuel(bloc, chemins_uniques[0])
                bloc[0]._p.addprevious(tbl._tbl)
                for p in bloc:
                    parent = p._element.getparent()
                    if parent is not None:
                        parent.remove(p._element)
                for p, cell, compact in marqueurs:
                    if compact.upper() in {"PHOTO", "PHOTOS"} or compact == MARQUEUR:
                        supprimer_paragraphe(p, cell)
                d.save(str(docx_path))
                return

        tbl = creer_grille_photos_ponctuel(groupes)
        if tbl is not None:
            if marqueur is not None:
                inserer_table_apres(tbl, marqueur, marqueur_cell)
            else:
                ancre = trouver_ancre_photos_ponctuel()
                if ancre is not None:
                    inserer_table_apres(tbl, ancre, None)
        for p, cell, compact in marqueurs:
            if compact.upper() in {"PHOTO", "PHOTOS"} or compact == MARQUEUR:
                supprimer_paragraphe(p, cell)
        d.save(str(docx_path))
        return

    repli = []
    travaux = []
    body_children = list(d.element.body)
    for grp in groupes:
        ancre, cell = trouver_ancre(grp["libelle"])
        if ancre is None:
            repli.append(grp)
            continue
        try:
            idx = body_children.index(ancre._p) if cell is None else -1
        except ValueError:
            idx = -1
        travaux.append((idx, grp, ancre, cell))

    for _idx, grp, ancre, cell in sorted(travaux, key=lambda x: x[0], reverse=True):
        a_deplacer = paragraphes_apres(ancre) if cell is None else []
        paragraphes_bloc = ([ancre] + a_deplacer) if cell is None else []
        tbl = creer_tableau(grp, paragraphes_bloc)
        inserer_table_apres_ancre(tbl, ancre, cell)
        for p in paragraphes_bloc:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)

    for grp in repli:
        if marqueur is None:
            continue
        # Si aucune ancre n'est retrouvee, on evite absolument de pousser les
        # photos apres les CGV : elles restent pres du marqueur technique.
        inserer_table_apres(creer_tableau(grp), marqueur, marqueur_cell)

    d.save(str(docx_path))


def _ajouter_signature_cgv(docx_path: Path):
    """Ajoute un encadre de signature client a la fin du document/CGV."""
    from docx import Document as _Doc
    from docx.shared import Mm, Pt
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    d = _Doc(str(docx_path))

    def iter_texts():
        for p in d.paragraphs:
            yield p.text
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p.text

    if any("Signature du client et date" in txt for txt in iter_texts()):
        return

    tbl = d.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl_pr = tbl._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "10000")
    tbl_w.set(qn("w:type"), "pct")

    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = tbl_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tbl_borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "6")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "000000")

    row = tbl.rows[0]
    row.height = Mm(34)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    cell = row.cells[0]
    cell.width = Mm(170)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name in ("top", "left", "bottom", "right"):
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "6")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "000000")
    p_sig = cell.paragraphs[0]
    p_sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sig.paragraph_format.keep_together = True
    p_sig.paragraph_format.space_before = Pt(0)
    p_sig.paragraph_format.space_after = Pt(0)
    r1 = p_sig.add_run("Signature du client et date")
    r1.bold = True
    r1.font.size = Pt(9)
    p_sig.add_run().add_break()
    r2 = p_sig.add_run('(Faire précéder de la mention « lu et approuvé »)')
    r2.font.size = Pt(8)

    d.save(str(docx_path))


import threading as _threading
_SOFFICE_LOCK = _threading.Lock()
_WARMUP_DONE = False
_WARMUP_LOCK = _threading.Lock()


def _run_soffice_pdf(docx_path: Path, pdf_dir: Path, profile: str):
    """Lance une conversion LibreOffice .docx -> .pdf dans un profil isolé."""
    return subprocess.run(
        [SOFFICE_BIN, "--headless",
         f"-env:UserInstallation=file://{profile}",
         "--convert-to", "pdf",
         "--outdir", str(pdf_dir), str(docx_path)],
        capture_output=True, text=True, timeout=90,
    )


def warmup_libreoffice() -> None:
    """Réchauffe LibreOffice une seule fois (démarrage à froid).

    La toute première conversion d'un process LibreOffice échoue de façon
    intermittente (SfxBaseModel::impl_store / Io-Abort). On paie donc ce coût
    une fois, sur un document jetable, pour que la première VRAIE conversion
    ne soit jamais celle qui essuie le démarrage à froid. Toute erreur ici est
    volontairement ignorée : même une tentative ratée « réchauffe » LibreOffice.
    """
    global _WARMUP_DONE
    if _WARMUP_DONE:
        return
    import tempfile
    with _WARMUP_LOCK:
        if _WARMUP_DONE:
            return
        try:
            from docx import Document as _Doc
            with tempfile.TemporaryDirectory(prefix="lo_warmup_") as d:
                dd = Path(d)
                src = dd / "warmup.docx"
                _doc = _Doc()
                _doc.add_paragraph("warmup")
                _doc.save(str(src))
                with _SOFFICE_LOCK:
                    with tempfile.TemporaryDirectory(prefix="lo_profile_") as profile:
                        _run_soffice_pdf(src, dd, profile)
        except Exception as exc:
            print(f"[warmup] LibreOffice non concluant (ignoré) : {exc}")
        finally:
            _WARMUP_DONE = True


def docx_to_pdf(docx_path: Path, pdf_dir: Optional[Path] = None) -> Path:
    """
    Convertit un .docx en .pdf via LibreOffice headless.

    Fiabilité (correctif) :
      - warm-up : LibreOffice est réchauffé une fois avant toute conversion,
        car la première conversion d'un process échoue de façon intermittente ;
      - retry unique : si la conversion échoue, on retente automatiquement une
        seule fois (un profil neuf par tentative) ;
      - erreur explicite : si la 2e tentative échoue aussi, on lève une
        RuntimeError contenant le message exact de LibreOffice — plus jamais de
        PDF manquant « en silence ».

    Concurrence : chaque conversion est isolée dans un profil temporaire dédié
    et sérialisée par un verrou.

    Args:
        docx_path : fichier .docx à convertir
        pdf_dir : dossier de sortie (par défaut : même dossier que le .docx)

    Returns:
        Le chemin du fichier PDF généré.
    """
    import tempfile
    pdf_dir = pdf_dir or docx_path.parent
    pdf_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = pdf_dir / (docx_path.stem + ".pdf")

    # Nettoyage défensif : un essai précédent interrompu peut laisser un verrou
    # LibreOffice « .~lock.<fichier>.pdf# » qui bloque toute réécriture du même
    # nom (SfxBaseModel::impl_store / Io-Abort). On le retire s'il traîne.
    stale_lock = pdf_dir / f".~lock.{docx_path.stem}.pdf#"
    try:
        if stale_lock.exists():
            stale_lock.unlink()
    except Exception:
        pass

    # 1) Warm-up (idempotent : une seule vraie exécution par process)
    warmup_libreoffice()

    last_error = ""
    # 2) Tentative + retry unique (2 essais au total)
    with _SOFFICE_LOCK:
        for tentative in (1, 2):
            # On convertit vers un dossier de sortie NEUF : impossible d'y trouver
            # un verrou résiduel ou un PDF préexistant qui ferait échouer l'écriture.
            with tempfile.TemporaryDirectory(prefix="lo_out_") as outdir:
                outdir_p = Path(outdir)
                with tempfile.TemporaryDirectory(prefix="lo_profile_") as profile:
                    result = _run_soffice_pdf(docx_path, outdir_p, profile)
                produced = outdir_p / (docx_path.stem + ".pdf")
                if result.returncode == 0 and produced.exists():
                    try:
                        if pdf_path.exists():
                            pdf_path.unlink()
                    except Exception:
                        pass
                    shutil.move(str(produced), str(pdf_path))
                    return pdf_path
                last_error = (result.stderr or result.stdout or "").strip()
                print(f"[docx_to_pdf] tentative {tentative}/2 échouée : {last_error[:200]}")

    # 3) Échec après retry -> erreur explicite (jamais de pdf_url null silencieux)
    raise RuntimeError(
        f"Conversion PDF échouée après 2 tentatives (LibreOffice) : "
        f"{last_error or 'aucun message renvoyé'}"
    )
