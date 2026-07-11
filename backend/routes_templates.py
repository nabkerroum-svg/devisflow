"""
Routes API pour la gestion des templates Word (back-office Modèles PDF).
"""
import json
import shutil
import uuid
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends
from fastapi.responses import FileResponse, Response
from sqlmodel import Session, select

from config import TEMPLATES_DIR
from models import Template, get_session
from template_service import analyser_template, annoter_auto, docx_to_pdf, convertir_en_docx

router = APIRouter(prefix="/templates", tags=["templates"])


def _get_template_or_404(code: str, session: Session) -> Template:
    t = session.exec(select(Template).where(Template.code == code)).first()
    if not t:
        raise HTTPException(404, f"Template '{code}' introuvable")
    return t


def _template_path_or_404(t: Template) -> Path:
    fp = TEMPLATES_DIR / t.fichier
    if not fp.exists():
        raise HTTPException(404, "Fichier Word du template introuvable")
    return fp


def _media_type_for_image(name: str) -> str:
    ext = Path(name).suffix.lower()
    if ext in {".jpg", ".jpeg"}:
        return "image/jpeg"
    if ext == ".png":
        return "image/png"
    if ext == ".webp":
        return "image/webp"
    if ext == ".gif":
        return "image/gif"
    return "application/octet-stream"


def _word_media_path(image_name: str) -> str:
    clean = Path(image_name).name
    if not clean or clean != image_name or "/" in image_name or "\\" in image_name:
        raise HTTPException(400, "Nom d'image invalide")
    return f"word/media/{clean}"


@router.get("")
def list_templates(session: Session = Depends(get_session)):
    """Liste tous les devis types (modèles). Renvoie toujours du JSON valide."""
    try:
        templates = session.exec(select(Template).order_by(Template.created_at.desc())).all()
        return [
            {
                "id": t.id, "code": t.code, "nom": t.nom, "famille": t.famille,
                "fichier": t.fichier, "type_intervention": t.type_intervention,
                "description": t.description,
                "is_default": t.is_default, "actif": t.actif,
                "variables": json.loads(t.variables) if t.variables else [],
                "created_at": t.created_at.isoformat() if t.created_at else None,
                "updated_at": t.updated_at.isoformat() if t.updated_at else None,
            } for t in templates
        ]
    except Exception as e:
        # Ne jamais renvoyer du HTML "Internal Server Error" : toujours du JSON.
        from fastapi.responses import JSONResponse
        return JSONResponse(status_code=500, content={
            "error": True,
            "message": "Erreur lors du chargement des devis types",
            "details": str(e),
        })


@router.post("/upload")
async def upload_template(
    fichier: UploadFile = File(...),
    code: str = Form(...),
    nom: str = Form(...),
    famille: str = Form("contrat"),
    type_intervention: Optional[str] = Form(None),
    annoter: bool = Form(False),
    session: Session = Depends(get_session),
):
    """
    Upload d'un nouveau template Word.

    - Si `annoter=true`, applique automatiquement la table DEFAULT_SUBSTITUTIONS
      pour remplacer les zones classiques (date, destinataire, numéro, site).
    - Sinon, prend le fichier tel quel (cas où vous l'avez déjà annoté à la main).
    """
    fname = fichier.filename.lower()
    if not (fname.endswith(".docx") or fname.endswith(".doc")):
        raise HTTPException(400, "Formats acceptés : .docx ou .doc (les .doc sont convertis automatiquement)")

    # Vérifier unicité du code
    existing = session.exec(select(Template).where(Template.code == code)).first()
    if existing:
        raise HTTPException(409, f"Un template avec le code '{code}' existe déjà")

    fichier_disque = f"{code}.docx"
    target_path = TEMPLATES_DIR / fichier_disque

    if fname.endswith(".doc"):
        # Sauver le .doc temporairement puis convertir en .docx
        tmp_doc = TEMPLATES_DIR / f"{code}_upload.doc"
        with open(tmp_doc, "wb") as f:
            shutil.copyfileobj(fichier.file, f)
        try:
            converted = convertir_en_docx(tmp_doc, TEMPLATES_DIR)
            shutil.move(str(converted), str(target_path))
        finally:
            if tmp_doc.exists():
                tmp_doc.unlink()
    else:
        with open(target_path, "wb") as f:
            shutil.copyfileobj(fichier.file, f)

    # Annoter si demandé
    if annoter:
        annoter_auto(target_path, target_path)

    # Analyser les variables
    variables = analyser_template(target_path)

    # Créer l'entrée DB
    t = Template(
        code=code, nom=nom, famille=famille,
        fichier=fichier_disque,
        type_intervention=type_intervention,
        variables=json.dumps(variables),
    )
    session.add(t)
    session.commit()
    session.refresh(t)

    return {
        "ok": True, "template_id": t.id, "code": code,
        "variables_detectees": variables,
        "fichier": fichier_disque,
    }


@router.put("/{code}")
def update_template(
    code: str,
    nom: Optional[str] = Form(None),
    famille: Optional[str] = Form(None),
    type_intervention: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    is_default: Optional[bool] = Form(None),
    actif: Optional[bool] = Form(None),
    session: Session = Depends(get_session),
):
    """Modifier les métadonnées d'un devis type."""
    t = session.exec(select(Template).where(Template.code == code)).first()
    if not t:
        raise HTTPException(404, "Template introuvable")

    if nom is not None: t.nom = nom
    if famille is not None: t.famille = famille
    if type_intervention is not None: t.type_intervention = type_intervention
    if description is not None: t.description = description
    if actif is not None: t.actif = actif

    # Si on définit comme défaut, retirer le défaut des autres
    if is_default is True:
        for other in session.exec(select(Template)).all():
            other.is_default = (other.code == code)
            session.add(other)

    from datetime import datetime as _dt
    t.updated_at = _dt.utcnow()
    session.add(t)
    session.commit()
    return {"ok": True}


@router.post("/{code}/duplicate")
def duplicate_template(
    code: str,
    nouveau_code: str = Form(...),
    nouveau_nom: str = Form(...),
    session: Session = Depends(get_session),
):
    """Duplique un devis type : copie le fichier Word + crée une nouvelle entrée.
    Le modèle original n'est PAS modifié."""
    src = session.exec(select(Template).where(Template.code == code)).first()
    if not src:
        raise HTTPException(404, f"Modèle source '{code}' introuvable")
    if session.exec(select(Template).where(Template.code == nouveau_code)).first():
        raise HTTPException(409, f"Le code '{nouveau_code}' existe déjà")

    src_path = TEMPLATES_DIR / src.fichier
    new_fichier = f"{nouveau_code}.docx"
    new_path = TEMPLATES_DIR / new_fichier
    if src_path.exists():
        shutil.copy(src_path, new_path)
    else:
        raise HTTPException(404, "Fichier Word source introuvable")

    t = Template(
        code=nouveau_code, nom=nouveau_nom, famille=src.famille,
        fichier=new_fichier, type_intervention=src.type_intervention,
        description=(src.description or "") + " (copie)",
        variables=src.variables, actif=True,
    )
    session.add(t)
    session.commit()
    session.refresh(t)
    return {"ok": True, "code": nouveau_code, "nom": nouveau_nom,
            "message": f"Modèle dupliqué depuis '{code}'. L'original est inchangé."}


@router.delete("/{code}")
def delete_template(code: str, session: Session = Depends(get_session)):
    """Supprimer un template."""
    t = session.exec(select(Template).where(Template.code == code)).first()
    if not t:
        raise HTTPException(404, "Template introuvable")
    # Supprimer le fichier disque
    fp = TEMPLATES_DIR / t.fichier
    if fp.exists():
        fp.unlink()
    session.delete(t)
    session.commit()
    return {"ok": True}


@router.get("/{code}/download")
def download_template(code: str, session: Session = Depends(get_session)):
    """Télécharger le fichier .docx du template (pour vérification ou édition)."""
    t = session.exec(select(Template).where(Template.code == code)).first()
    if not t:
        raise HTTPException(404, "Template introuvable")
    fp = TEMPLATES_DIR / t.fichier
    if not fp.exists():
        raise HTTPException(404, "Fichier disque introuvable")
    return FileResponse(fp, filename=f"{code}.docx",
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@router.get("/{code}/images")
def list_template_images(code: str, session: Session = Depends(get_session)):
    """Liste les images embarquées dans le .docx du template.

    V1 back-office : on expose les médias existants pour remplacement contrôlé,
    sans recréer la mise en page Word.
    """
    t = _get_template_or_404(code, session)
    fp = _template_path_or_404(t)
    images = []
    try:
        with zipfile.ZipFile(fp) as zin:
            for name in sorted(n for n in zin.namelist() if n.startswith("word/media/")):
                data = zin.read(name)
                width = height = None
                try:
                    from PIL import Image
                    with Image.open(BytesIO(data)) as img:
                        width, height = img.size
                except Exception:
                    pass
                short = Path(name).name
                images.append({
                    "name": short,
                    "path": name,
                    "bytes": len(data),
                    "width": width,
                    "height": height,
                    "url": f"/api/templates/{code}/images/{short}",
                })
    except zipfile.BadZipFile:
        raise HTTPException(400, "Le template n'est pas un fichier .docx valide")
    return {"ok": True, "code": code, "images": images}


@router.get("/{code}/images/{image_name}")
def get_template_image(code: str, image_name: str, session: Session = Depends(get_session)):
    """Retourne une image embarquée dans le template Word."""
    t = _get_template_or_404(code, session)
    fp = _template_path_or_404(t)
    media_path = _word_media_path(image_name)
    try:
        with zipfile.ZipFile(fp) as zin:
            if media_path not in zin.namelist():
                raise HTTPException(404, "Image introuvable dans le template")
            return Response(content=zin.read(media_path), media_type=_media_type_for_image(media_path))
    except zipfile.BadZipFile:
        raise HTTPException(400, "Le template n'est pas un fichier .docx valide")


@router.post("/{code}/images/{image_name}")
async def replace_template_image(
    code: str,
    image_name: str,
    fichier: UploadFile = File(...),
    session: Session = Depends(get_session),
):
    """Remplace une image existante du template sans modifier le reste du DOCX."""
    t = _get_template_or_404(code, session)
    fp = _template_path_or_404(t)
    media_path = _word_media_path(image_name)
    data = await fichier.read()
    if not data:
        raise HTTPException(400, "Image vide")
    upload_ext = Path(fichier.filename or "").suffix.lower()
    if upload_ext not in {".jpg", ".jpeg", ".png", ".webp"}:
        raise HTTPException(400, "Formats acceptés : JPG, JPEG, PNG ou WEBP")

    tmp_path = fp.with_suffix(".images.tmp.docx")
    backup = fp.with_suffix(".images.backup.docx")
    try:
        with zipfile.ZipFile(fp, "r") as zin:
            names = zin.namelist()
            if media_path not in names:
                raise HTTPException(404, "Image introuvable dans le template")
            shutil.copy(fp, backup)
            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for info in zin.infolist():
                    payload = data if info.filename == media_path else zin.read(info.filename)
                    zout.writestr(info, payload)
        shutil.move(tmp_path, fp)
        from datetime import datetime as _dt
        t.updated_at = _dt.utcnow()
        session.add(t)
        session.commit()
    except HTTPException:
        raise
    except zipfile.BadZipFile:
        raise HTTPException(400, "Le template n'est pas un fichier .docx valide")
    except Exception as e:
        if backup.exists():
            shutil.copy(backup, fp)
        raise HTTPException(500, f"Remplacement de l'image impossible : {e}")
    finally:
        if tmp_path.exists():
            tmp_path.unlink()

    return {"ok": True, "code": code, "image": image_name, "message": "Image du modèle remplacée."}


@router.post("/{code}/replace")
async def replace_template_file(
    code: str,
    fichier: UploadFile = File(...),
    session: Session = Depends(get_session),
):
    """Remplace le fichier Word d'un template EXISTANT par votre version corrigée.

    C'est LA fonction "Importer / remplacer le modèle" : vous éditez le .docx dans
    Word (en gardant les variables), vous l'importez ici, et tous les devis suivants
    utilisent VOTRE mise en page. Le code ne reconstruit rien — il remplit seulement
    les variables présentes dans votre fichier.
    """
    t = session.exec(select(Template).where(Template.code == code)).first()
    if not t:
        raise HTTPException(404, f"Template '{code}' introuvable")

    fname = (fichier.filename or "").lower()
    if not (fname.endswith(".docx") or fname.endswith(".doc")):
        raise HTTPException(400, "Formats acceptés : .docx ou .doc")

    target_path = TEMPLATES_DIR / t.fichier  # garde le même nom de fichier
    backup = TEMPLATES_DIR / f"{code}.backup.docx"
    # Sauvegarde de l'ancien modèle (au cas où)
    if target_path.exists():
        shutil.copy(target_path, backup)

    try:
        if fname.endswith(".doc"):
            tmp_doc = TEMPLATES_DIR / f"{code}_upload.doc"
            with open(tmp_doc, "wb") as f:
                shutil.copyfileobj(fichier.file, f)
            try:
                converted = convertir_en_docx(tmp_doc, TEMPLATES_DIR)
                shutil.move(str(converted), str(target_path))
            finally:
                if tmp_doc.exists():
                    tmp_doc.unlink()
        else:
            with open(target_path, "wb") as f:
                shutil.copyfileobj(fichier.file, f)

        # Ré-analyser les variables réellement présentes dans VOTRE fichier
        variables = analyser_template(target_path)
        t.variables = json.dumps(variables)
        from datetime import datetime as _dt
        t.updated_at = _dt.utcnow()
        session.add(t)
        session.commit()
    except Exception as e:
        # Restaurer la sauvegarde en cas d'échec
        if backup.exists():
            shutil.copy(backup, target_path)
        raise HTTPException(500, f"Échec du remplacement : {e}")

    return {
        "ok": True,
        "code": code,
        "fichier": t.fichier,
        "variables_detectees": variables,
        "message": "Modèle remplacé. Les prochains devis utiliseront votre mise en page.",
    }


@router.get("/{code}/preview")
def preview_template(code: str, session: Session = Depends(get_session)):
    """
    Aperçu PDF du template (avec marqueurs Jinja2 visibles).
    Utile pour vérifier visuellement que le template est bien annoté.
    """
    t = session.exec(select(Template).where(Template.code == code)).first()
    if not t:
        raise HTTPException(404, "Template introuvable")
    fp = TEMPLATES_DIR / t.fichier
    if not fp.exists():
        raise HTTPException(404, "Fichier disque introuvable")
    try:
        pdf_path = docx_to_pdf(fp)
        return FileResponse(pdf_path, filename=f"{code}_preview.pdf", media_type="application/pdf")
    except Exception as e:
        raise HTTPException(500, f"Conversion PDF impossible : {e}")


@router.get("/{code}/preview-sample")
def preview_template_sample(code: str, session: Session = Depends(get_session)):
    """Aperçu du devis TYPE avec des données d'exemple (pas un devis client).
    Permet de vérifier mise en page, titres, textes, prestations, sauts de page."""
    from config import GENERATED_DIR
    from template_service import generer_devis as _gen, docx_to_pdf as _pdf
    t = session.exec(select(Template).where(Template.code == code)).first()
    if not t:
        raise HTTPException(404, "Modèle introuvable")
    fp = TEMPLATES_DIR / t.fichier
    if not fp.exists():
        raise HTTPException(404, "Fichier Word introuvable")
    # Données d'exemple génériques
    sample = {
        "NUMERO_DEVIS": "EXEMPLE-001", "DATE_EMISSION": "01/01/2026",
        "DEST_LIGNE1": "Client exemple", "DEST_LIGNE2": "Société exemple",
        "DEST_LIGNE3": "1 rue Exemple", "DEST_LIGNE4": "13000 Marseille",
        "TYPE_PRESTATION": "Exemple de prestation",
        "SITE_ADRESSE": "1 rue Exemple", "SITE_CP_VILLE": "13000 Marseille",
        "DATE_PRISE_EFFET": "01/01/2026", "DUREE_CONTRAT": "12 mois",
        "FREQ_HALL": "2 fois par semaine", "FREQ_ASCENSEUR": "1 fois par semaine",
        "FREQ_ESCALIERS": "1 fois par semaine", "FREQ_CAVES": "1 fois par mois",
        "FREQ_GARAGE": "1 fois par mois", "FREQ_ABORDS": "1 fois par semaine",
        "FREQ_CONTENEUR": "1 fois par semaine", "FREQ_OM": "1 fois par semaine",
        "SHOW_HALL": True, "SHOW_ASCENSEUR": True, "SHOW_ESCALIERS": True,
        "SHOW_CAVES": False, "SHOW_GARAGE": False, "SHOW_ABORDS": False,
        "SHOW_CONTENEUR": False, "SHOW_OM": False,
        "OPTIONS": [
            {"libelle": "Hall d'entrée — 2 fois par semaine", "ht": "416,00 €", "tva": "83,20 €", "ttc": "499,20 €"},
            {"libelle": "Cabine d'ascenseur — 1 fois par semaine", "ht": "208,00 €", "tva": "41,60 €", "ttc": "249,60 €"},
        ],
        "PRESTATIONS": [{"libelle": "Prestation exemple", "detail": "Détail exemple"}],
        "FORFAIT_LIBELLE": "Forfait exemple", "FORFAIT_HT": "500,00 €",
        "FORFAIT_TVA": "100,00 €", "FORFAIT_TTC": "600,00 €",
        "NOM_OPPORTUNITE": "Exemple",
    }
    out_docx = GENERATED_DIR / f"_apercu_type_{code}.docx"
    try:
        _gen(fp, sample, out_docx)
        pdf_path = _pdf(out_docx, GENERATED_DIR)
        return FileResponse(pdf_path, filename=f"{code}_apercu_type.pdf", media_type="application/pdf")
    except Exception as e:
        raise HTTPException(500, f"Aperçu impossible : {e}")


# ============================================================
#  ÉDITEUR VISUEL DU DEVIS TYPE
#  On expose les paragraphes/cellules ÉDITABLES du .docx.
#  L'édition se fait EN PLACE dans le .docx (python-docx) :
#  la mise en page (logos, ligne rouge, images, marges, sauts
#  de page, pagination) n'est JAMAIS reconstruite.
# ============================================================
import re as _re
from html.parser import HTMLParser as _HTMLParser



def _bloc_protege(texte: str) -> bool:
    """Protege uniquement les balises structurelles, pas les variables simples."""
    return bool(_re.search(r'\{%|@@', texte or ""))


def _paragraph_is_title(p) -> bool:
    txt = (p.text or "").strip()
    if not txt or "{{" in txt:
        return False
    try:
        if len(txt) < 85 and any(r.bold or r.underline for r in p.runs if r.text.strip()):
            return True
    except Exception:
        pass
    return False


def _set_paragraph_text(p, new_text):
    new_text = (new_text or "").replace("\r\n", "\n").replace("\r", "\n")
    if p.runs:
        run0 = p.runs[0]
        for r in p.runs[1:]:
            r.text = ""
    else:
        run0 = p.add_run("")
    parts = new_text.split("\n")
    run0.text = parts[0] if parts else ""
    for part in parts[1:]:
        run0.add_break()
        run0.add_text(part)


class _SimpleDocxHtmlParser(_HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.stack = [{"bold": False, "italic": False, "underline": False, "size": None, "font": None, "color": None, "highlight": None}]
        self.segments = []
        self.align = None
        self.list_stack = []
        self.ol_counters = []

    def _state(self):
        return dict(self.stack[-1])

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        attrs = dict(attrs or [])
        st = self._state()
        style = attrs.get("style", "")
        def css_value(name):
            m2 = _re.search(r"(?:^|;)\s*" + _re.escape(name) + r"\s*:\s*([^;]+)", style, _re.I)
            return m2.group(1).strip().strip("'\"") if m2 else None
        m = _re.search(r"text-align\s*:\s*(left|center|right)", style, _re.I)
        if m:
            self.align = m.group(1).lower()
        css_font = css_value("font-family")
        css_color = css_value("color")
        css_bg = css_value("background-color") or css_value("background")
        css_size = css_value("font-size")
        if css_font:
            st["font"] = css_font.split(",")[0].strip().strip("'\"")
        if css_color:
            st["color"] = css_color
        if css_bg and css_bg.lower() not in ("transparent", "inherit", "initial", "none"):
            st["highlight"] = css_bg
        if css_size:
            m_size = _re.match(r"([0-9]+(?:\.[0-9]+)?)\s*(pt|px)?", css_size, _re.I)
            if m_size:
                value = float(m_size.group(1))
                st["size"] = int(round(value * 0.75 if (m_size.group(2) or "").lower() == "px" else value))
        if tag in ("b", "strong"):
            st["bold"] = True
        elif tag in ("i", "em"):
            st["italic"] = True
        elif tag == "u":
            st["underline"] = True
        elif tag in ("h1", "h2", "h3"):
            st["bold"] = True
            st["size"] = {"h1": 16, "h2": 14, "h3": 12}[tag]
        elif tag == "font":
            size = attrs.get("size")
            st["size"] = {"1": 8, "2": 10, "3": 11, "4": 12, "5": 14, "6": 16, "7": 18}.get(str(size), st.get("size"))
            if attrs.get("face"):
                st["font"] = attrs.get("face").split(",")[0].strip().strip("'\"")
            if attrs.get("color"):
                st["color"] = attrs.get("color")
        elif tag == "br":
            self.segments.append(("\n", st))
        elif tag in ("div", "p") and self.segments and not self.segments[-1][0].endswith("\n"):
            self.segments.append(("\n", st))
        elif tag == "ul":
            self.list_stack.append("ul")
        elif tag == "ol":
            self.list_stack.append("ol")
            self.ol_counters.append(0)
        elif tag == "li":
            if self.segments and not self.segments[-1][0].endswith("\n"):
                self.segments.append(("\n", st))
            if self.list_stack and self.list_stack[-1] == "ol":
                self.ol_counters[-1] += 1
                self.segments.append((f"{self.ol_counters[-1]}. ", st))
            else:
                self.segments.append(("- ", st))
        self.stack.append(st)

    def handle_endtag(self, tag):
        tag = tag.lower()
        if len(self.stack) > 1:
            self.stack.pop()
        if tag in ("p", "div", "h1", "h2", "h3", "li"):
            st = self._state()
            if self.segments and not self.segments[-1][0].endswith("\n"):
                self.segments.append(("\n", st))
        elif tag == "ul" and self.list_stack:
            self.list_stack.pop()
        elif tag == "ol" and self.list_stack:
            self.list_stack.pop()
            if self.ol_counters:
                self.ol_counters.pop()

    def handle_data(self, data):
        if data:
            self.segments.append((data.replace("\xa0", " "), self._state()))


def _set_paragraph_html(p, html_text):
    """Applique un HTML volontairement simple en runs Word."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    parser = _SimpleDocxHtmlParser()
    parser.feed(html_text or "")
    text_flat = "".join(txt for txt, _ in parser.segments).strip("\n")
    if "<" not in (html_text or ""):
        _set_paragraph_text(p, html_text)
        return

    for child in list(p._p):
        if child.tag == qn("w:r"):
            p._p.remove(child)
    if parser.align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif parser.align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif parser.align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if not text_flat:
        p.add_run("")
        return
    def parse_hex_color(value):
        value = (value or "").strip().lower()
        names = {
            "black": "000000", "gray": "666666", "grey": "666666",
            "red": "c00000", "blue": "1f4e79", "green": "2e7d32",
            "orange": "c55a11", "yellow": "fff200", "white": "ffffff",
        }
        if value in names:
            value = names[value]
        m = _re.search(r"#?([0-9a-f]{6})", value)
        if not m:
            m3 = _re.search(r"#?([0-9a-f]{3})(?![0-9a-f])", value)
            if m3:
                short = m3.group(1)
                return "".join(ch * 2 for ch in short)
            return None
        return m.group(1)
    def highlight_index(value):
        hx = parse_hex_color(value)
        if not hx:
            return None
        r, g, b = int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16)
        if r > 230 and g > 220:
            return WD_COLOR_INDEX.YELLOW
        if g >= r and g >= b:
            return WD_COLOR_INDEX.BRIGHT_GREEN
        if b >= r and b >= g:
            return WD_COLOR_INDEX.TURQUOISE
        if r >= g and r >= b:
            return WD_COLOR_INDEX.PINK
        return WD_COLOR_INDEX.YELLOW
    for txt, st in parser.segments:
        if txt == "":
            continue
        parts = txt.split("\n")
        for i, part in enumerate(parts):
            if i:
                p.add_run().add_break()
            if part:
                r = p.add_run(part)
                r.bold = bool(st.get("bold"))
                r.italic = bool(st.get("italic"))
                r.underline = bool(st.get("underline"))
                if st.get("size"):
                    r.font.size = Pt(st["size"])
                if st.get("font"):
                    r.font.name = st["font"]
                    r_pr = r._element.get_or_add_rPr()
                    r_fonts = r_pr.find(qn("w:rFonts"))
                    if r_fonts is None:
                        r_fonts = OxmlElement("w:rFonts")
                        r_pr.append(r_fonts)
                    r_fonts.set(qn("w:ascii"), st["font"])
                    r_fonts.set(qn("w:hAnsi"), st["font"])
                color = parse_hex_color(st.get("color"))
                if color:
                    r.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
                hi = highlight_index(st.get("highlight"))
                if hi:
                    r.font.highlight_color = hi


def _build_grouped_blocks(d):
    blocs = []
    current = []

    def flush():
        nonlocal current
        if not current:
            return
        refs = [ref for ref, _ in current]
        paragraphs = [p for _, p in current]
        text = "\n".join((p.text or "").strip() for p in paragraphs if (p.text or "").strip())
        has_title = len(paragraphs) == 1 and _paragraph_is_title(paragraphs[0])
        blocs.append({
            "ref": "g:" + ",".join(refs),
            "refs": refs,
            "type": "titre" if has_title else "document",
            "texte": text,
            "editable": True,
            "protege_raison": None,
        })
        current = []

    for i, p in enumerate(d.paragraphs):
        txt = (p.text or "").strip()
        if not txt:
            continue
        ref = f"p{i}"
        if _bloc_protege(txt):
            flush()
            blocs.append({
                "ref": ref,
                "type": "structure",
                "texte": txt,
                "editable": False,
                "protege_raison": "balise technique du mod?le",
            })
            continue
        if _paragraph_is_title(p):
            flush()
            current.append((ref, p))
            flush()
            continue
        current.append((ref, p))
        if len(current) >= 10:
            flush()
    flush()

    for ti, table in enumerate(d.tables):
        for ri, row in enumerate(table.rows):
            row_items = []
            for ci, cell in enumerate(row.cells):
                txt = (cell.text or "").strip()
                if not txt:
                    continue
                ref = f"t{ti}r{ri}c{ci}"
                if _bloc_protege(txt):
                    blocs.append({
                        "ref": ref,
                        "type": "structure",
                        "texte": txt,
                        "editable": False,
                        "protege_raison": "balise technique du mod?le",
                    })
                else:
                    row_items.append((ref, txt))
            if row_items:
                blocs.append({
                    "ref": "g:" + ",".join(ref for ref, _ in row_items),
                    "refs": [ref for ref, _ in row_items],
                    "type": "tableau",
                    "texte": "\n".join(txt for _, txt in row_items),
                    "editable": True,
                    "protege_raison": None,
                })
    return blocs


def _build_editor_paragraphs(d):
    items = []
    for i, p in enumerate(d.paragraphs):
        txt = (p.text or "").strip()
        if not txt:
            continue
        items.append({
            "ref": f"p{i}",
            "type": "titre" if _paragraph_is_title(p) else "paragraphe",
            "texte": txt,
            "editable": not _bloc_protege(txt),
            "protege_raison": "balise technique du modele" if _bloc_protege(txt) else None,
        })
    for ti, table in enumerate(d.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                txt = (cell.text or "").strip()
                if not txt:
                    continue
                items.append({
                    "ref": f"t{ti}r{ri}c{ci}",
                    "type": "tableau",
                    "texte": txt,
                    "editable": not _bloc_protege(txt),
                    "protege_raison": "balise technique du modele" if _bloc_protege(txt) else None,
                })
    return items


@router.get("/{code}/contenu")
def contenu_template(code: str, session: Session = Depends(get_session)):
    """Retourne de grands blocs ?ditables, proches d'un document Word."""
    from docx import Document
    t = session.exec(select(Template).where(Template.code == code)).first()
    if not t:
        return _json_err("Mod?le introuvable", 404)
    fp = TEMPLATES_DIR / t.fichier
    if not fp.exists():
        return _json_err("Fichier Word introuvable", 404)
    try:
        d = Document(str(fp))
        return {
            "code": code,
            "nom": t.nom,
            "blocs": _build_grouped_blocks(d),
            "paragraphs": _build_editor_paragraphs(d),
        }
    except Exception as e:
        return _json_err(f"Lecture impossible : {e}", 500)


@router.post("/{code}/contenu")
async def enregistrer_contenu(code: str, payload: dict, session: Session = Depends(get_session)):
    """Enregistre les grands blocs ?ditables dans le .docx sans retirer les variables."""
    from docx import Document
    from datetime import datetime as _dt
    t = session.exec(select(Template).where(Template.code == code)).first()
    if not t:
        return _json_err("Mod?le introuvable", 404)
    fp = TEMPLATES_DIR / t.fichier
    if not fp.exists():
        return _json_err("Fichier Word introuvable", 404)

    mods = (payload or {}).get("modifications", {})
    html_mods = (payload or {}).get("html_modifications", {})
    if isinstance(html_mods, dict) and html_mods:
        mods = html_mods
    if not isinstance(mods, dict) or not mods:
        return _json_err("Aucune modification fournie", 400)

    try:
        d = Document(str(fp))
        appliquees = 0
        refusees = []

        def get_target(ref):
            m = _re.match(r"^p(\d+)$", ref)
            if m:
                idx = int(m.group(1))
                return d.paragraphs[idx] if 0 <= idx < len(d.paragraphs) else None
            m = _re.match(r"^t(\d+)r(\d+)c(\d+)$", ref)
            if m:
                ti, ri, ci = map(int, m.groups())
                try:
                    cell = d.tables[ti].rows[ri].cells[ci]
                    return cell.paragraphs[0] if cell.paragraphs else None
                except Exception:
                    return None
            return None

        def apply_group(refs, text, html_mode=False):
            targets = [get_target(r) for r in refs]
            targets = [tgt for tgt in targets if tgt is not None]
            if not targets:
                return False
            if any(_bloc_protege(tgt.text) for tgt in targets):
                return False
            if html_mode and len(targets) == 1:
                _set_paragraph_html(targets[0], text)
                return True
            lines = (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
            if not lines:
                lines = [""]
            for i, tgt in enumerate(targets):
                if i < len(targets) - 1:
                    _set_paragraph_text(tgt, lines[i] if i < len(lines) else "")
                else:
                    rest = lines[i:] if i < len(lines) else [""]
                    _set_paragraph_text(tgt, "\n".join(rest))
            return True

        for ref, new_text in mods.items():
            refs = []
            if ref.startswith("g:"):
                refs = [r for r in ref[2:].split(",") if r]
            else:
                refs = [ref]
            if apply_group(refs, new_text, html_mode=isinstance(html_mods, dict) and bool(html_mods)):
                appliquees += 1
            else:
                refusees.append(ref)

        backup = TEMPLATES_DIR / f"{code}.backup.docx"
        shutil.copy(fp, backup)
        d.save(str(fp))
        t.variables = json.dumps(analyser_template(fp))
        t.updated_at = _dt.utcnow()
        session.add(t)
        session.commit()
        return {"ok": True, "appliquees": appliquees, "refusees": refusees,
                "message": f"{appliquees} bloc(s) modifi?(s). Les prochains devis utiliseront ces textes."}
    except Exception as e:
        return _json_err(f"Enregistrement impossible : {e}", 500)


def _json_err(message, status):
    from fastapi.responses import JSONResponse
    return JSONResponse(status_code=status, content={"error": True, "message": message})
